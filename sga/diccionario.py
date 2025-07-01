# -*- coding: UTF-8 -*-
import os
from django.contrib.auth.decorators import login_required
from django.http import JsonResponse
from django.shortcuts import render
import sga.models
import sagest.models
import med.models
import bib.models
import socioecon.models
import mobile.models
from decorators import secure_module, last_access
from django.db import transaction
from settings import SITE_ROOT
from sga.commonviews import adduserdata
from sga.funciones import fields_model, field_default_value_model, generar_nombre


@login_required(redirect_field_name='ret', login_url='/loginsga')
@secure_module
@last_access
@transaction.atomic()
def view(request):
    if request.method == 'POST':
        action = request.POST['action']

        if action == 'documentacion':
            try:
                modelo = request.POST['modelo']
                from docx import Document
                from docx.shared import Inches, Pt
                document = Document()
                document.add_heading('Modelos del: %s' % modelo, 0)
                run = document.add_paragraph().add_run()
                modelos = []
                if modelo == 'sagest':
                    modelos = sagest.models.sagest_list_classes()
                    libreria = 'sagest'
                if modelo == 'sga':
                    modelos = sga.models.sga_list_classes()
                    libreria = 'sga'
                if modelo == 'bib':
                    modelos = bib.models.bib_list_classes()
                    libreria = 'bib'
                if modelo == 'mobile':
                    modelos = mobile.models.mobile_list_classes()
                    libreria = 'mobile'
                if modelo == 'sociecon':
                    modelos = socioecon.models.socioecon_list_classes()
                    libreria = 'sociecon'
                for dicmodelo in modelos:
                    hd = document.add_heading('Modelo: %s' % dicmodelo, 0)
                    hd.style.font.size = Pt(10)
                    table = document.add_table(rows=1, cols=5)
                    table.style = 'TableGrid'
                    table.style.font.size = Pt(8)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = u'Campo'
                    hdr_cells[1].text = u'Tipo'
                    hdr_cells[2].text = u'Valor por defecto'
                    hdr_cells[3].text = u'Largo'
                    hdr_cells[4].text = u'Descripción'
                    for campo in fields_model(dicmodelo, libreria):
                        nombre = ''

                        row_cells = table.add_row().cells
                        if hasattr(campo, 'attname'):
                            row_cells[0].text = campo.attname
                        if hasattr(campo, 'description'):
                            row_cells[1].text = campo.description
                        if hasattr(campo, 'default'):
                            row_cells[2].text = field_default_value_model(campo.default)
                        if hasattr(campo, 'max_length'):
                            row_cells[3].text = str(campo.max_length) if campo.max_length else ''
                        if hasattr(campo, 'verbose_name'):
                            row_cells[4].text = campo.verbose_name
                    document.add_paragraph()
                direccion = os.path.join(SITE_ROOT, 'media', 'diccionario')
                nombre = generar_nombre('diccionario', 'diccionario.docx')
                filename = os.path.join(direccion, nombre)
                doc = document.save(filename)
                return JsonResponse({"result": "ok", "documento": 'https://sga.unemi.edu.ec/media/diccionario/%s' % nombre})
            except Exception as ex:
                return JsonResponse({"result": "bad", "mensaje": u"Error al generar los datos."})

        return JsonResponse({"result": "bad", "mensaje": u"Solicitud Incorrecta."})
    else:
        data = {}
        adduserdata(request, data)
        data['title'] = 'Diccionario de datos'
        data['sgaclases'] = sga.models.sga_list_classes()
        data['sagestclases'] = sagest.models.sagest_list_classes()
        data['medclases'] = med.models.med_list_classes()
        data['bibclases'] = bib.models.bib_list_classes()
        data['socieconclases'] = socioecon.models.socioecon_list_classes()
        data['mobileclases'] = mobile.models.mobile_list_classes()
        try:
            return render(request, "diccionario/view.html", data)
        except Exception as ex:
            pass
