# -*- coding: latin-1 -*-}

import io
import sys
import json

import pandas as pd
from django.contrib import messages
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.table import  _Row

import os
from datetime import datetime, timedelta,date

import xlsxwriter
import xlwt
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User, Group
from django.forms import model_to_dict

from xlwt import *
import random
from django.db import transaction
from django.http import HttpResponseRedirect, JsonResponse, HttpResponse
from django.shortcuts import render, redirect
from django.template.loader import get_template
from django.template import Context
from xlwt import easyxf, XFStyle

from decorators import secure_module, last_access
from empleo.models import ResponsableConvenio
from sagest.funciones import formatear_cabecera_pd
from sagest.models import DistributivoPersona, DenominacionPuesto
from settings import DEFAULT_PASSWORD, SEXO_MASCULINO, EMPLEADORES_GRUPO_ID, MODELO_EVALUACION, EMAIL_DOMAIN, BASE_DIR
from sga.commonviews import adduserdata
from sga.forms import ConvenioEmpresaForm, TipoConvenioForm, TipoArchivoConvenioForm, ArchivoConvenioForm, \
    EmpleadorForm, EmpresaForm, MovilidadSolicitudForm, DepartamentoConvenioForm, ActividadConvenioForm, \
    NotificacionconvenioForm, NotificarManualForm, NotificacionconvenioIndividualForm, DetalleResponsableInternoForm
from sga.funciones import log, MiPaginador, generar_nombre, resetear_clave_empresa,notificacion
from django.db.models import Q
from sga.models import ConvenioEmpresa, TipoConvenio, TipoArchivoConvenio, ArchivoConvenio, EmpresaEmpleadora, Persona, \
    Empleador, miinstitucion, ConvenioCarrera, Carrera, CUENTAS_CORREOS, MovilidadTipoSolicitud, MovilidadBaseLegal, \
    MovilidadTipoEstancia, MESES_CHOICES, ActividadConvenio, NotificacionConvenio, DetalleNotificacionCovenio, Coordinacion, EstadoCarrera
from sga.tasks import send_html_mail, conectar_cuenta
from typing import Any, Hashable, Iterable, Optional

from utils.filtros_genericos import filtro_persona_select, consultarPersona


def buscar_dicc(it: Iterable[dict], clave: Hashable, valor: Any) -> Optional[dict]:
    for dicc in it:
        if dicc[clave] == valor:
            return dicc
    return None

@login_required(redirect_field_name='ret', login_url='/loginsga')
@secure_module
@last_access
@transaction.atomic()
def view(request):
    global ex
    data = {}
    adduserdata(request, data)
    persona = request.session['persona']
    if request.method == 'POST':
        action = request.POST['action']
        if action == 'addconvenio':
            try:
                convenio = None
                form = ConvenioEmpresaForm(request.POST, request.FILES)


                denominacionPuestos = None
                if 'lista_items1' in request.POST:
                    denominacionPuestos = json.loads(request.POST['lista_items1'])

                if form.is_valid():

                    convenio = ConvenioEmpresa(empresaempleadora=form.cleaned_data['empresaempleadora'],
                                             tipoconvenio=form.cleaned_data['tipoconvenio'],
                                             objetivo=form.cleaned_data['objetivo'],
                                             fechainicio=form.cleaned_data['fechainicio'],
                                             responsableexterno=form.cleaned_data['responsableexterno'],
                                             cargoresponsableexterno=form.cleaned_data['cargoresponsableexterno'],
                                             dias_notificacion=form.cleaned_data['dias_notificacion'],
                                             para_practicas=form.cleaned_data['para_practicas'],
                                             para_pasantias=form.cleaned_data['para_pasantias'],
                                             fechafinalizacion=form.cleaned_data['fechafinalizacion'],
                                             numocas=form.cleaned_data['numocas'],
                                             fechaocas=form.cleaned_data['fechaocas']
                                           )


                    convenio.save(request)
                    if 'archivoocas' in request.FILES:
                        arch = request.FILES['archivoocas']
                        extension = arch._name.split('.')
                        tam = len(extension)
                        exte = extension[tam - 1].lower()
                        permitidas = "doc,docx,xls,xlsx,pdf,ppt,pptx,rar,zip,txt"

                        if permitidas.find(exte) < 0:
                            return JsonResponse({"result": "bad", "mensaje": u"Solo se permiten archivos %s" % (permitidas)})

                        if arch.size > 4194304:
                            return JsonResponse({"result": "bad", "mensaje": u"Error, el tamaño del archivo es mayor a 4 Mb."})
                        arch._name = generar_nombre(f"archivo_ocas_{convenio.id}", arch._name)
                        convenio.archivoocas = arch
                        convenio.save()

                    # convenio.departamentoresponsable=form.cleaned_data['departamentoresponsable']

                    if denominacionPuestos:
                        for id_denominacionPuesto in denominacionPuestos:
                            distributivo = DistributivoPersona.objects.get(pk=id_denominacionPuesto['id'], status=True)
                            responsable_convenio = ResponsableConvenio(convenio=convenio,persona=distributivo.persona, cargo=distributivo.denominacionpuesto)
                            responsable_convenio.save(request)
                            log('Adiciono responsable de convenio: %s' % responsable_convenio, request, 'add')

                            mensaje = 'Estimado/a %s \n %s ' \
                                      'Con base a la Resolución OCAS %s de fecha %s,informo a usted que ha sido' \
                                      ' designado como Coordinador interno del %s entre %s y la UNIVERSIDAD ESTATAL DE MILAGRO ' \
                                      'cuya vigencia es de %s al %s.\n ' \
                                      'De igual manera le recordamos que como Coordinador interno será encargado de ' \
                                      'supervisar la ejecución de los acuerdos de compromisos, por cuanto se solicita de carácter ' \
                                      'indispensable registrar periódicamente las actividades del convenio que se encuentre ' \
                                      'realizando, se sugiere registrar al menos una actividad al mes y en casos excepcionales puede ser bimensual.\n ' \
                                      'La Dirección de Relaciones Interinstitucionales realizará un informe semestralmente ' \
                                      'compendio donde se evidenciará la ejecución de todo lo informado por usted como ' \
                                      'Coordinador/a de cada acuerdo de compromisos.\n  ' \
                                      'En caso de no poder asistir a las reuniones de trabajo convocadas, tiene la potestad ' \
                                      'administrativa de designar a un(a) funcionario(a) y/o servidor(a) para que lo represente en la ' \
                                      'actividad convocada debiendo presentando un documento que avale su presencia.\n  ' \
                                      'Solicitamos a usted ingresar al sistema de Gestión Técnica Académica (SGA) para ' \
                                      'mayor información del Convenio.' % (distributivo.persona, distributivo.denominacionpuesto.descripcion, convenio.numocas, convenio.fechaocas,
                                                                           convenio.tipoconvenio, convenio.empresaempleadora.nombre, convenio.fechainicio, convenio.fechafinalizacion)

                            notificacion('Responsable de convenio',
                                         mensaje, distributivo.persona, None, 'dir_convenios',
                                         convenio.id, 3, 'sga', ConvenioEmpresa, request)

                            enviar_email_convenio_departamento( distributivo.persona, convenio,distributivo.denominacionpuesto.descripcion)


                    log(u'Adiciono convenio con empresa : %s[%s]' % (convenio,convenio.id), request, "add")
                if not convenio is None:
                    if request.FILES:
                        form2 = ArchivoConvenioForm(request.POST, request.FILES)
                        if 'archivo' in request.FILES:
                            arch = request.FILES['archivo']
                            extension = arch._name.split('.')
                            tam = len(extension)
                            exte = extension[tam - 1].lower()
                            permitidas = "doc,docx,xls,xlsx,pdf,ppt,pptx,rar,zip,txt"

                            if permitidas.find(exte) < 0:
                                return JsonResponse({"result": "bad", "mensaje": u"Solo se permiten archivos %s" % (permitidas)})

                        if arch.size > 4194304:
                            return JsonResponse({"result": "bad", "mensaje": u"Error, el tamaño del archivo es mayor a 4 Mb."})

                        if form2.is_valid():
                            # convenioempresa = ConvenioEmpresa.objects.get(pk=conv.id, status=True)
                            if 'archivo' in request.FILES:
                                nfile = request.FILES['archivo']
                                nfile._name = generar_nombre("archivo_convenio_", nfile._name)
                                if ArchivoConvenio.objects.filter(
                                                                # tipoarchivoconvenio=form2.cleaned_data['tipoarchivoconvenio'],
                                                                  convenioempresa=convenio, status=True).exists():
                                   return JsonResponse({"result": "bad", "mensaje": u"Ya existe un registro."})
                                archivoconvenio = ArchivoConvenio(
                                                                # tipoarchivoconvenio=form2.cleaned_data['tipoarchivoconvenio'],
                                                                  convenioempresa=convenio, archivo=nfile)
                                archivoconvenio.save(request)
                                log(u'Adiciono archivo convenio : %s[%s]' % (archivoconvenio, archivoconvenio.id), request, "add")
                else:
                    raise NameError('Error')
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": f"Error: {ex}"})

        elif action == 'editconvenio':
            try:
                denominacionPuestos = json.loads(request.POST['lista_items1'])
                convenio = None
                form = ConvenioEmpresaForm(request.POST, request.FILES)

                if form.is_valid():
                    convenio = ConvenioEmpresa.objects.get(pk=int(request.POST['id']))
                    convenio.empresaempleadora = form.cleaned_data['empresaempleadora']
                    convenio.tipoconvenio = form.cleaned_data['tipoconvenio']
                    convenio.objetivo = form.cleaned_data['objetivo']
                    convenio.fechainicio = form.cleaned_data['fechainicio']
                    convenio.responsableexterno = form.cleaned_data['responsableexterno']
                    convenio.dias_notificacion = form.cleaned_data['dias_notificacion']
                    convenio.cargoresponsableexterno = form.cleaned_data['cargoresponsableexterno']
                    # convenio.responsableinterno.id = form.cleaned_data['responsableinterno']
                    convenio.fechafinalizacion = form.cleaned_data['fechafinalizacion']
                    convenio.para_practicas = form.cleaned_data['para_practicas']
                    convenio.para_pasantias = form.cleaned_data['para_pasantias']
                    convenio.numocas = form.cleaned_data['numocas']
                    convenio.fechaocas = form.cleaned_data['fechaocas']

                    # convenio.cargo_denominaciones = form.cleaned_data['denominacionPuesto']
                    convenio.save(request)

                    if 'archivoocas' in request.FILES:
                        arch = request.FILES['archivoocas']
                        extension = arch._name.split('.')
                        tam = len(extension)
                        exte = extension[tam - 1].lower()
                        permitidas = "doc,docx,xls,xlsx,pdf,ppt,pptx,rar,zip,txt"

                        if permitidas.find(exte) < 0:
                            return JsonResponse({"result": "bad", "mensaje": u"Solo se permiten archivos %s" % (permitidas)})

                        if arch.size > 4194304:
                            return JsonResponse({"result": "bad", "mensaje": u"Error, el tamaño del archivo es mayor a 4 Mb."})
                        convenio.archivoocas = arch
                        convenio.save()
                    ids_excluir = [elem['idresponsableconvenio'] for elem in denominacionPuestos if 'idresponsableconvenio' in elem]
                    ResponsableConvenio.objects.filter(status=True, convenio=convenio).exclude(id__in=ids_excluir).update(status=False)
                    for elemento in denominacionPuestos:
                        if 'id' in elemento and not 'idresponsableconvenio' in elemento:
                            distributivo = DistributivoPersona.objects.get(pk=int(elemento['id']), status=True)
                            responsable_convenio = ResponsableConvenio(convenio=convenio, persona=distributivo.persona, cargo=distributivo.denominacionpuesto)
                            responsable_convenio.save(request)
                            log('Adiciono responsable de convenio: %s' % responsable_convenio, request, 'add')
                            if form.cleaned_data['notificar']:
                                mensaje = 'Estimado/a %s \n %s ' \
                                          'Con base a la Resolución OCAS %s de fecha %s,informo a usted que ha sido' \
                                          ' designado como Coordinador interno del %s entre %s y la UNIVERSIDAD ESTATAL DE MILAGRO ' \
                                          'cuya vigencia es de %s al %s.\n ' \
                                          'De igual manera le recordamos que como Coordinador interno será encargado de ' \
                                          'supervisar la ejecución de los acuerdos de compromisos, por cuanto se solicita de carácter ' \
                                          'indispensable registrar periódicamente las actividades del convenio que se encuentre ' \
                                          'realizando, se sugiere registrar al menos una actividad al mes y en casos excepcionales puede ser bimensual.\n ' \
                                          'La Dirección de Relaciones Interinstitucionales realizará un informe semestralmente ' \
                                          'compendio donde se evidenciará la ejecución de todo lo informado por usted como ' \
                                          'Coordinador/a de cada acuerdo de compromisos.\n  ' \
                                          'En caso de no poder asistir a las reuniones de trabajo convocadas, tiene la potestad ' \
                                          'administrativa de designar a un(a) funcionario(a) y/o servidor(a) para que lo represente en la ' \
                                          'actividad convocada debiendo presentando un documento que avale su presencia.\n  ' \
                                          'Solicitamos a usted ingresar al sistema de Gestión Técnica Académica (SGA) para ' \
                                          'mayor información del Convenio.' % (distributivo.persona, distributivo.denominacionpuesto.descripcion, convenio.numocas,convenio.fechaocas, convenio.tipoconvenio,
                                                                                       convenio.empresaempleadora.nombre, convenio.fechainicio, convenio.fechafinalizacion)

                                notificacion('Responsable de convenio',
                                             mensaje, distributivo.persona, None, 'dir_convenios',
                                             convenio.id, 3, 'sga', ConvenioEmpresa, request)

                                enviar_email_convenio_departamento(distributivo.persona, convenio,distributivo.denominacionpuesto.descripcion)

                    log(u'Edito convenio con empresa: %s[%s]' % (convenio, convenio.id),request, "edit")
                if not convenio is None:
                    form2 = ArchivoConvenioForm(request.POST, request.FILES)
                    if 'archivo' in request.FILES:
                        arch = request.FILES['archivo']
                        extension = arch._name.split('.')
                        tam = len(extension)
                        exte = extension[tam - 1].lower()
                        permitidas = "doc,docx,xls,xlsx,pdf,ppt,pptx,rar,zip,txt"

                        if permitidas.find(exte) < 0:
                            return JsonResponse(
                                {"result": "bad", "mensaje": u"Solo se permiten archivos %s" % (permitidas)})

                        if arch.size > 4194304:
                            return JsonResponse(
                                {"result": "bad", "mensaje": u"Error, el tamaño del archivo es mayor a 4 Mb."})

                    if form2.is_valid():
                        # archivoconvenio = ArchivoConvenio.objects.get(pk=int(request.POST['id']))
                        if 'archivo' in request.FILES:
                                form2 = ArchivoConvenioForm(request.POST, request.FILES)
                                arch = request.FILES['archivo']
                                extension = arch._name.split('.')
                                tam = len(extension)
                                exte = extension[tam - 1].lower()
                                permitidas = "doc,docx,xls,xlsx,pdf,ppt,pptx,rar,zip,txt"

                                if permitidas.find(exte) < 0:
                                    return JsonResponse({"result": "bad",
                                                            "mensaje": u"Solo se permiten archivos %s" % (permitidas)})

                                    if arch.size > 4194304:
                                        return JsonResponse({"result": "bad",
                                                             "mensaje": u"Error, el tamaño del archivo es mayor a 4 Mb."})

                                if form2.is_valid():
                                    # convenioempresa = ConvenioEmpresa.objects.get(pk=conv.id, status=True)
                                    nfile = request.FILES['archivo']
                                    nfile._name = generar_nombre("archivo_convenio_", nfile._name)
                                    if ArchivoConvenio.objects.filter(convenioempresa=convenio, status=True).exists():
                                        archivoconvenio = ArchivoConvenio.objects.filter(convenioempresa=convenio, status=True).first()
                                    else:
                                        archivoconvenio = ArchivoConvenio(convenioempresa=convenio)
                                    # archivoconvenio.tipoarchivoconvenio = form2.cleaned_data['tipoarchivoconvenio']
                                    archivoconvenio.archivo = nfile
                                    archivoconvenio.save(request)
                                    log(u'Adiciono archivo convenio : %s[%s]' % (
                                    archivoconvenio, archivoconvenio.id), request, "add")
                    else:
                         mensaje = [{k: v[0]} for k, v in form2.errors.items()][0]['archivo']
                         transaction.set_rollback(True)
                         return JsonResponse({"result": "bad", "mensaje": mensaje})
                else:
                    return JsonResponse({"result": "bad", "mensaje": u"Faltan campos de llenar."})


                return JsonResponse({"result": "ok"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'detalle_objetivo':
            try:
                if 'id' in request.POST:
                    data['detalleObjetivo'] = ConvenioEmpresa.objects.get(id=int(request.POST['id']))
                    template = get_template("adm_convenioempresa/modal/detalle_objetivo.html")
                    return JsonResponse({"result": 'ok', 'data': template.render(data)})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al obtener los datos."})

        elif action == 'adicionarcarrerasconvenios':
            try:
                valor = 0
                if ConvenioCarrera.objects.filter(convenioempresa_id=request.POST['convenioid'], carrera_id= request.POST['carreraid'], status=True):
                    conveniocarrera = ConvenioCarrera.objects.get(convenioempresa_id=request.POST['convenioid'], carrera_id=request.POST['carreraid'], status=True)
                    conveniocarrera.delete()
                    log(u'Elimino una carrera de convenio: %s' % conveniocarrera, request, "del")
                else:
                    conveniocarrera = ConvenioCarrera(convenioempresa_id=int(request.POST['convenioid']),
                                                      carrera_id=int(request.POST['carreraid']))
                    conveniocarrera.save(request)
                    valor = 1
                    log(u'Agrego una carrera a convenio: %s' % conveniocarrera, request, "add")
                return JsonResponse({"result": "ok", "valor": valor})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": "Error al eliminar los datos."})

        elif action == 'activartodascarreras':
            try:
                accion = request.POST['accion']
                idconvenio = request.POST['convenioid']
                if accion == 'activar':
                    convenioempresa = ConvenioEmpresa.objects.get(id=idconvenio)
                    convenioscarreras = convenioempresa.conveniocarrera_set.filter(status=True).values_list('carrera_id')
                    carreras = Carrera.objects.filter(status=True, coordinacion__isnull=False).exclude(pk__in=convenioscarreras).exclude(coordinacion__id__in=[6, 8, 9, 10]).order_by('nombre')
                    for carrera in carreras:
                        conveniocarrera = ConvenioCarrera(convenioempresa_id=int(request.POST['convenioid']),
                                                          carrera=carrera)
                        conveniocarrera.save(request)
                        log(u'Agrego una carrera a convenio: %s' % conveniocarrera, request, "add")
                else:
                    conveniocarrera = ConvenioCarrera.objects.filter(convenioempresa_id=idconvenio, status=True).update(status=False)
                    log(u'Elimino una carrera de convenio: %s' % conveniocarrera, request, "del")
                return JsonResponse({"result": "ok"}, safe=True)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": "Error al eliminar los datos."})

        elif action == 'delconvenio':
            try:
                convenio = ConvenioEmpresa.objects.get(pk=int(request.POST['id']))
                convenio.status = False
                convenio.save(request)
                log(u'Elimino convenio con empresa: %s[%s]' % (convenio, convenio.id),request, "del")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})

        elif action == 'adddepartamento':
            try:
                postar = ConvenioEmpresa.objects.get(id=int(request.POST['id']))
                form = DepartamentoConvenioForm(request.POST)
                if form.is_valid():
                    postar.departamentoresponsable = form.cleaned_data['departamentoresponsable']
                    postar.save(request)
                    for dep in postar.departamentoresponsable.all():
                        if dep.responsable:
                            mensaje = 'Se comunica que ha sido designado como responsable del convenio: %s' %(postar)
                            notificacion('Responsable de convenio',
                                         mensaje, dep.responsable, None, 'dir_convenios',
                                         postar.id, 3, 'sga', ConvenioEmpresa, request)

                            enviar_email_convenio_departamento(dep.responsable, postar, '')
                    log(u'Adiciono Departamento Responsable Convenios : %s' % (postar), request, "add")
                    return JsonResponse({"result": False,}, safe=False)
                else:
                    return JsonResponse({'result': True, "form": [{k: v[0]} for k, v in form.errors.items()], "mensaje": "Error en el formulario"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": u"Error al obtener los datos. {}".format(ex)})
        # TIPO DE CONVENIO
        elif action == 'addtipoconv':
            try:
                form = TipoConvenioForm(request.POST)
                if form.is_valid():
                    if TipoConvenio.objects.filter(nombre=form.cleaned_data['nombre'], status=True).exists():
                        return JsonResponse({"result": "bad", "mensaje": u"El nombre ya existe."})
                    tipo = TipoConvenio(nombre=form.cleaned_data['nombre'])
                    tipo.save(request)
                    log(u'Adiciono un tipo de convenio : %s[%s]' % (tipo,tipo.id), request, "add")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'edittipoconv':
            try:
                form = TipoConvenioForm(request.POST)
                if form.is_valid():
                    tipo = TipoConvenio.objects.get(pk=int(request.POST['id']))
                    tipo.nombre=form.cleaned_data['nombre']
                    tipo.save(request)
                    log(u'Edito un tipo de convenio : %s[%s]' % (tipo, tipo.id),request, "edit")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    return JsonResponse({"result": "bad", "mensaje": u"Faltan campos de llenar."})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action=='deltipoconv':
            try:
                tipo = TipoConvenio.objects.get(pk=int(request.POST['id']))
                if tipo.nopuedeeliminar():
                    return JsonResponse({"result": "bad", "mensaje": u"No puede eliminar, esta usando en otro registro."})
                tipo.delete()
                log(u'Elimino tipo de convenio : %s[%s]' % (tipo, tipo.id),request, "del")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})

        # TIPO DE ARCHIVO
        elif action == 'addtipoarch':
            try:
                form = TipoArchivoConvenioForm(request.POST)
                if form.is_valid():
                    if TipoArchivoConvenio.objects.filter(nombre=form.cleaned_data['nombre'], status=True).exists():
                        return JsonResponse({"result": "bad", "mensaje": u"El nombre ya existe."})
                    tipoarchivo = TipoArchivoConvenio(nombre=form.cleaned_data['nombre'])
                    tipoarchivo.save(request)
                    log(u'Adiciono un tipo de archivo convenio : %s[%s]' % (tipoarchivo, tipoarchivo.id), request, "add")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'edittipoarch':
            try:
                form = TipoArchivoConvenioForm(request.POST)
                if form.is_valid():
                    tipoarchivo = TipoArchivoConvenio.objects.get(pk=int(request.POST['id']))
                    tipoarchivo.nombre = form.cleaned_data['nombre']
                    tipoarchivo.save(request)
                    log(u'Edito un tipo de convenio : %s[%s]' % (tipoarchivo, tipoarchivo.id), request, "edit")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'deltipoarch':
            try:
                tipoarchivo = TipoArchivoConvenio.objects.get(pk=int(request.POST['id']))
                if tipoarchivo.nopuedeeliminar():
                    return JsonResponse({"result": "bad", "mensaje": u"No puede eliminar, esta usando en otro registro."})
                tipoarchivo.delete()
                log(u'Elimino tipo de archivo convenio : %s[%s]' % (tipoarchivo, tipoarchivo.id), request, "del")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})

        # ARCHIVO CONVENIO
        elif action == 'addarchivoconv':
            try:
                form = ArchivoConvenioForm(request.POST, request.FILES)

                arch = request.FILES['archivo']
                extension = arch._name.split('.')
                tam = len(extension)
                exte = extension[tam - 1].lower()
                permitidas = "doc,docx,xls,xlsx,pdf,ppt,pptx,rar,zip,txt"

                if permitidas.find(exte) < 0:
                    return JsonResponse({"result": "bad", "mensaje": u"Solo se permiten archivos %s" % (permitidas)})

                if arch.size > 4194304:
                    return JsonResponse({"result": "bad", "mensaje": u"Error, el tamaño del archivo es mayor a 4 Mb."})

                if form.is_valid():
                    convenioempresa = ConvenioEmpresa.objects.get(pk=int(request.POST['id']), status=True)
                    nfile = request.FILES['archivo']
                    nfile._name = generar_nombre("archivo_convenio_", nfile._name)
                    if ArchivoConvenio.objects.filter(tipoarchivoconvenio=form.cleaned_data['tipoarchivoconvenio'], convenioempresa=convenioempresa, status=True).exists():
                        return JsonResponse({"result": "bad", "mensaje": u"Ya existe un registro."})
                    archivoconvenio = ArchivoConvenio(tipoarchivoconvenio=form.cleaned_data['tipoarchivoconvenio'], convenioempresa=convenioempresa, archivo=nfile)
                    archivoconvenio.save(request)
                    log(u'Adiciono archivo convenio : %s[%s]' % (archivoconvenio, archivoconvenio.id), request, "add")
                    return JsonResponse({"result": "ok"})
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'editarchivoconv':
            try:
                form = ArchivoConvenioForm(request.POST, request.FILES)

                arch = request.FILES['archivo']
                extension = arch._name.split('.')
                tam = len(extension)
                exte = extension[tam - 1].lower()
                permitidas = "doc,docx,xls,xlsx,pdf,ppt,pptx,rar,zip,txt"

                if permitidas.find(exte) < 0:
                    return JsonResponse({"result": "bad", "mensaje": u"Solo se permiten archivos %s" % (permitidas)})

                if arch.size > 4194304:
                    return JsonResponse({"result": "bad", "mensaje": u"Error, el tamaño del archivo es mayor a 4 Mb."})

                if form.is_valid():
                    archivoconvenio = ArchivoConvenio.objects.get(pk=int(request.POST['id']))
                    if ArchivoConvenio.objects.filter(tipoarchivoconvenio=form.cleaned_data['tipoarchivoconvenio'], convenioempresa=archivoconvenio.convenioempresa, status=True).exclude(pk=archivoconvenio.id).exists():
                        return JsonResponse({"result": "bad", "mensaje": u"Ya existe un registro."})
                    if 'archivo' in request.FILES:
                        nfile = request.FILES['archivo']
                        nfile._name = generar_nombre("archivo_convenio_", nfile._name)
                        archivoconvenio.archivo = nfile
                    archivoconvenio.tipoarchivoconvenio = form.cleaned_data['tipoarchivoconvenio']
                    archivoconvenio.save(request)
                    log(u'Edito archivo de convenio : %s[%s]' % (archivoconvenio, archivoconvenio.id), request, "edit")
                    return JsonResponse({"result": "ok"})
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'delarchivoconv':
            try:
                archivoconvenio = ArchivoConvenio.objects.get(pk=int(request.POST['id']))
                archivoconvenio.delete()
                log(u'Elimino archivo de convenio : %s[%s]' % (archivoconvenio, archivoconvenio.id), request, "del")
                return JsonResponse({"error": False})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"error": "bad", "mensaje": u"Error al eliminar los datos."})

        # EMPRESA
        elif action == 'addempresa':
            try:
                form = EmpleadorForm(request.POST)
                # form.desbloquear()
                if form.is_valid():
                    ruc = form.cleaned_data['ruc'].strip()
                    if ruc:
                        if EmpresaEmpleadora.objects.filter(status=True, ruc=ruc).exists():
                            return JsonResponse({"result": "bad", "mensaje": u"Empresa ya registrada."})
                    # if (form.cleaned_data['convenio'] == True):
                    #     empresa = EmpresaEmpleadora(nombre=form.cleaned_data['nombre'],
                    #                                 ruc=ruc,
                    #                                 provincia_id=request.POST['provincia'],
                    #                                 direccion=form.cleaned_data['direccion'],
                    #                                 telefonos=form.cleaned_data['telefonos'],
                    #                                 pais=form.cleaned_data['pais'],
                    #                                 autorizada=True,
                    #                                 representante=form.cleaned_data['representante'],
                    #                                 fechainicio=form.cleaned_data['fechainicio'],
                    #                                 fechafin=form.cleaned_data['fechafin'],
                    #                                 objetivo=form.cleaned_data['observacion'],
                    #                                 # tipoconvenio_id=request.POST['tipoconvenio'],
                    #                                 # convenio=form.cleaned_data['convenio'],
                    #                                 tipoinstitucion=int(request.POST['tipoinstitucion']),
                    #                                 sectoreconomico=int(request.POST['sectoreconomico'])
                    #                                 )
                    # else:
                    empresa = EmpresaEmpleadora(nombre=form.cleaned_data['nombre'],
                                                ruc=ruc,
                                                email=form.cleaned_data['email'],
                                                provincia_id=request.POST['provincia'],
                                                direccion=form.cleaned_data['direccion'],
                                                telefonos=form.cleaned_data['telefonos'],
                                                telefonoconv=form.cleaned_data['telefonoconv'],
                                                autorizada=True,
                                                pais=form.cleaned_data['pais'],
                                                representante=form.cleaned_data['representante'],
                                                # fechainicio=form.cleaned_data['fechainicio'],
                                                # fechafin=form.cleaned_data['fechafin'],
                                                objetivo=form.cleaned_data['observacion'],
                                                # convenio=form.cleaned_data['convenio'],
                                                #tipoinstitucion=int(request.POST['tipoinstitucion']),
                                                sectoreconomico=int(request.POST['sectoreconomico']),
                                                cargo=form.cleaned_data['cargo']
                                                )
                    empresa.save(request)
                    # nombre = form.cleaned_data['contacto']
                    # nombres = ""
                    # apellido1 = ""
                    # apellido2 = ""
                    # if nombre.split().__len__() == 1:
                    #     nombres = nombre
                    # elif nombre.split().__len__() == 2:
                    #     nombres = nombre.split()[0]
                    #     apellido1 = nombre.split()[1]
                    # elif nombre.split().__len__() >= 3:
                    #     nombres = nombre.split()[0]
                    #     apellido1 = nombre.split()[1]
                    #     apellido2 = " ".join(nombre.split()[2:])
                    # persona = Persona(nombres=nombres,
                    #                   apellido1=apellido1,
                    #                   apellido2=apellido2,
                    #                   email=form.cleaned_data['email'],
                    #                   nacimiento=datetime.now().date() - timedelta(days=1),
                    #                   sexo_id=SEXO_MASCULINO)
                    # persona.save(request)
                    # username = ruc
                    # if ruc:
                    #     password = DEFAULT_PASSWORD
                    #     user = User.objects.create_user(username, "", password)
                    #     user.save()
                    #     persona.usuario = user
                    #     persona.save(request)
                    #     grupo = Group.objects.get(pk=EMPLEADORES_GRUPO_ID)
                    #     grupo.user_set.add(user)
                    #     grupo.save()
                    # empleador = Empleador(empresa_id=empresa.id,
                    #                       persona_id=persona.id,
                    #                       cargo=form.cleaned_data['cargo'])
                    # empleador.save(request)
                    # if not persona.tiene_perfil():
                    #     persona.crear_perfil(empleador=empleador)

                    # log(u'Adiciono nuevo empleador: %s' % empleador, request, "add")
                    # send_html_mail("Registro a bolsa laboral", "emails/nuevoregistro.html",
                    #                {'sistema': request.session['nombresistema'], 'e': empleador,
                    #                 'clave': DEFAULT_PASSWORD, 't': miinstitucion(), 'modelo': MODELO_EVALUACION,
                    #                 'autorizar': False, 'dominio': EMAIL_DOMAIN}, persona.lista_emails_envio(), [])
                    return JsonResponse({"result": "ok"})
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al registrarse."})

        elif action == 'editempresa':
            try:
                # empleador = Empleador.objects.get(pk=request.POST['id'])
                # persona = empleador.persona
                # empresa = empleador.empresa
                # f = Empleador2Form(request.POST)
                # if f.is_valid():
                #     empresa.nombre = f.cleaned_data['nombre']
                #     empresa.provincia = f.cleaned_data['provincia']
                #     empresa.direccion = f.cleaned_data['direccion']
                #     empresa.telefonos = f.cleaned_data['telefonos']
                #     persona.nombres = f.cleaned_data['nombres']
                #     persona.apellido1 = f.cleaned_data['apellido1']
                #     persona.apellido2 = f.cleaned_data['apellido2']
                #     empleador.cargo = f.cleaned_data['cargo']
                #     persona.email = f.cleaned_data['email']
                #     persona.save()
                #     empresa.save()
                #     empleador.save()
                #     log(u'Modifico empleador: %s' % empleador, request, "edit")
                #     return JsonResponse({"result": "ok", "id": empleador.id})
                empresa = EmpresaEmpleadora.objects.get(pk=request.POST['id'])
                # empleador = Empleador.objects.get(empresa=empresa)
                # persona = empleador.persona
                f = EmpleadorForm(request.POST)
                # f.desbloquear()
                if f.is_valid():
                    ruc = f.cleaned_data['ruc']
                    if ruc:
                        if EmpresaEmpleadora.objects.filter(status=True, ruc=ruc).exclude(id=empresa.id).exists():
                            return JsonResponse({"result": "bad", "mensaje": u"Empresa ya registrada."})
                    empresa.nombre = f.cleaned_data['nombre']
                    empresa.nombrecorto = f.cleaned_data['nombrecorto']
                    empresa.email = f.cleaned_data['email']
                    empresa.ruc = f.cleaned_data['ruc']
                    empresa.pais = f.cleaned_data['pais']
                    empresa.provincia = f.cleaned_data['provincia']
                    empresa.direccion = f.cleaned_data['direccion']
                    empresa.telefonos = f.cleaned_data['telefonos']

                    empresa.representante = f.cleaned_data['representante']
                    # empresa.fechainicio = f.cleaned_data['fechainicio']
                    # empresa.fechafin = f.cleaned_data['fechafin']
                    empresa.objetivo = f.cleaned_data['observacion']
                    # empresa.convenio = f.cleaned_data['convenio']
                    # if (f.cleaned_data['convenio'] == True):
                    #     empresa.tipoconvenio = f.cleaned_data['tipoconvenio']
                    # empresa.tipoinstitucion = f.cleaned_data['tipoinstitucion']
                    empresa.sectoreconomico = f.cleaned_data['sectoreconomico']
                    empresa.telefonoconv = f.cleaned_data['telefonoconv']
                    empresa.cargo = f.cleaned_data['cargo']
                    # persona.email = f.cleaned_data['email']
                    empresa.save(request)
                    # nombre = f.cleaned_data['contacto']
                    # nombres = ""
                    # apellido1 = ""
                    # apellido2 = ""
                    # if nombre.split().__len__() == 1:
                    #     apellido1 = nombre
                    # elif nombre.split().__len__() == 2:
                    #     apellido1 = nombre.split()[0]
                    #     apellido2 = nombre.split()[1]
                    # elif nombre.split().__len__() >= 3:
                    #     apellido1 = nombre.split()[0]
                    #     apellido2 = nombre.split()[1]
                    #     nombres = " ".join(nombre.split()[2:])


                    # persona.nombres = nombres
                    # persona.apellido1 = apellido1
                    # persona.apellido2 = apellido2

                    # empleador.save(request)
                    # persona.save(request)



                    # momentaneo
                    # if not empresa.persona().tiene_perfil():
                    #     persona.crear_perfil(empleador=empresa.empleador())
                    return JsonResponse({"result": "ok"})
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'resetear':
            try:
                empleador = Persona.objects.get(pk=int(request.POST['id']))
                resetear_clave_empresa(empleador)
                log(u'Reseteo clave de empleador: %s' % empleador, request, "edit")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al registrarse."})

        elif action == 'delempresa':
            try:
                empresa = EmpresaEmpleadora.objects.get(pk=int(request.POST['id']))
                if not empresa.no_puede_eliminar():
                    empresa.status = False
                    empresa.save(request, update_fields=['status'])
                    log(u'Elimino empresa empleadora : %s[%s]' % (empresa, empresa.id), request, "del")
                    return JsonResponse({"result": "ok"})
                else:
                    return JsonResponse({"result": "bad", "mensaje": u"No se puede eliminar la empresa."})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})

        elif action == 'addtiposolicitud':
            try:
                nombre = (request.POST['nombre'])
                tipo = MovilidadTipoSolicitud(nombre=nombre)
                tipo.save(request)
                log(u'Registro tipo de solicitud: %s' % tipo, request, "add")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                pass

        elif action == 'edittiposolicitud':
            try:
                tipo =  MovilidadTipoSolicitud.objects.get(pk=request.POST['id'])
                tipo.nombre=request.POST['nombre']
                tipo.save(request)
                log(u'Edito tipo de solicitud: %s' % tipo, request, "add")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                pass

        elif action == 'deltiposolicitud':
            try:
                tipo = MovilidadTipoSolicitud.objects.get(pk=int(request.POST['id']))
                tipo.delete()
                log(u'eliminó tipo de solicitud: %s' % tipo, request, "add")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})


        elif action == 'addbaselegal':
            try:
                descripcion = (request.POST['descripcion'])
                base = MovilidadBaseLegal(descripcion=descripcion)
                base.save(request)
                log(u'Registro base legal de solicitud: %s' % base, request, "add")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                pass

        elif action == 'editbaselegal':
            try:
                base =  MovilidadBaseLegal.objects.get(pk=request.POST['id'])
                base.descripcion=request.POST['descripcion']
                base.save(request)
                log(u'Edito base legal: %s' % base, request, "add")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                pass

        elif action == 'delbaselegal':
            try:
                base = MovilidadBaseLegal.objects.get(pk=int(request.POST['id']))
                base.delete()
                log(u'eliminó Base Legal : %s' % base, request, "add")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})


        elif action == 'addtipoestancia':
            try:
                descripcion = (request.POST['descripcion'])
                tipo = MovilidadTipoEstancia(descripcion=descripcion)
                tipo.save(request)
                log(u'Registro tipo de estancia: %s' % tipo, request, "add")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                pass

        elif action == 'edittipoestancia':
            try:
                tipo =  MovilidadTipoEstancia.objects.get(pk=request.POST['id'])
                tipo.descripcion=request.POST['descripcion']
                tipo.save(request)
                log(u'Edito tipo de estancia: %s' % tipo, request, "add")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                pass

        elif action == 'deltipoestancia':
            try:
                tipo = MovilidadTipoEstancia.objects.get(pk=int(request.POST['id']))
                tipo.delete()
                log(u'eliminó tipo de estancia: %s' % tipo, request, "add")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})

        elif action == 'cambiaestado':
            try:
                convenio = ConvenioEmpresa.objects.get(pk=request.POST['idconvenio'])
                if convenio.solicitud:
                    convenio.solicitud = False
                else:
                    convenio.solicitud = True
                    convenio.save(request)
                return JsonResponse({'result': 'ok', 'valor': convenio.solicitud})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'practicas':
            try:
                convenio = ConvenioEmpresa.objects.get(pk=request.POST['idconvenio'])
                if convenio.para_practicas:
                    convenio.para_practicas = False
                else:
                    convenio.para_practicas = True
                convenio.save(request)
                return JsonResponse({'result': 'ok', 'valor': convenio.para_practicas})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'pasantias':
            try:
                convenio = ConvenioEmpresa.objects.get(pk=request.POST['idconvenio'])
                if convenio.para_pasantias:
                    convenio.para_pasantias = False
                else:
                    convenio.para_pasantias = True
                convenio.save(request)
                return JsonResponse({'result': 'ok', 'valor': convenio.para_pasantias})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'notificar':
            try:
                f = NotificarManualForm(request.POST)
                if f.is_valid():
                    noti = f.cleaned_data['notificacion']
                    fecha_actual = datetime.now().date()
                    fecha_notificar = fecha_actual + timedelta(days=noti.dias)
                    notificados = DetalleNotificacionCovenio.objects.values_list('convenioempresa_id').filter(status=True).distinct('convenioempresa')
                    convenios = ConvenioEmpresa.objects.filter(status=True,fechafinalizacion__gte= fecha_actual).exclude(id__in=notificados)
                    for convenio in convenios:
                        if convenio.fechainicio + timedelta(days=noti.dias) <= fecha_actual:
                            departamentos = convenio.departamentoresponsable.all()
                            for departamento in departamentos:
                                if departamento.responsable:
                                    detalle = DetalleNotificacionCovenio(notificacion=noti,
                                                               convenioempresa=convenio,
                                                               fecha = fecha_actual
                                    )
                                    detalle.save(request)

                                    notificacion(detalle.notificacion.nombre,
                                             detalle.notificacion.mensaje, departamento.responsable, None, 'dir_convenios',
                                             convenio.id, 2, 'sga', ConvenioEmpresa, request)

                    log(u'Notificó masivo: %s' % noti, request, "add")

                return JsonResponse({"result": False}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'notificarind':
            try:
                f = NotificacionconvenioIndividualForm(request.POST)
                if f.is_valid():
                    nombre = f.cleaned_data['nombre']
                    mensaje = f.cleaned_data['mensaje']
                    convenio = ConvenioEmpresa.objects.get(pk=request.POST['id'])
                    departamentos = convenio.departamentoresponsable.all()
                    for departamento in departamentos:
                        if departamento.responsable:
                            detalle = DetalleNotificacionCovenio(
                                                       convenioempresa=convenio,
                                                       nombre=nombre,
                                                       mensaje=mensaje,
                                                       individual = True,
                                                       fecha = datetime.now().date(),
                            )
                            detalle.save(request)

                            notificacion(detalle.nombre,
                                     detalle.mensaje, departamento.responsable, None, 'dir_convenios',
                                     convenio.id, 2, 'sga', ConvenioEmpresa, request)

                    log(u'Notificó individual: %s' % detalle, request, "add")

                return JsonResponse({"result": False}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'addactividad':
            try:
                f = ActividadConvenioForm(request.POST)
                convenio = ConvenioEmpresa.objects.get(pk=request.POST['id'])
                fecha = datetime.strptime((request.POST['fecha']), '%Y-%m-%d').date()

                if f.is_valid():
                    actividad = ActividadConvenio(convenioempresa = convenio,
                                              actividad= f.cleaned_data['actividad'],
                                              fecha = fecha)

                    if 'archivo' in request.FILES:
                        newfile = request.FILES['archivo']
                        newfile._name = generar_nombre("actividadconvenio_", newfile._name)
                        actividad.archivo = newfile

                    actividad.save(request)
                    log(u'Adiciono actividad: %s' % actividad, request, "add")

                return JsonResponse({"result": False}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'editactividad':
            try:
                f = ActividadConvenioForm(request.POST)
                actividad = ActividadConvenio.objects.get(pk=request.POST['id'])
                if f.is_valid():
                    actividad.actividad = f.cleaned_data['actividad']

                    if 'archivo' in request.FILES:
                        newfile = request.FILES['archivo']
                        newfile._name = generar_nombre("actividadconvenio_", newfile._name)
                        actividad.archivo = newfile
                    actividad.save(request)
                    log(u'Editó actividad: %s' % actividad, request, "edit")

                return JsonResponse({"result": False}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'delactividad':
            try:
                # f = ActividadConvenioForm(request.POST)
                actividad = ActividadConvenio.objects.get(pk=request.POST['id'])
                actividad.status=False
                actividad.save(request)
                log(u'Eliminó actividad: %s' % actividad, request, "del")
                res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": str(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'addnotificacion':
            try:
                f = NotificacionconvenioForm(request.POST)
                if f.is_valid():
                    notificar = NotificacionConvenio(nombre = f.cleaned_data['nombre'],
                                              mensaje= f.cleaned_data['mensaje'],
                                              dias= f.cleaned_data['dias'])

                    notificar.save(request)
                    log(u'Adicionó notificacion para convenio: %s' % notificar, request, "add")

                return JsonResponse({"result": False}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'editnotificacion':
            try:
                f = NotificacionconvenioForm(request.POST)
                notificar = NotificacionConvenio.objects.get(pk=request.POST['id'])
                if f.is_valid():
                    notificar.nombre = f.cleaned_data['nombre']
                    notificar.mensaje = f.cleaned_data['mensaje']
                    notificar.dias = f.cleaned_data['dias']

                    notificar.save(request)
                    log(u'Editó notificación para convenio: %s' % notificar, request, "edit")

                return JsonResponse({"result": False}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos."})

        elif action == 'delnotificacion':
            try:
                notificar = NotificacionConvenio.objects.get(pk=request.POST['id'])
                notificar.status=False
                notificar.save(request)
                log(u'Eliminó notificación para convenio: %s' % notificar, request, "del")
                res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": str(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'moverregistros':
            try:
                idenvia = int(request.POST['id'])
                idrecibe= int(request.POST['empresa'])
                empresa = EmpresaEmpleadora.objects.get(id=idenvia)
                recibe = EmpresaEmpleadora.objects.get(id=idrecibe)
                empresa.convenioempresa_set.filter().update(empresaempleadora_id=idrecibe)
                empresa.detallepreinscripcionpracticaspp_set.filter().update(empresaempleadora_id=idrecibe)
                empresa.practicaspreprofesionalesinscripcion_set.filter().update(empresaempleadora_id=idrecibe)
                empresa.acuerdocompromiso_set.filter().update(empresa_id=idrecibe)
                empresa.empleador_set.filter().update(empresa_id=idrecibe)
                empresa.ofertalaboral_set.filter().update(empresa_id=idrecibe)
                empresa.ofertaspracticas_set.filter().update(empresa_id=idrecibe)
                empresa.movilidadsolicitud_set.filter().update(empresa_id=idrecibe)
                empresa.cartavinculacionpracticaspreprofesionales_set.filter().update(empresa_id=idrecibe)
                log(f'Actualización de registros de empresa empleadora {empresa} a {recibe} en 9 tablas', request,'edit')
                return JsonResponse({'result':False}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, 'mensaje': f'Error: {ex}'})

        elif action == 'importarconvenios':
            try:
                archivo_ = request.FILES['archivo']
                cabecera = ['codigoconvenio', 'codigoempresa',
                            'codigotipoconvenio', 'fechainicio',
                            'fechafin', 'cargo', 'responsableexterno',
                            'objetivo', 'parapracticasprofesionales',
                            'parapasantias', 'numeroocas/memorando', 'fechaocas/memorando',
                            'urlarchivo', 'cedularesponsableinterno']
                name_hoja = pd.ExcelFile(archivo_).sheet_names[0]
                df = pd.read_excel(archivo_, sheet_name=name_hoja)
                df.columns = formatear_cabecera_pd(df)
                num_columnas = len(df.columns)
                lista_sc, cont = [], 0
                for c in cabecera:
                    if not c in df.columns:
                        raise NameError(f'Formato de archivo erróneo: La columna {c} no se encuentra en el documento.')

                for index, row in df.iterrows():
                    codigo_convenio = str(row['codigoconvenio']).strip().split('.')[0]
                    convenio = ConvenioEmpresa.objects.filter(pk=codigo_convenio).first()
                    codigo_tipo_convenio = str(row['codigotipoconvenio']).strip().split('.')[0]
                    tipoconvenio = TipoConvenio.objects.filter(pk=codigo_tipo_convenio).first()
                    if not tipoconvenio:
                        raise NameError(f'Código:{codigo_tipo_convenio}, de tipo de convenio no existe, por favor llene su matriz con un código válido')
                    codigo_empresa = str(row['codigoempresa']).strip().split('.')[0]
                    empresa = EmpresaEmpleadora.objects.filter(pk=codigo_empresa).first()
                    if not empresa:
                        raise NameError(f'Código:{codigo_empresa}, de tipo de empresa no existe, por favor llene su matriz con un código válido')
                    fecha_inicio = str(row['fechainicio']).strip()
                    fecha_fin = str(row['fechafin']).strip()
                    cargo = str(row['cargo']).strip()
                    responsable_externo = str(row['responsableexterno']).strip()
                    para_practicas_profesionales = str(row['parapracticasprofesionales']).strip()
                    parapasantias = str(row['parapasantias']).strip()
                    numero_ocas = str(row['numeroocas/memorando']).strip().split('.')[0]
                    fecha_ocas = str(row['fechaocas/memorando'])
                    url_archivo = str(row['urlarchivo']).strip()
                    cedula_responsable_interno = str(row['cedularesponsableinterno']).strip().split('.')[0]
                    responsable_interno = consultarPersona(cedula_responsable_interno)
                    if not responsable_interno:
                        raise NameError(f'Cédula:{cedula_responsable_interno}, No pertenece a ningún funcionario registrado, por favor digite una cédula valida.')

                    if convenio:
                        convenio.tipoconvenio = tipoconvenio
                        convenio.empresaempleadora = empresa
                        convenio.fechainicio = fecha_inicio
                        convenio.fechafinalizacion = fecha_fin
                        convenio.cargoresponsableexterno = cargo
                        convenio.responsableexterno = responsable_externo
                        convenio.para_practicas = para_practicas_profesionales
                        convenio.para_pasantias = parapasantias
                        convenio.numocas = numero_ocas
                        convenio.fechaocas = fecha_ocas
                        convenio.url = url_archivo
                        convenio.save(request)
                        log(f'Edito convenio de empresa - {convenio}', request, 'edit')
                    elif codigo_convenio == '':
                        cont += 1
                        convenio = ConvenioEmpresa(empresaempleadora=empresa,
                                                   tipoconvenio=tipoconvenio,
                                                   fechainicio=fecha_inicio,
                                                   fechafinalizacion=fecha_fin,
                                                   cargoresponsableexterno=cargo,
                                                   responsableexterno=responsable_externo,
                                                   para_practicas=para_practicas_profesionales,
                                                   para_pasantias=parapasantias,
                                                   numocas=numero_ocas,
                                                   fechaocas=fecha_ocas,
                                                   url=url_archivo)
                        convenio.save(request)
                        log(f'Creo convenio con empresa - {convenio}', request, 'add')
                    else:
                        lista_sc.append(codigo_convenio)

                    if convenio:
                        inicio = 15
                        for index in range(inicio, num_columnas):
                            idcarrera = df.columns[index]
                            value = df[idcarrera].values
                            if not Carrera.objects.filter(id=idcarrera).exists():
                                raise NameError(f'Código de carrera {idcarrera}: no existe, por favor revise su matriz y digite un codigo de carrera válido')
                            conv_carrera = convenio.conveniocarrera_set.filter(status=True, carrera_id=idcarrera)
                            if value and not conv_carrera:
                                conv_carrera = ConvenioCarrera(convenioempresa=convenio, carrera_id=idcarrera)
                                conv_carrera.save(request)
                                log(f'Creo convenio de carrera - {conv_carrera}', request, 'add')
                            elif conv_carrera and not value:
                                conv_carrera.update(status=False)

                total_registros = df.shape[0]
                no_encontrados = len(lista_sc)
                actuaizados = total_registros - no_encontrados
                mensaje = f'Se actualizo {actuaizados} convenios con éxito, Convenios creados: {cont}, Convenios no encontrados: {no_encontrados}-{lista_sc}'
                return JsonResponse({'result': False, 'modalsuccess': True, 'mensaje': mensaje}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, 'mensaje': f'Error: {ex}'})

        return JsonResponse({"result": "bad", "mensaje": u"Solicitud Incorrecta."})


        return HttpResponseRedirect(request.path)
    else:
        if 'action' in request.GET:
            data['action'] = action = request.GET['action']

            if action == 'addconvenio':
                try:
                    data['title'] = u'Adicionar convenio con empresa'
                    form = ConvenioEmpresaForm()
                    form.fields['empresaempleadora'].queryset = EmpresaEmpleadora.objects.none()
                    data['form'] = form
                    # data['form'] = form.adicionar()
                    data['form2'] = ArchivoConvenioForm()
                    form3 = DetalleResponsableInternoForm()
                    form3.fields['responsable'].queryset = Persona.objects.none()
                    form3.fields['cargo'].queryset = DenominacionPuesto.objects.none()
                    data['form3'] = form3
                    return render(request, "adm_convenioempresa/add.html", data)
                except Exception as ex:
                    pass

            elif action == 'editconvenio':
                try:
                    data['title'] = u'Editar convenio con empresa'
                    data['convenio'] = convenio = ConvenioEmpresa.objects.get(pk=int(request.GET['id']))
                    # form = ConvenioEmpresaForm(initial=model_to_dict(convenio))
                    form = ConvenioEmpresaForm(initial=
                    {
                        'empresaempleadora': convenio.empresaempleadora,
                        'tipoconvenio': convenio.tipoconvenio,
                        'fechainicio': convenio.fechainicio,
                        'fechafinalizacion': convenio.fechafinalizacion,
                        'objetivo': convenio.objetivo,
                        # 'departamentoresponsable': convenio.departamentoresponsable.all(),
                        'responsableexterno': convenio.responsableexterno,
                        'dias_notificacion': convenio.dias_notificacion,
                        'para_practicas': convenio.para_practicas,
                        'para_pasantias': convenio.para_pasantias,
                        'denominacionPuesto': convenio.cargo_denominaciones.all(),
                        'numocas': convenio.numocas,
                        'fechaocas': convenio.fechaocas,
                        'archivoocas': convenio.archivoocas,
                    }
                    )
                    if convenio.archivoconvenio_set.filter(status=True).exists():
                        archivoconvenio = convenio.archivoconvenio_set.filter(status=True).first()
                        # 'archivo': archivoconvenio.archivo
                        data['form2'] = form2 = ArchivoConvenioForm(initial={'tipoarchivoconvenio': archivoconvenio.tipoarchivoconvenio,})
                    else:
                        data['form2'] = form2 = ArchivoConvenioForm()
                    # form.cargar_responsableinterno(convenio.responsableinterno)
                    form.fields['empresaempleadora'].queryset = EmpresaEmpleadora.objects.filter(id=convenio.empresaempleadora.id)
                    data['form'] = form
                    form3=DetalleResponsableInternoForm()
                    form3.fields['responsable'].queryset = Persona.objects.none()
                    form3.fields['cargo'].queryset = DenominacionPuesto.objects.none()
                    data['form3'] = form3
                    data['responsablesConvenio'] = convenio.responsableconvenio_set.filter(status=True)

                    return render(request, "adm_convenioempresa/editconvenio.html", data)
                except Exception as ex:
                    pass

            elif action == 'adddepartamento':
                try:
                    data['filtro'] = filtro = ConvenioEmpresa.objects.get(pk=int(request.GET['id']))
                    form = DepartamentoConvenioForm(initial=model_to_dict(filtro))
                    data['form2'] = form
                    template = get_template("adm_convenioempresa/modal/departamentoform.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action=='delconvenio':
                try:
                    data['title'] = u'Eliminar convenio con empresa'
                    data['convenio'] = ConvenioEmpresa.objects.get(pk=int(request.GET['id']))
                    return render(request, "adm_convenioempresa/delconvenio.html", data)
                except Exception as ex:
                    pass

            # TIPO DE CONVENIO
            elif action == 'tipoconvenio':
                try:
                    data['title'] = u'Tipo convenio'
                    search = None
                    ids = None
                    if 's' in request.GET:
                        search = request.GET['s'].strip()
                        ss = search.split(' ')
                        if len(ss) == 1:
                            convenio = TipoConvenio.objects.filter(Q(nombre__icontains=search), Q(status=True)).distinct()
                        elif len(ss) == 2:
                            convenio = TipoConvenio.objects.filter(Q(nombre__icontains=ss[0]), Q(nombre__icontains=ss[1]), Q(status=True)).distinct()
                        elif len(ss) == 3:
                            convenio = TipoConvenio.objects.filter(Q(nombre__icontains=ss[0]), Q(nombre__icontains=ss[1]), Q(nombre__icontains=ss[2]), Q(status=True)).distinct()
                        else:
                            convenio = TipoConvenio.objects.filter(Q(nombre__icontains=ss[0]), Q(nombre__icontains=ss[1]), Q(nombre__icontains=ss[2]), Q(nombre__icontains=ss[3]), Q(status=True)).distinct()
                    else:
                        convenio = TipoConvenio.objects.filter(status=True)
                    paging = MiPaginador(convenio, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['search'] = search if search else ""
                    data['ids'] = ids if ids else ""
                    data['convenioempresas'] = page.object_list
                    return render(request, "adm_convenioempresa/viewtipoconvenio.html", data)
                except Exception as ex:
                    pass

            elif action == 'addtipoconv':
                try:
                    data['title'] = u'Adicionar tipo de convenio'
                    data['form'] = TipoConvenioForm()
                    template = get_template("adm_convenioempresa/formtipoconvenio.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})

                except Exception as ex:
                    pass

            elif action == 'edittipoconv':
                try:
                    data['title'] = u'Editar tipo de convenio'
                    data['tipo'] = convenio = TipoConvenio.objects.get(pk=int(request.GET['id']))
                    form = TipoConvenioForm(initial=model_to_dict(convenio))
                    data['form']= form
                    template = get_template("adm_convenioempresa/formtipoconvenio.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})

                except Exception as ex:
                    pass

            elif action=='deltipoconv':
                try:
                    data['title'] = u'Eliminar tipo de convenio'
                    data['tipo'] = TipoConvenio.objects.get(pk=int(request.GET['id']))
                    return render(request, "adm_convenioempresa/deltipoconvenio.html", data)
                except Exception as ex:
                    pass

            # TIPO DE ARCHIVO
            elif action == 'tipoarchivo':
                try:
                    data['title'] = u'Tipo de archivo'
                    search = None
                    ids = None
                    if 's' in request.GET:
                        search = request.GET['s'].strip()
                        ss = search.split(' ')
                        if len(ss) == 1:
                            tipoarchivo = TipoArchivoConvenio.objects.filter(Q(nombre__icontains=search), Q(status=True)).distinct()
                        elif len(ss) == 2:
                            tipoarchivo = TipoArchivoConvenio.objects.filter(Q(nombre__icontains=ss[0]), Q(nombre__icontains=ss[1]), Q(status=True)).distinct()
                        elif len(ss) == 3:
                            tipoarchivo = TipoArchivoConvenio.objects.filter(Q(nombre__icontains=ss[0]), Q(nombre__icontains=ss[1]), Q(nombre__icontains=ss[2]), Q(status=True)).distinct()
                        else:
                            tipoarchivo = TipoArchivoConvenio.objects.filter(Q(nombre__icontains=ss[0]), Q(nombre__icontains=ss[1]), Q(nombre__icontains=ss[2]), Q(nombre__icontains=ss[3]), Q(status=True)).distinct()
                    else:
                        tipoarchivo = TipoArchivoConvenio.objects.filter(status=True)
                    paging = MiPaginador(tipoarchivo, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['search'] = search if search else ""
                    data['ids'] = ids if ids else ""
                    data['tiposarchivo'] = page.object_list
                    return render(request, "adm_convenioempresa/viewtipoarchivo.html", data)
                except Exception as ex:
                    pass

            elif action == 'addtipoarch':
                try:
                    data['title'] = u'Adicionar tipo de archivo'
                    data['form'] = TipoArchivoConvenioForm()
                    template = get_template("adm_convenioempresa/formtipoarchivo.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'edittipoarch':
                try:
                    data['title'] = u'Editar tipo de archivo'
                    data['tipoarchivo'] = convenio = TipoArchivoConvenio.objects.get(pk=int(request.GET['id']))
                    form = TipoArchivoConvenioForm(initial=model_to_dict(convenio))
                    data['form'] = form
                    template = get_template("adm_convenioempresa/formtipoarchivo.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'deltipoarch':
                try:
                    data['title'] = u'Eliminar tipo de archivo'
                    data['tipoarchivo'] = TipoArchivoConvenio.objects.get(pk=int(request.GET['id']))
                    return render(request, "adm_convenioempresa/deltipoarchivo.html", data)
                except Exception as ex:
                    pass

            # CONVENIO ARCHIVO
            elif action == 'archivoconvenio':
                try:
                    data['title'] = u'Archivos de convenio'
                    search = None
                    ids = None
                    convenioempresa = ConvenioEmpresa.objects.get(pk=int(request.GET['id']), status=True)
                    if 's' in request.GET:
                        search = request.GET['s'].strip()
                        ss = search.split(' ')
                        if len(ss) == 1:
                            archivoconvenio = ArchivoConvenio.objects.filter(Q(tipoarchivoconvenio__nombre__icontains=search), Q(status=True), Q(convenioempresa=convenioempresa)).distinct()
                        elif len(ss) == 2:
                            archivoconvenio = ArchivoConvenio.objects.filter(Q(tipoarchivoconvenio__nombre__icontains=ss[0]), Q(tipoarchivoconvenio__nombre__icontains=ss[1]), Q(status=True), Q(convenioempresa=convenioempresa)).distinct()
                        elif len(ss) == 3:
                            archivoconvenio = ArchivoConvenio.objects.filter(Q(tipoarchivoconvenio__nombre__icontains=ss[0]), Q(tipoarchivoconvenio__nombre__icontains=ss[1]), Q(tipoarchivoconvenio__nombre__icontains=ss[2]), Q(status=True), Q(convenioempresa=convenioempresa)).distinct()
                        else:
                            archivoconvenio = ArchivoConvenio.objects.filter(Q(tipoarchivoconvenio__nombre__icontains=ss[0]), Q(tipoarchivoconvenio__nombre__icontains=ss[1]), Q(tipoarchivoconvenio__nombre__icontains=ss[2]), Q(tipoarchivoconvenio__nombre__icontains=ss[3]), Q(status=True), Q(convenioempresa=convenioempresa)).distinct()
                    else:
                        archivoconvenio = ArchivoConvenio.objects.filter(status=True, convenioempresa=convenioempresa)
                    paging = MiPaginador(archivoconvenio, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['search'] = search if search else ""
                    data['ids'] = ids if ids else ""
                    data['archivosconvenio'] = page.object_list
                    data['convenioempresa'] = convenioempresa
                    return render(request, "adm_convenioempresa/viewarchivoconvenio.html", data)
                except Exception as ex:
                    pass

            elif action == 'carrerasconvenios_ori':
                try:
                    data['title'] = u'Listado Carreras'
                    data['convenioempresa'] = convenioempresa = ConvenioEmpresa.objects.get(pk=int(request.GET['id']), status=True)
                    data['listaconveniocarreras'] = listaconveniocarreras = convenioempresa.conveniocarrera_set.filter(status=True)
                    listaconveniocarreras = listaconveniocarreras.values_list('carrera')
                    data['listacarreras'] = Carrera.objects.filter(status=True).exclude(pk__in=listaconveniocarreras).exclude(coordinacion__id__in=[6,8,9,10]).order_by('nombre')
                    return render(request, "adm_convenioempresa/listadoconveniocarreras.html", data)
                except Exception as ex:
                    pass

            elif action == 'carrerasconvenios':
                try:
                    data['title'] = u'Listado de Programas'
                    data['convenioempresa'] = convenioempresa = ConvenioEmpresa.objects.get(pk=int(request.GET['id']), status=True)
                    listaconveniocarreras = convenioempresa.conveniocarrera_set.filter(status=True).values_list('carrera_id')
                    mod = request.GET.get('mod', '')
                    search = request.GET.get('s', '')
                    url_vars = ""
                    filtros, coordinacion, estado = Q(status=True, coordinacion__isnull=False), \
                                                    request.GET.get('coordinacion',''), \
                                                    request.GET.get('estado','')
                    if estado:
                        data['estado'] = estado = int(estado)
                        url_vars += "&estado={}".format(estado)
                        if estado == 0:
                            filtros = filtros & Q(estadocarrera__isnull=True)
                        else:
                            filtros = filtros & Q(estadocarrera_id=estado)
                    if mod:
                        data['mod'] = mod = int(mod)
                        url_vars += "&mod={}".format(mod)
                        filtros = filtros & Q(modalidad=int(mod))

                    if coordinacion:
                        data['coordinacion'] = coordinacion = int(coordinacion)
                        url_vars += "&coordinacion={}".format(coordinacion)
                        filtros = filtros & Q(coordinacion__id=int(coordinacion))

                    if search:
                        url_vars += "&s={}".format(search)
                        if search.isdigit():
                            filtros = filtros & Q(id=search)
                        else:
                            filtros = filtros & Q(nombre__icontains=search)
                    data["url_vars"] = url_vars
                    carreras = Carrera.objects.filter(filtros).exclude(coordinacion__id__in=[6, 8, 9, 10]).order_by('nombre')
                    paging = MiPaginador(carreras, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)

                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['listacarreras'] = page.object_list
                    data['seleccionados'] = seleccion = len(listaconveniocarreras)
                    data['total'] = total = len(Carrera.objects.filter(status=True).exclude(coordinacion__id__in=[6, 8, 9, 10]).values_list('id'))
                    data['sinseleccionar'] = total - seleccion
                    data['coordinaciones'] = Coordinacion.objects.filter(status=True).exclude(id__in=[6, 8, 9, 10])
                    data['estados'] = EstadoCarrera.objects.filter(status=True)
                    return render(request, "adm_convenioempresa/listadoconveniocarreras.html", data)
                except Exception as ex:
                    pass

            elif action == 'actividades':
                try:
                    data['title'] = u'Actividades'
                    data['convenioempresa'] = convenioempresa = ConvenioEmpresa.objects.get(pk=int(request.GET['id']), status=True)
                    fecha = datetime.now().date()
                    panio = fecha.year
                    pmes = fecha.month
                    if 'mover' in request.GET:
                        mover = request.GET['mover']

                        if mover == 'anterior':
                            mes = 10
                            anio = 2021
                            pmes = mes - 1
                            if pmes == 0:
                                pmes = 12
                                panio = anio - 1
                            else:
                                panio = anio

                        elif mover == 'proximo':
                            mes = 11
                            anio = 2021
                            pmes = mes + 1
                            if pmes == 13:
                                pmes = 1
                                panio = anio + 1
                            else:
                                panio = anio
                    s_anio = panio
                    s_mes = pmes
                    s_dia = 1
                    data['mes'] = MESES_CHOICES[s_mes - 1]
                    data['ws'] = [0, 7, 14, 21, 28, 35]
                    lista = {}
                    listaactividades = {}
                    for i in range(1, 43, 1):
                        dia = {i: 'no'}
                        actividaddia = {i: None}
                        lista.update(dia)
                        listaactividades.update(actividaddia)
                    comienzo = False
                    fin = False
                    for i in lista.items():
                        try:
                            fecha = date(s_anio, s_mes, s_dia)
                            if fecha.isoweekday() == i[0] and fin is False and comienzo is False:
                                comienzo = True
                        except Exception as ex:
                            pass
                        if comienzo:
                            try:
                                fecha = date(s_anio, s_mes, s_dia)
                            except Exception as ex:
                                fin = True
                        if comienzo and fin is False:
                            dia = {i[0]: s_dia}
                            s_dia += 1
                            lista.update(dia)
                            actividaddia = ActividadConvenio.objects.filter(fecha=fecha,
                                                                            status=True,convenioempresa=convenioempresa)
                            diaact = []
                            if actividaddia.exists():
                                valor = str(actividaddia[0].actividad)
                            else:
                                valor = ""
                            act = [valor, (fecha < datetime.now().date() and valor == ""), actividaddia.count(),
                                   fecha.strftime('%d-%m-%Y')]
                            diaact.append(act)
                            listaactividades.update({i[0]: diaact})
                    data['dias_mes'] = lista
                    data['s_anio'] = s_anio
                    data['s_mes'] = s_mes
                    data['lista'] = lista
                    data['listaactividades'] = listaactividades
                    data['mostrar_dia_actual'] = fecha.month == datetime.now().date().month and fecha.year == datetime.now().date().year

                    data['actividades'] = actividades = convenioempresa.actividadconvenio_set.filter(status=True,fecha__month=s_mes, fecha__year=s_anio)


                    data['dwnm'] = ['lunes', 'martes', 'miercoles', 'jueves', 'viernes', 'sabado', 'domingo']
                    data['dwn'] = [1, 2, 3, 4, 5, 6, 7]

                    return render(request, "adm_convenioempresa/viewactividades.html", data)
                except Exception as ex:
                    pass

            elif action == 'excelconvenios':
                try:
                    fecha_actual = datetime.now().date()
                    filtro = Q(status=True)
                    if 'desde' in request.GET:
                        filtro = filtro & Q(fechainicio__gte=request.GET['desde'])
                    if 'hasta' in request.GET:
                        filtro = filtro & Q(fechafinalizacion__lte=request.GET['hasta'])
                    if 'vigencia' in request.GET:
                        if request.GET['vigencia'] == '1':
                            filtro = filtro & Q(fechafinalizacion__gt=fecha_actual)
                        else:
                            filtro = filtro & Q(fechafinalizacion__lte=fecha_actual)

                    __author__ = 'Unemi'
                    style0 = easyxf('font: name Times New Roman, color-index blue, bold off', num_format_str='#,##0.00')
                    style_nb = easyxf('font: name Times New Roman, color-index blue, bold on', num_format_str='#,##0.00')
                    style_sb = easyxf('font: name Times New Roman, color-index blue, bold on')
                    title = easyxf('font: name Times New Roman, color-index blue, bold on , height 350; alignment: horiz centre')
                    style1 = easyxf(num_format_str='D-MMM-YY')
                    font_style = XFStyle()
                    font_style.font.bold = True
                    font_style2 = XFStyle()
                    font_style2.font.bold = False
                    wb = Workbook(encoding='utf-8')
                    ws = wb.add_sheet('convenios')
                    ws.write_merge(0, 0, 0, 12, 'UNIVERSIDAD ESTATAL DE MILAGRO', title)
                    response = HttpResponse(content_type="application/ms-excel")
                    response[
                        'Content-Disposition'] = 'attachment; filename=listado_convenios' + random.randint(
                        1, 10000).__str__() + '.xls'

                    columns = [
                        (u"NRO.", 1000),
                        (u"EMPRESA", 10000),
                        (u"TIPO CONVENIO", 10000),
                        (u"FECHA INICIO", 5000),
                        (u"FECHA FIN", 5000),
                        (u"CARRERAS", 10000),
                        (u"RESPONSABLE INTERNO", 10000),
                        (u"EMAIL RESP INTERNO", 10000),
                        (u"RESPONSABLE EXTERNO", 10000),
                        (u"EMAIL EMPRESA", 10000),
                        (u"TELF. EMPRESA", 10000),
                        (u"OBJETIVO", 25000),
                        (u"VIGENCIA", 4500),
                    ]
                    row_num = 3
                    campo5=''
                    for col_num in range(len(columns)):
                        ws.write(row_num, col_num, columns[col_num][0], font_style)
                        ws.col(col_num).width = columns[col_num][1]
                    date_format = xlwt.XFStyle()
                    date_format.num_format_str = 'yyyy/mm/dd'
                    listaconvenios = ConvenioEmpresa.objects.filter(filtro)
                    if filtro:
                        listaconvenios.filter(filtro)
                    row_num = 4
                    cont = 1
                    i = 0
                    for lista in listaconvenios:
                        cadena = []
                        empresa = lista.empresaempleadora.nombre
                        tipo = lista.tipoconvenio.nombre
                        finicio = lista.fechainicio
                        ffin = lista.fechafinalizacion
                        resp=''
                        mails=''
                        rinterno=''
                        for resp in lista.departamentoresponsable.all():
                            if resp.responsable:
                                rinterno = rinterno + str(resp.responsable)+'\n'
                                mails = mails + str(resp.responsable.emailinst)+'\n'
                        #campo5 = lista.responsableinterno.nombre_completo().__str__()
                        rexterno = lista.responsableexterno
                        objetivo = lista.objetivo
                        vigencia = 'VIGENTE' if lista.fechafinalizacion > fecha_actual else 'NO VIGENTE'
                        carreras = ''
                        if lista.conveniocarrera_set.filter(status=True):
                            carreras = lista.conveniocarrera_set.values('carrera__nombre').filter(status=True)
                            for cam in carreras:
                                cadena.append(cam['carrera__nombre'])
                        if cadena:
                            carreras = cadena
                        ws.write(row_num, 0, cont, font_style2)
                        ws.write(row_num, 1, empresa, font_style2)
                        ws.write(row_num, 2, tipo, font_style2)
                        ws.write(row_num, 3, finicio, date_format)
                        ws.write(row_num, 4, ffin, date_format)
                        ws.write(row_num, 5, carreras.__str__(), font_style2)
                        ws.write(row_num, 6, rinterno, font_style2)
                        ws.write(row_num, 7, mails, font_style2)
                        ws.write(row_num, 8, rexterno, font_style2)
                        ws.write(row_num, 9, lista.empresaempleadora.email, font_style2)
                        ws.write(row_num, 10, lista.empresaempleadora.telefonos, font_style2)
                        ws.write(row_num, 11, objetivo, font_style2)
                        ws.write(row_num, 12, vigencia, font_style2)
                        row_num += 1
                        cont += 1
                    wb.save(response)
                    return response
                except Exception as ex:
                    print(u'Error on line {}'.format(sys.exc_info()[-1].tb_lineno))
                    pass

            elif action == 'excelconvenioempresa':
                try:
                    __author__ = 'Unemi'
                    style0 = easyxf('font: name Times New Roman, color-index blue, bold off', num_format_str='#,##0.00')
                    style_nb = easyxf('font: name Times New Roman, color-index blue, bold on', num_format_str='#,##0.00')
                    style_sb = easyxf('font: name Times New Roman, color-index blue, bold on')
                    title = easyxf('font: name Times New Roman, color-index blue, bold on , height 350; alignment: horiz centre')
                    style1 = easyxf(num_format_str='D-MMM-YY')
                    font_style = XFStyle()
                    font_style.font.bold = True
                    font_style2 = XFStyle()
                    font_style2.font.bold = False
                    wb = Workbook(encoding='utf-8')
                    ws = wb.add_sheet('empresas')
                    ws.write_merge(0, 0, 0, 10, 'UNIVERSIDAD ESTATAL DE MILAGRO', title)
                    response = HttpResponse(content_type="application/ms-excel")
                    response[
                        'Content-Disposition'] = 'attachment; filename=listado_empresas' + random.randint(
                        1, 10000).__str__() + '.xls'

                    columns = [
                        (u"EMPRESA", 10000),
                        (u"RUC", 4000),
                        (u"PAIS", 5000),
                        (u"PROVINCIA", 5000),
                        (u"DIRECCION", 8000),
                        (u"TELEFONO", 8000),
                        (u"AUTORIZADA", 5000),
                        (u"REPRESENTANTE", 7000),
                        (u"FECHA  INICIO", 3500),
                        (u"FECHA  FIN", 3500),
                        (u"OBJETIVO", 3500),
                        (u"TIPO CONVENIO", 3500),
                        (u"SECTOR ECONOMICO", 3500),
                        (u"EMAIL", 3500),
                    ]
                    row_num = 3
                    for col_num in range(len(columns)):
                        ws.write(row_num, col_num, columns[col_num][0], font_style)
                        ws.col(col_num).width = columns[col_num][1]
                    date_format = xlwt.XFStyle()
                    date_format.num_format_str = 'yyyy/mm/dd'
                    listaconvenios = EmpresaEmpleadora.objects.filter(convenioempresa__status=True, status=True).distinct()
                    row_num = 4
                    i = 0
                    for lista in listaconvenios:
                        campo1 = lista.nombre
                        campo2 = lista.ruc
                        campo3 = ''
                        campo4 = ''
                        if lista.pais:
                            campo3 = lista.pais.nombre
                        if lista.provincia:
                            campo4 = lista.provincia.nombre
                        campo5 = lista.direccion
                        campo6 = lista.telefonos
                        campo7 = 'NO'
                        if lista.autorizada:
                            campo7 = 'SI'
                        campo8 = lista.representante
                        campo9 = lista.fechainicio
                        campo10 = lista.fechafin
                        campo11 = lista.objetivo
                        campo12 = ''
                        if lista.tipoconvenio:
                            campo12 = lista.tipoconvenio.nombre
                        campo13 = lista.get_sectoreconomico_display()
                        campo14 = lista.email
                        ws.write(row_num, 0, campo1, font_style2)
                        ws.write(row_num, 1, campo2, font_style2)
                        ws.write(row_num, 2, campo3, font_style2)
                        ws.write(row_num, 3, campo4, font_style2)
                        ws.write(row_num, 4, campo5, font_style2)
                        ws.write(row_num, 5, campo6, font_style2)
                        ws.write(row_num, 6, campo7, font_style2)
                        ws.write(row_num, 7, campo8, font_style2)
                        ws.write(row_num, 8, campo9, date_format)
                        ws.write(row_num, 9, campo10, date_format)
                        ws.write(row_num, 10, campo11, date_format)
                        ws.write(row_num, 11, campo12, date_format)
                        ws.write(row_num, 12, campo13, date_format)
                        ws.write(row_num, 13, campo14, date_format)
                        row_num += 1
                    wb.save(response)
                    return response
                except Exception as ex:
                    print(u'Error on line {}'.format(sys.exc_info()[-1].tb_lineno))
                    pass

            elif action == 'vercarreras':
                try:
                    data = {}
                    convenio = ConvenioEmpresa.objects.get(pk=request.GET['id'], status=True)
                    data['listaconveniocarreras'] = carreras = convenio.conveniocarrera_set.filter(status=True)
                    data['total'] = len(carreras)
                    template = get_template("adm_convenioempresa/vercarrera.html")
                    json_content = template.render(data)
                    return JsonResponse({"result": "ok", 'data': json_content})
                except Exception as ex:
                    return JsonResponse({"result": "bad", "mensaje": u"Error al obtener los datos."})

            elif action == 'addarchivoconv':
                try:
                    data['title'] = u'Adicionar archivo de convenio'
                    data['form'] = ArchivoConvenioForm()
                    data['convenioempresa'] = ConvenioEmpresa.objects.get(pk=int(request.GET['id']), status=True)
                    return render(request, "adm_convenioempresa/addarchivoconvenio.html", data)
                except Exception as ex:
                    pass

            elif action == 'editarchivoconv':
                try:
                    data['title'] = u'Editar archivo de convenio'
                    data['archivoconvenio'] = archivoconvenio = ArchivoConvenio.objects.get(pk=int(request.GET['id']))
                    data['form'] = ArchivoConvenioForm(initial={'tipoarchivoconvenio': archivoconvenio.tipoarchivoconvenio})
                    return render(request, "adm_convenioempresa/editarchivoconvenio.html", data)
                except Exception as ex:
                    pass

            elif action == 'delarchivoconv':
                try:
                    data['title'] = u'Eliminar archivo de convenio'
                    data['archivoconvenio'] = ArchivoConvenio.objects.get(pk=int(request.GET['id']))
                    return render(request, "adm_convenioempresa/delarchivoconvenio.html", data)
                except Exception as ex:
                    pass

            # EMPRESA
            elif action == 'empresa':
                try:
                    data['title'] = u'Empresas'
                    search = None
                    ids = None
                    if 'se' in request.GET:
                        search = request.GET['se']
                        empresa = EmpresaEmpleadora.objects.filter(Q(nombre__icontains=search) | Q(ruc__icontains=search))
                    elif 'ide' in request.GET:
                        ids = request.GET['ide']
                        empresa = EmpresaEmpleadora.objects.filter(id=ids)
                    else:
                        empresa = EmpresaEmpleadora.objects.filter(status=True).order_by('-convenio', 'nombre')
                    paging = MiPaginador(empresa.filter(status=True), 25)
                    p = 1
                    try:
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        page = paging.page(p)
                    except Exception as ex:
                        page = paging.page(p)
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['search'] = search if search else ""
                    data['ids'] = ids if ids else ""
                    data['empresas'] = page.object_list
                    data['clave'] = DEFAULT_PASSWORD
                    return render(request, "adm_convenioempresa/viewempresa.html", data)
                except Exception as ex:
                    pass

            elif action == 'addempresa':
                try:
                    data['title'] = u"Adicionar empresa"
                    form = EmpleadorForm()
                    form.adicionar()
                    # form.bloquear()
                    data['form'] = form
                    return render(request, "adm_convenioempresa/addempresa.html", data)
                except Exception as ex:
                    pass

            elif action == 'editempresa':
                try:
                    data['title'] = u"Editar datos del Empleador"
                    # empresa = EmpresaEmpleadora.objects.get(pk=request.GET['id'])
                    # empleador = Empleador.objects.get(empresa=empresa)
                    # form = Empleador2Form(initial={'nombre': empleador.empresa.nombre,
                    #                                'ruc': empleador.empresa.ruc,
                    #                                'provincia': empleador.empresa.provincia,
                    #                                'fechainicio': empleador.empresa.fechainicio,
                    #                                'fechafin': empleador.empresa.fechafin,
                    #                                'direccion': empleador.empresa.direccion,
                    #                                'telefonos': empleador.empresa.telefonos,
                    #                                'nombres': empleador.persona.nombres,
                    #                                'apellido1': empleador.persona.apellido1,
                    #                                'apellido2': empleador.persona.apellido2,
                    #                                'cargo': empleador.cargo,
                    #                                'email': empleador.persona.email})
                    # data['form'] = form
                    # data['empleador'] = empleador
                    # return render(request, "adm_ofertalaboral/editempleador.html", data)
                    empresa = EmpresaEmpleadora.objects.get(pk=request.GET['id'])
                    # empleador = Empleador.objects.get(empresa=empresa)
                    data['empresa'] = empresa
                   # data['empleador'] = empleador
                    form = EmpleadorForm(initial={'nombre': empresa.nombre,
                                                'ruc': empresa.ruc,
                                                'nombrecorto':empresa.nombrecorto,
                                                'sectoreconomico':empresa.sectoreconomico,
                                                'pais': empresa.pais,
                                                'provincia': empresa.provincia,
                                                # 'tipoconvenio': empresa.tipoconvenio,
                                                'direccion': empresa.direccion,
                                                'telefonos': empresa.telefonos,
                                                'representante': empresa.representante,
                                                # 'fechainicio': empresa.fechainicio,
                                                # 'fechafin': empresa.fechafin,
                                                'email': empresa.email,
                                                # 'convenio': empresa.convenio,
                                                'observacion': empresa.objetivo,
                                                'cargo' : empresa.cargo,
                                              #  'contacto' : empleador.persona,
                                                'telefonoconv': empresa.telefonoconv
                                                })
                    form.editar(empresa.pais)
                    # if empresa.convenio==False:
                    #     form.bloquear()
                    data['form'] = form
                    return render(request,"adm_convenioempresa/editempresa.html", data)
                except Exception as ex:
                    pass

            elif action == 'delempresa':
                try:
                    data['title'] = u'Eliminar empresa empleadora'
                    data['empresa'] = EmpresaEmpleadora.objects.get(pk=int(request.GET['id']))
                    return render(request, "adm_convenioempresa/delempresa.html", data)
                except Exception as ex:
                    pass

            elif action == 'reporteEmpresas':
                try:
                    __author__ = 'Unemi'
                    output = io.BytesIO()
                    workbook = xlsxwriter.Workbook(output)
                    ws = workbook.add_worksheet('empresas')
                    formatoceldatitulo = workbook.add_format({'text_wrap': True, 'bg_color': 'silver', 'align': 'center'})
                    formatoceldacenter = workbook.add_format({'valign': 'vcenter'})

                    ws.set_column(0, 20, 50)
                    ws.write('A1', 'TIPO CONVENIO ', formatoceldatitulo)
                    ws.write('B1', 'CONVENIO', formatoceldatitulo)
                    ws.write('C1', 'TIPO INSTITUCION', formatoceldatitulo)
                    ws.write('D1', 'SECTOR ECONOMICO', formatoceldatitulo)
                    ws.write('E1', 'CODIGO EMPRESA', formatoceldatitulo)
                    ws.write('F1', 'EMPRESA', formatoceldatitulo)
                    ws.write('G1', 'RUC', formatoceldatitulo)
                    ws.write('H1', 'EMAIL', formatoceldatitulo)
                    ws.write('I1', 'AUTORIZADA ', formatoceldatitulo)
                    ws.write('J1', 'REPRESENTANTE ', formatoceldatitulo)
                    ws.write('K1', 'CARGO', formatoceldatitulo)
                    ws.write('L1', 'TELEFONOS', formatoceldatitulo)
                    ws.write('M1', 'TELEFONO CONVENCIONAL', formatoceldatitulo)
                    ws.write('N1', 'DIRECCION', formatoceldatitulo)
                    ws.write('O1', 'PARROQUIA', formatoceldatitulo)
                    ws.write('P1', 'CANTON', formatoceldatitulo)
                    ws.write('Q1', 'PROVINCIA', formatoceldatitulo)
                    ws.write('R1', 'PAIS', formatoceldatitulo)
                    ws.write('S1', 'REFERENCIA', formatoceldatitulo)
                    ws.write('T1', 'OBJETIVO ', formatoceldatitulo)

                    empresas = EmpresaEmpleadora.objects.filter( status= True)
                    row_num = 1
                    for empresa in empresas:
                        if empresa.nombre is None:
                            nombre = ""
                        else:
                            nombre = empresa.nombre

                        if empresa.ruc is None or empresa.ruc == '':
                            ruc = ""
                        else:
                            ruc = empresa.ruc

                        if empresa.pais is None:
                            pais = ""
                        else:
                            pais = empresa.pais

                        if empresa.provincia is None:
                            provincia = ""
                        else:
                            provincia = empresa.provincia

                        if empresa.direccion is None or empresa.direccion == '':
                            direccion = ""
                        else:
                            direccion = empresa.direccion.upper()

                        if empresa.telefonos is None or empresa.telefonos == '':
                            telefonos = ""
                        else:
                            telefonos = empresa.telefonos

                        if empresa.telefonoconv is None:
                            telefonoconv = ""
                        else:
                            telefonoconv = empresa.telefonoconv

                        if empresa.autorizada:
                            autorizada = "SI"
                        else:
                            autorizada = "NO"

                        if empresa.representante is None or empresa.representante == '':
                            representante = ""
                        else:
                            representante = empresa.representante.upper()

                        if empresa.cargo is None or empresa.cargo == '':
                            cargo = ""
                        else:
                            cargo = empresa.cargo.upper()

                        if empresa.objetivo is None:
                            objetivo = ""
                        else:
                            objetivo = empresa.objetivo.upper()

                        if empresa.tipoconvenio is None:
                            tipoconvenio = ""
                        else:
                            tipoconvenio = empresa.tipoconvenio

                        if empresa.convenio:
                            convenio = "SI"
                        else:
                            convenio = "NO"

                        if empresa.tipoinstitucion is None:
                            tipoinstitucion = ""
                        else:
                            tipoinstitucion = empresa.get_tipoinstitucion_display()

                        if empresa.sectoreconomico is None:
                            sectoreconomico = ""
                        else:
                            sectoreconomico = empresa.get_sectoreconomico_display()

                        if empresa.parroquia is None:
                            parroquia = ""
                        else:
                            parroquia = empresa.parroquia

                        if empresa.referencia is None:
                            referencia = ""
                        else:
                            referencia = empresa.referencia

                        if empresa.email is None or empresa.email == '':
                            email = ""
                        else:
                            email = empresa.email

                        if empresa.canton is None:
                            canton = ""
                        else:
                            canton = empresa.canton

                        if empresa.pk is None:
                            codigoempresa = ""
                        else:
                            codigoempresa = empresa.pk

                        ws.write(row_num, 0, str(tipoconvenio), formatoceldacenter)
                        ws.write(row_num, 1, str(convenio), formatoceldacenter)
                        ws.write(row_num, 2, str(tipoinstitucion), formatoceldacenter)
                        ws.write(row_num, 3, str(sectoreconomico), formatoceldacenter)
                        ws.write(row_num, 4, int(codigoempresa), formatoceldacenter)
                        ws.write(row_num, 5, str(nombre), formatoceldacenter)
                        ws.write(row_num, 6, str(ruc), formatoceldacenter)
                        ws.write(row_num, 7, str(email), formatoceldacenter)
                        ws.write(row_num, 8, str(autorizada), formatoceldacenter)
                        ws.write(row_num, 9, str(representante), formatoceldacenter)
                        ws.write(row_num, 10, str(cargo), formatoceldacenter)
                        ws.write(row_num, 11, str(telefonos), formatoceldacenter)
                        ws.write(row_num, 12, str(telefonoconv), formatoceldacenter)
                        ws.write(row_num, 13, str(direccion), formatoceldacenter)
                        ws.write(row_num, 14, str(parroquia), formatoceldacenter)
                        ws.write(row_num, 15, str(canton), formatoceldacenter)
                        ws.write(row_num, 16, str(provincia), formatoceldacenter)
                        ws.write(row_num, 17, str(pais), formatoceldacenter)
                        ws.write(row_num, 18, str(referencia), formatoceldacenter)
                        ws.write(row_num, 19, str(objetivo), formatoceldacenter)






                        row_num += 1

                    workbook.close()
                    output.seek(0)
                    filename = 'reporte_empresas' + random.randint(1, 10000).__str__() + '.xlsx'
                    response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    response['Content-Disposition'] = 'attachment; filename=%s' % filename
                    return response



                except Exception as ex:
                    pass

            elif action == 'autorizar':
                try:
                    empresa = EmpresaEmpleadora.objects.get(pk=request.GET['id'])
                    empresa.autorizada = True
                    empresa.save(request)
                    log(u'Autorizo empresa empleadora: %s' % empresa, request, "edit")
                    empleador = empresa.empleador_set.all()[0]
                    send_html_mail("Registro a bolsa laboral autorizado", "emails/autorizacionnuevoregistro.html", {'sistema': request.session['nombresistema'], 'e': empleador, 't': miinstitucion(), 'modelo': MODELO_EVALUACION, 'dominio': EMAIL_DOMAIN}, empleador.persona.lista_emails_envio(), [])
                    return HttpResponseRedirect("/adm_convenioempresa?action=empresa&id" + request.GET['id'])
                except Exception as ex:
                    return HttpResponseRedirect("/adm_convenioempresa?action=empresa")

            elif action == 'desautorizar':
                try:
                    empresa = EmpresaEmpleadora.objects.get(pk=request.GET['id'])
                    empresa.autorizada = False
                    empresa.save()
                    log(u'Desautoriza empresa empleadora: %s' % empresa, request, "edit")
                    return HttpResponseRedirect("/adm_convenioempresa?action=empresa&id" + request.GET['id'])
                except Exception as ex:
                    transaction.set_rollback(True)
                    return HttpResponseRedirect("/adm_convenioempresa?action=empresa")

            elif action == 'resetear':
                try:
                    data['title'] = u'Resetear clave'
                    empleador = Persona.objects.get(pk=int(request.GET['id']))
                    data['empleador'] = empleador
                    return render(request, "adm_convenioempresa/resetear.html", data)
                except Exception as ex:
                    pass

            elif action == 'addactividad':
                try:
                    data['filtro'] = filtro = ConvenioEmpresa.objects.get(pk=int(request.GET['id']))
                    data['fecha'] = request.GET['fecha']
                    data['id'] = request.GET['id']
                    form = ActividadConvenioForm()
                    data['form2'] = form
                    template = get_template("adm_convenioempresa/modal/actividad.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'editactividad':
                try:
                    data['actividad'] = actividad = ActividadConvenio.objects.get(pk=int(request.GET['id']))
                    data['id'] = request.GET['id']
                    form = ActividadConvenioForm(initial=model_to_dict(actividad))
                    data['form2'] = form
                    template = get_template("adm_convenioempresa/modal/actividad.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'notificar':
                try:
                    form = NotificarManualForm()
                    data['form2'] = form
                    template = get_template("adm_convenioempresa/modal/notificacion.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'notificarind':
                try:
                    form = NotificacionconvenioIndividualForm()
                    data['form2'] = form
                    data['id'] = request.GET['id']
                    template = get_template("adm_convenioempresa/modal/notificacion.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'addnotificacion':
                try:
                    form = NotificacionconvenioForm()
                    data['form2'] = form
                    template = get_template("adm_convenioempresa/modal/notificacion.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'editnotificacion':
                try:
                    data['actividad'] = notificar = NotificacionConvenio.objects.get(pk=int(request.GET['id']))
                    data['id'] = request.GET['id']
                    form = NotificacionconvenioForm(initial=model_to_dict(notificar))
                    data['form2'] = form
                    template = get_template("adm_convenioempresa/modal/notificacion.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'edittiposolicitud':
                try:
                    tipo = MovilidadTipoSolicitud.objects.get(pk=request.GET['id'])
                    data['tipo'] = tipo.nombre
                    return JsonResponse({"result": "ok", 'tipo': data['tipo']})
                except Exception as ex:
                    pass

            elif action == 'deltiposolicitud':
                try:
                    data['title'] = u'Eliminar tipo de solicitud'
                    data['tipo'] = MovilidadTipoSolicitud.objects.get(pk=int(request.GET['id']))
                    return render(request, "adm_convenioempresa/deltiposolicitud.html", data)
                except Exception as ex:
                    pass

            elif action == 'tiposolicitud':
                try:
                    data['title'] = u'Tipos de solicitud'
                    search = None
                    ids = None
                    if 's' in request.GET:
                        search = request.GET['s'].strip()
                        ss = search.split(' ')
                        if len(ss) == 1:
                            solicitud = MovilidadTipoSolicitud.objects.filter(Q(nombre__icontains=search),
                                                                   Q(status=True)).distinct()
                        elif len(ss) == 2:
                            solicitud = MovilidadTipoSolicitud.objects.filter(Q(nombre__icontains=ss[0]),
                                                                   Q(nombre__icontains=ss[1]),
                                                                   Q(status=True)).distinct()
                        elif len(ss) == 3:
                            solicitud = MovilidadTipoSolicitud.objects.filter(Q(nombre__icontains=ss[0]),
                                                                   Q(nombre__icontains=ss[1]),
                                                                   Q(nombre__icontains=ss[2]),
                                                                   Q(status=True)).distinct()
                        else:
                            solicitud = MovilidadTipoSolicitud.objects.filter(Q(nombre__icontains=ss[0]),
                                                                   Q(nombre__icontains=ss[1]),
                                                                   Q(nombre__icontains=ss[2]),
                                                                   Q(nombre__icontains=ss[3]),
                                                                   Q(status=True)).distinct()
                    else:
                        solicitud = MovilidadTipoSolicitud.objects.filter(status=True).order_by('id')
                    paging = MiPaginador(solicitud, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['search'] = search if search else ""
                    data['ids'] = ids if ids else ""
                    data['tiposolicitud'] = page.object_list
                    return render(request, "adm_convenioempresa/viewtiposolicitud.html", data)
                except Exception as ex:
                    pass

            elif action == 'editbaselegal':
                try:
                    data['title'] = u'Editar base legal'
                    base = MovilidadBaseLegal.objects.get(pk=request.GET['id'])
                    data['base'] = base.descripcion
                    return JsonResponse(
                        {"result": "ok", 'title': data['title'], 'base': data['base']})
                except Exception as ex:
                    pass

            elif action == 'delbaselegal':
                try:
                    data['title'] = u'Eliminar base legal'
                    data['base'] = MovilidadBaseLegal.objects.get(pk=int(request.GET['id']))
                    return render(request, "adm_convenioempresa/delbaselegal.html", data)
                except Exception as ex:
                    pass

            elif action == 'baselegal':
                try:
                    data['title'] = u'Base legal'
                    search = None
                    ids = None
                    if 's' in request.GET:
                        search = request.GET['s'].strip()
                        ss = search.split(' ')
                        if len(ss) == 1:
                            solicitud = MovilidadBaseLegal.objects.filter(Q(descripcion__icontains=search),
                                                                              Q(status=True)).distinct()
                        elif len(ss) == 2:
                            solicitud = MovilidadBaseLegal.objects.filter(Q(descripcion__icontains=ss[0]),
                                                                              Q(descripcion__icontains=ss[1]),
                                                                              Q(status=True)).distinct()
                        elif len(ss) == 3:
                            solicitud = MovilidadBaseLegal.objects.filter(Q(descripcion__icontains=ss[0]),
                                                                              Q(descripcion__icontains=ss[1]),
                                                                              Q(descripcion__icontains=ss[2]),
                                                                              Q(status=True)).distinct()
                        else:
                            solicitud = MovilidadBaseLegal.objects.filter(Q(descripcion__icontains=ss[0]),
                                                                              Q(descripcion__icontains=ss[1]),
                                                                              Q(descripcion__icontains=ss[2]),
                                                                              Q(descripcion__icontains=ss[3]),
                                                                              Q(status=True)).distinct()
                    else:
                        solicitud = MovilidadBaseLegal.objects.filter(status=True).order_by('id')
                    paging = MiPaginador(solicitud, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['search'] = search if search else ""
                    data['ids'] = ids if ids else ""
                    data['bases'] = page.object_list
                    return render(request, "adm_convenioempresa/viewbaselegal.html", data)
                except Exception as ex:
                    pass

            elif action == 'edittipoestancia':
                try:
                    data['title'] = u'Editar tipo de estancia'
                    base = MovilidadTipoEstancia.objects.get(pk=request.GET['id'])
                    data['base'] = base.descripcion
                    return JsonResponse(
                        {"result": "ok", 'title': data['title'], 'base': data['base']})
                except Exception as ex:
                    pass

            elif action == 'deltipoestancia':
                try:
                    data['title'] = u'Eliminar tipo de estancia'
                    data['tipo'] = MovilidadTipoEstancia.objects.get(pk=int(request.GET['id']))
                    return render(request, "adm_convenioempresa/deltipoestancia.html", data)
                except Exception as ex:
                    pass

            elif action == 'tipoestancia':
                try:
                    data['title'] = u'Tipo de estancia'
                    search = None
                    ids = None
                    if 's' in request.GET:
                        search = request.GET['s'].strip()
                        ss = search.split(' ')
                        if len(ss) == 1:
                            tipos = MovilidadTipoEstancia.objects.filter(Q(descripcion__icontains=search),
                                                                              Q(status=True)).distinct()
                        elif len(ss) == 2:
                            tipos = MovilidadTipoEstancia.objects.filter(Q(descripcion__icontains=ss[0]),
                                                                              Q(descripcion__icontains=ss[1]),
                                                                              Q(status=True)).distinct()
                        elif len(ss) == 3:
                            tipos = MovilidadTipoEstancia.objects.filter(Q(descripcion__icontains=ss[0]),
                                                                              Q(descripcion__icontains=ss[1]),
                                                                              Q(descripcion__icontains=ss[2]),
                                                                              Q(status=True)).distinct()
                        else:
                            tipos = MovilidadTipoEstancia.objects.filter(Q(descripcion__icontains=ss[0]),
                                                                              Q(descripcion__icontains=ss[1]),
                                                                              Q(descripcion__icontains=ss[2]),
                                                                              Q(descripcion__icontains=ss[3]),
                                                                              Q(status=True)).distinct()
                    else:
                        tipos = MovilidadTipoEstancia.objects.filter(status=True).order_by('id')
                    paging = MiPaginador(tipos, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['search'] = search if search else ""
                    data['ids'] = ids if ids else ""
                    data['tipos'] = page.object_list
                    return render(request, "adm_convenioempresa/viewtipoestancia.html", data)
                except Exception as ex:
                    pass


            if action == 'addsolicitud':
                try:
                    data['title'] = u'Adicionar Solicitud'
                    data['form'] = MovilidadSolicitudForm()
                    return render(request, "adm_convenioempresa/addsolicitud.html", data)
                except Exception as ex:
                    pass


            elif action == 'solicitudes':
                try:
                    data['title'] = u'Solicitudes'
                    search = None
                    ids = None
                    if 'se' in request.GET:
                        search = request.GET['se']
                        empresa = EmpresaEmpleadora.objects.filter(
                            Q(nombre__icontains=search) | Q(ruc__icontains=search))
                    elif 'ide' in request.GET:
                        ids = request.GET['ide']
                        empresa = EmpresaEmpleadora.objects.filter(id=ids)
                    else:
                        empresa = EmpresaEmpleadora.objects.filter(status=True).order_by('-convenio', 'nombre')
                    paging = MiPaginador(empresa.filter(status=True), 25)
                    p = 1
                    try:
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        page = paging.page(p)
                    except Exception as ex:
                        page = paging.page(p)
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['search'] = search if search else ""
                    data['ids'] = ids if ids else ""
                    data['empresas'] = page.object_list
                    data['clave'] = DEFAULT_PASSWORD
                    return render(request, "adm_convenioempresa/viewsolicitud.html", data)
                except Exception as ex:
                    pass

            elif action == 'reporteactividades':
                try:
                    convenio = ConvenioEmpresa.objects.get(pk=request.GET['id'])
                    actividades = convenio.actividadconvenio_set.filter(status=True)
                    __author__ = 'Unemi'

                    output = io.BytesIO()
                    workbook = xlsxwriter.Workbook(output)
                    ws = workbook.add_worksheet('plantilla_')
                    ws.set_column(0, 100, 60)

                    formatoceldagris = workbook.add_format(
                        {'align': 'center', 'border': 1, 'text_wrap': True, 'fg_color': '#B6BFC0'})
                    formatocelda = workbook.add_format(
                        {'align': 'center', 'border': 1, 'text_wrap': True})


                    ws.write(0, 0, 'CONVENIO')
                    ws.merge_range('B1:D1', str(convenio))
                    ws.write(1, 0, 'USUARIO', formatoceldagris)
                    ws.write(1, 1, 'FECHA', formatoceldagris)
                    ws.write(1, 2, 'DETALLE', formatoceldagris)
                    ws.write(1, 3, 'TIENE ARCHIVO', formatoceldagris)
                    cont=2
                    for actividad in actividades:
                        ws.write(cont,0, str(actividad.usuario_creacion),formatocelda)
                        ws.write(cont,1, str(actividad.fecha),formatocelda)
                        ws.write(cont,2, str(actividad.actividad),formatocelda)
                        ws.write(cont,3, str('SI' if actividad.archivo else 'NO' ),formatocelda)
                        cont+=1
                    workbook.close()
                    output.seek(0)
                    filename = 'reporte_actividades.xlsx'
                    response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    response['Content-Disposition'] = 'attachment; filename=%s' % filename
                    return response
                except Exception as ex:
                    pass

            elif action == 'reporteactividadesres':
                try:
                    responsable = Persona.objects.get(pk=request.GET['responsable'])
                    convenios = ConvenioEmpresa.objects.filter(status=True,departamentoresponsable__responsable=responsable)
                    __author__ = 'Unemi'

                    output = io.BytesIO()
                    workbook = xlsxwriter.Workbook(output)
                    ws = workbook.add_worksheet('actividades')
                    ws.set_column(0, 100, 60)

                    formatoceldagris = workbook.add_format(
                        {'align': 'center', 'border': 1, 'text_wrap': True, 'fg_color': '#B6BFC0'})
                    formatocelda = workbook.add_format(
                        {'align': 'center', 'border': 1, 'text_wrap': True})
                    ws.write(1, 0, 'CONVENIO', formatoceldagris)
                    ws.write(1, 1, 'USUARIO', formatoceldagris)
                    ws.write(1, 2, 'FECHA', formatoceldagris)
                    ws.write(1, 3, 'DETALLE', formatoceldagris)
                    ws.write(1, 4, 'TIENE ARCHIVO', formatoceldagris)
                    cont = 2

                    for convenio in convenios:
                        if convenio.tiene_actividades():
                            val = convenio.cantidad_actividades()-1
                            ws.merge_range('A%s:A%s' % (cont + 1, cont + 1 + val), str(convenio), formatocelda)
                            for actividad in convenio.actividades():

                                ws.write(cont,1, str(actividad.usuario_creacion),formatocelda)
                                ws.write(cont,2, str(actividad.fecha),formatocelda)
                                ws.write(cont,3, str(actividad.actividad),formatocelda)
                                ws.write(cont,4, str('SI' if actividad.archivo else 'NO' ),formatocelda)
                                cont+=1
                    workbook.close()
                    output.seek(0)
                    filename = 'reporte_actividades.xlsx'
                    response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    response['Content-Disposition'] = 'attachment; filename=%s' % filename
                    return response
                except Exception as ex:
                    pass
            #NOTIFICACIONES
            elif action == 'notificacion':
                try:
                    data['title'] = u'Notificaciones'
                    search = None
                    ids = None
                    if 's' in request.GET:
                        search = request.GET['s'].strip()
                        ss = search.split(' ')
                        if len(ss) == 1:
                            notificar = NotificacionConvenio.objects.filter(Q(nombre__icontains=search), Q(status=True)).distinct()
                        elif len(ss) == 2:
                            notificar = NotificacionConvenio.objects.filter(Q(nombre__icontains=ss[0]), Q(nombre__icontains=ss[1]), Q(status=True)).distinct()
                        elif len(ss) == 3:
                            notificar = NotificacionConvenio.objects.filter(Q(nombre__icontains=ss[0]), Q(nombre__icontains=ss[1]), Q(nombre__icontains=ss[2]), Q(status=True)).distinct()
                        else:
                            notificar = NotificacionConvenio.objects.filter(Q(nombre__icontains=ss[0]), Q(nombre__icontains=ss[1]), Q(nombre__icontains=ss[2]), Q(nombre__icontains=ss[3]), Q(status=True)).distinct()
                    else:
                        notificar = NotificacionConvenio.objects.filter(status=True)
                    paging = MiPaginador(notificar, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['search'] = search if search else ""
                    data['ids'] = ids if ids else ""
                    data['notificaciones'] = page.object_list
                    return render(request, "adm_convenioempresa/viewnotificacion.html", data)
                except Exception as ex:
                    pass


            elif action == 'word':
                try:

                    output = io.BytesIO()
                    document = Document()
                    section = document.sections[0]
                    document.styles['Normal'].font.name='Times New Roman'
                    header = section.header
                    encabezado = header.paragraphs[0]
                    tabla = header.add_table(rows=2, cols=3,width=Inches(6.25))
                    celdas = tabla.rows[0].cells
                    tabla.columns[0].width=Cm(2)
                    tabla.columns[1].width=Cm(10)

                    tabla.cell(0, 0).paragraphs[0].add_run().add_picture("media/reportes/encabezados_pies/logo.png", width=Inches(1))
                    tabla.cell(1, 0).paragraphs[0].add_run().add_picture("media/reportes/encabezados_pies/membreteinformesevolucion.jpg", width=Inches(6))
                    tabla.cell(1,0).merge(tabla.cell(1,1))
                    tabla.cell(1,0).merge(tabla.cell(1,2))

                    celdas[1].text = 'INFORME TÉCNICO \n INSTITUCIONAL'
                    celfont = celdas[1].paragraphs[0].runs[0].font
                    celfont.size = Pt(14)
                    celdas[1].paragraphs[0].style.font.cs_bold=True
                    celdas[2].text = 'N° \nFECHA DE EMISIÓN: %s' % (str(datetime.now().date()))
                    celfont = celdas[2].paragraphs[0].runs[0].font
                    celfont.size = Pt(8)
                    tablad = document.add_table(3,3)
                    tablad.columns[0].width = Cm(2)
                    tablad.style='TableGrid'

                    tablad.cell(0,0).text = 'Para:\n \n '
                    tablad.cell(1,0).text = 'Elaborado por: \n \n'
                    tablad.cell(0,2).text = 'Contenido\n ' \
                                            '1. Antecedentes\n 2. Motivación jurídica \n 3. Motivación técnica\n ' \
                                            '4. Conclusiones\n 5. Recomendaciones\n 6. Anexos'
                    tablad.cell(2,0).text = 'Objeto: \n \n'
                    tablad.cell(2,1).text = 'Informar acerca de los resultados de ejecución de los convenios vigentes en la institución'
                    tablad.cell(2,1).merge(tablad.cell(2,2))
                    tablad.cell(0,2).merge(tablad.cell(1,2))
                    tablad.columns[1].width = Cm(10)
                    tablad.columns[2].width = Cm(4)
                    document.add_heading('  Antecedentes', level=1)

                    p = document.add_paragraph('El informe de resultados corresponde a la evaluación de convenios vigentes hasta la presente fecha '
                                               'elaborado por este departamento con la finalidad de que posteriormente se realice socialización del mismo '
                                               'con las diversas Unidades Organizaciones correspondientes y que se ejecuten para el desarrollo de nuestra '
                                               'institución y de la comunidad universitaria.')
                    p = document.add_paragraph('De acuerdo al Plan Operativo Anual 2021 del Departamento de Relaciones Interinstitucionales, '
                        'el mismo que establece los siguientes objetivos:')

                    p = document.add_paragraph('Objetivo Estratégico: Consolidar la oferta académica en función del desarrollo regional y los desafíos de la UNEMI.\n'
                        'Objetivo Operativo: Promover acercamientos interinstitucionales a nivel nacional e internacional que contribuyan al desarrollo académico y científico de la Universidad.\n'
                        'Actividad: Campo: Coordinar con las Unidades Académicas y Administrativas la ejecución de convenios institucionales vigentes.')
                    p = document.add_paragraph('Medio de verificación: Informe de resultados.')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    document.add_heading('  Motivación Jurídica', level=1)
                    p = document.add_paragraph(
                        'Que, en el Capítulo 1, Artículo 4, Numeral 1 y 4 de las "Objetivos Estratégicos Institucionales".\n'
                        '1. Fortalecer una educación de calidad, desarrollando saberes y consolidando la identidad cultural de la sociedad.\n '
                        '4. Mejorar la administración de los recursos y del talento humano con eficiencia y eficacia '
                        'PLAN NACIONAL DEL BUEN VIVIR, a los objetivos N° 4 que es el de fortalecer las capacidades y potencialidades de la ciudadanía, en concordancia ' \
                        'con Constitución de la República del Ecuador en su Art. 350 que establece "(...) están llamados a consolidar las capacidades y oportunidades de ' \
                        'la población y a formar académica y profesionalmente a las personas bajo una visión científica y humanista, que incluye los saberes y las culturas ' \
                        'de nuestro pueblo (...)".')
                    p.style.font.size = Pt(10)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    document.add_heading('  Motivación Técnica (Administrativa o Académica)', level=1)
                    p = document.add_paragraph(
                        ' El informe muestra la cantidad de convenios vigentes, ejecutados y sin ejecutar que mantiene nuestra institución. Además de su detalle como lo es:')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    document.add_paragraph('Nombre de la institución.', style='List Bullet')
                    document.add_paragraph('Fecha de vigencia.', style='List Bullet')
                    document.add_paragraph('Personas responsables del convenio.', style='List Bullet')
                    document.add_paragraph('Objeto del convenio.', style='List Bullet')
                    document.add_paragraph('Tipo de convenio.', style='List Bullet')
                    document.add_paragraph('Ejecución.', style='List Bullet')
                    fecha_actual = datetime.now().date()
                    convenios = ConvenioEmpresa.objects.filter(status=True,departamentoresponsable__isnull=False,fechafinalizacion__gte=fecha_actual).distinct()
                    conveniosid = ConvenioEmpresa.objects.values_list('id').filter(status=True,departamentoresponsable__isnull=False,fechafinalizacion__gte=fecha_actual).distinct()
                    ejecutados = 0
                    for con in convenios:
                        if con.tiene_actividades():
                            ejecutados+=1

                    tablac = document.add_table(1,8)
                    tablac.style = 'TableGrid'
                    tablac.cell(0, 0).text = 'INSTITUCIÓN '
                    tablac.cell(0, 0).paragraphs[0].add_run()
                    tablac.cell(0, 0).paragraphs[0].runs[0].font.size = Pt(6)
                    tablac.cell(0, 1).text = 'FECHA DE INICIO '
                    tablac.cell(0, 1).paragraphs[0].add_run()
                    tablac.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(6)
                    tablac.cell(0, 2).text = 'FECHA FIN'
                    tablac.cell(0, 2).paragraphs[0].add_run()
                    tablac.cell(0, 2).paragraphs[0].runs[0].font.size = Pt(6)
                    tablac.cell(0, 3).text = 'RESPONSABLE CONTRAPARTE'
                    tablac.cell(0, 3).paragraphs[0].add_run()
                    tablac.cell(0, 3).paragraphs[0].runs[0].font.size = Pt(6)
                    tablac.cell(0, 4).text = 'RESPONSABLE UNEMI'
                    tablac.cell(0, 4).paragraphs[0].add_run()
                    tablac.cell(0, 4).paragraphs[0].runs[0].font.size = Pt(6)
                    tablac.cell(0, 5).text = 'OBJETIVO'
                    tablac.cell(0, 5).paragraphs[0].add_run()
                    tablac.cell(0, 5).paragraphs[0].runs[0].font.size = Pt(6)
                    tablac.cell(0, 6).text = 'TIPO DE CONVENIO'
                    tablac.cell(0, 6).paragraphs[0].add_run()
                    tablac.cell(0, 6).paragraphs[0].runs[0].font.size = Pt(6)
                    tablac.cell(0, 7).text = 'EJECUCIÓN'
                    tablac.cell(0, 7).paragraphs[0].add_run()
                    tablac.cell(0, 7).paragraphs[0].runs[0].font.size = Pt(6)
                    cont=1
                    for convenio in convenios:
                        txtresp=''
                        txtact=''
                        for resp in convenio.departamentoresponsable.all():
                            txtresp+=str(resp)+'\n'
                        for act in convenio.actividades():
                            txtact+=str(act)+'\n'

                        tablac.add_row().cells
                        tablac.cell(cont,0).text = str(convenio.empresaempleadora)
                        tablac.cell(cont, 0).paragraphs[0].add_run()
                        tablac.cell(cont, 0).paragraphs[0].runs[0].font.size = Pt(6)
                        tablac.cell(cont,1).text = str(convenio.fechainicio)
                        tablac.cell(cont, 1).paragraphs[0].add_run()
                        tablac.cell(cont, 1).paragraphs[0].runs[0].font.size = Pt(6)
                        tablac.cell(cont,2).text = str(convenio.fechafinalizacion)
                        tablac.cell(cont, 2).paragraphs[0].add_run()
                        tablac.cell(cont, 2).paragraphs[0].runs[0].font.size = Pt(6)
                        tablac.cell(cont,3).text = str(convenio.responsableexterno)
                        tablac.cell(cont, 3).paragraphs[0].add_run()
                        tablac.cell(cont, 3).paragraphs[0].runs[0].font.size = Pt(6)
                        tablac.cell(cont,4).text = txtresp
                        tablac.cell(cont, 4).paragraphs[0].add_run()
                        tablac.cell(cont, 4).paragraphs[0].runs[0].font.size = Pt(6)
                        tablac.cell(cont,5).text = str(convenio.objetivo)
                        tablac.cell(cont, 5).paragraphs[0].add_run()
                        tablac.cell(cont, 5).paragraphs[0].runs[0].font.size = Pt(6)
                        tablac.cell(cont,6).text = str(convenio.tipoconvenio)
                        tablac.cell(cont, 6).paragraphs[0].add_run()
                        tablac.cell(cont, 6).paragraphs[0].runs[0].font.size = Pt(6)
                        tablac.cell(cont,7).text = txtact
                        tablac.cell(cont, 7).paragraphs[0].add_run()
                        tablac.cell(cont, 7).paragraphs[0].runs[0].font.size = Pt(6)
                        cont+=1
                    p = document.add_paragraph('Este es el resultado de las designaciones a las responsables realizadas '
                                               'mediante memorandos, notificaciones y correos electrónicos solicitando la ejecución de los convenios además del detalle de la importancia de mantener vigentes convenios que se encuentran sin ejecución, a continuación, se detallan los memorandos enviados:')
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    tablan = document.add_table(2, 4)
                    tablan.style = 'TableGrid'
                    tablan.cell(0, 0).text = 'FECHA '
                    tablan.cell(0, 1).text = 'MEDIO '
                    tablan.cell(0, 2).text = 'DIRIGIDO A '
                    tablan.cell(0, 3).text = 'ACTIVIDAD REALIZADA'
                    contn=1

                    for noti in DetalleNotificacionCovenio.objects.filter(status=True,convenioempresa_id__in=conveniosid):
                        tablan.add_row().cells
                        tablan.cell(contn,0).text = str(noti.fecha)
                        tablan.cell(contn,1).text = 'NOTIFICACIÓN DEL SISTEMA SGA'
                        tablan.cell(contn,2).text = str(noti.convenioempresa)
                        tablan.cell(contn,3).text = str(noti.info_notificacion())

                    document.add_heading('  Conclusiones', level=1)
                    p = document.add_paragraph('Realizada la revisión de ejecución de convenios vigentes en UNEMI se cuenta con:')

                    document.add_paragraph('La Universidad Estatal de Milagro cuenta con %s convenios vigentes, de los cuales %s convenios se han ejecutado desde su firma hasta el presente momento.' %(convenios.count(),ejecutados), style='List Bullet')
                    document.add_paragraph('Existe un total de %s convenios sin ejecución.' %(convenios.count()-ejecutados), style='List Bullet')

                    document.add_heading('  Recomendaciones', level=1)
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    p = document.add_paragraph('Se recomienda a los responsables de los convenios realizar actividades enmarcadas en los convenios tanto en los ejecutados como en los no ejecutados.')

                    document.add_heading('  Anexos', level=1)
                    tablaa = document.add_table(2, 4)
                    tablaa.style = 'TableGrid'
                    tablaa.cell(0, 0).text = 'No. '
                    tablaa.cell(0, 1).text = 'NOMBRE DEL DOCUMENTO '
                    tablaa.cell(0, 2).text = 'FECHA DE GENERACIÓNN'
                    tablaa.cell(0, 3).text = '# PÁGINAS'
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                    p = document.add_paragraph('\n')

                    tablaf = document.add_table(3, 3)
                    tablaf.style = 'TableGrid'
                    tablaf.cell(0, 0).text = 'ROL/CARGO'
                    tablaf.cell(0, 1).text = 'FECHA/HORA'
                    tablaf.cell(0, 2).text = 'FIRMA'

                    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    document.save(output)
                    response = HttpResponse(output.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                    filename = 'demo.docx'
                    response['Content-Disposition'] = 'attachment; filename=%s' % filename

                    return response
                except Exception as ex:
                    print(u'Error on line {}'.format(sys.exc_info()[-1].tb_lineno))
                    pass

            elif action == 'buscarCargo':
                try:
                    # lista = []
                    idscargos = DistributivoPersona.objects.values_list('denominacionpuesto_id', flat=True).order_by('denominacionpuesto_id').distinct()
                    filtro = Q(pk__in=idscargos, status=True)
                    if 'q' in request.GET:
                        search =request.GET['q']
                    if search:
                        filtro = filtro & Q(descripcion__icontains=search)
                    denominacionPuesto = DenominacionPuesto.objects.filter(filtro)
                    resp = [{"id": c.id, "text": c.descripcion} for c in denominacionPuesto]
                    return JsonResponse(resp, safe=False)
                except Exception as ex:
                    return JsonResponse({"result": "bad", "mensaje": u"Error al obtener los datos."})

            elif action == 'cargar_responsable':
                try:
                    lista = []
                    fecha_actual = datetime.now().date()
                    convenios = ConvenioEmpresa.objects.filter(status=True,departamentoresponsable__isnull=False,fechafinalizacion__gte=fecha_actual)
                    for convenio in convenios:
                        for responsable in convenio.departamentoresponsable.all():
                            if responsable.responsable:
                                if not buscar_dicc(lista,'id',responsable.responsable.id):
                                    lista.append({'id':responsable.responsable.id, 'nombre':responsable.responsable.nombre_completo_inverso()})
                    return JsonResponse({'result': 'ok', 'lista': lista})
                except Exception as ex:
                    return JsonResponse({"result": "bad", "mensaje": u"Error al obtener los datos."})

            elif action == 'cargarresponsables':
                try:
                    lista = []
                    responsables = DistributivoPersona.objects.filter(denominacionpuesto=request.GET['id'], status=True)
                    for responsable in responsables:
                        lista.append({'value': responsable.pk, 'text': str(responsable.persona)})
                    return JsonResponse({'result': True, 'data': lista})
                except Exception as ex:
                    return JsonResponse({"result": "bad", "mensaje": u"Error al obtener los datos."})

            elif action == 'cargarcargos':
                try:
                    pers = Persona.objects.get(id=request.GET['id'])
                    cargos = pers.mis_cargos_vigente()
                    resp = [{"id": c.id, "text": c.denominacionpuesto.descripcion} for c in cargos]
                    return JsonResponse({'result': True, 'data': resp})
                except Exception as ex:
                    pass

            elif action == 'buscarpersonas':
                try:
                    resp = filtro_persona_select(request)
                    return HttpResponse(json.dumps({'status': True, 'results': resp}))
                except Exception as ex:
                    pass

            elif action == 'buscarempresa':
                try:
                    idexc=int(request.GET.get('idexc',0))
                    q = request.GET['q'].strip()
                    filtro = Q(nombre__unaccent__icontains=q)
                    if q.isdigit():
                        filtro = Q(id=q)
                    empresas = EmpresaEmpleadora.objects.filter(filtro, status=True).exclude(nombre='').exclude(id=idexc)
                    resp = [{"id": c.id, "text": f'COD:{c.id} - {c}'} for c in empresas]
                    return JsonResponse(resp, safe=False)
                except Exception as ex:
                    pass

            elif action == 'importarconvenios':
                try:
                    template = get_template("adm_convenioempresa/modal/importarconvenios.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, 'mensaje': f'Error: {ex}'})

            elif action == 'moverregistros':
                try:
                    id = request.GET['id']
                    data['empresa'] = EmpresaEmpleadora.objects.get(id=id)
                    template = get_template("adm_convenioempresa/modal/formmoverregistros.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, 'mensaje': f'Error: {ex}'})

            return HttpResponseRedirect(request.path)
        else:
            fecha_actual = datetime.now().date()
            data['title'] = u'Convenios institucionales'
            id, tipo, estado, desde, hasta, search, filtros, url_vars = request.GET.get('id', ''), request.GET.get('tipo', ''), request.GET.get('estado', ''), request.GET.get('desde', ''), request.GET.get('hasta', ''), request.GET.get('search', ''), Q(status=True), ''
            if id:
                filtros = filtros & Q(id=id)
            if tipo:
                data['tipo'] = tipo = int(tipo)
                url_vars += "&tipo={}".format(tipo)
                if tipo == 1:
                    filtros = filtros & Q(para_practicas=True)
                if tipo == 2:
                    filtros = filtros & Q(para_pasantias=True)
                if tipo == 3:
                    filtros = filtros & Q(para_practicas=True) | Q(para_pasantias=True)
            if estado:
                data['estado'] = estado = int(estado)
                url_vars += "&estado={}".format(estado)
                if estado == 1:
                    filtros = filtros & Q(fechafinalizacion__gte=fecha_actual)
                elif estado == 2:
                    filtros = filtros & Q(fechafinalizacion__lte=fecha_actual)
            if desde:
                data['desde'] = desde
                url_vars += "&desde={}".format(desde)
                filtros = filtros & Q(fechainicio__gte=desde)
            if hasta:
                data['hasta'] = hasta
                url_vars += "&hasta={}".format(hasta)
                filtros = filtros & Q(fechafinalizacion__lte=hasta)
            if search:
                data['search'] = search
                s = search.split()
                filtros = filtros & (Q(empresaempleadora__nombre__icontains=search) |
                                     Q(tipoconvenio__nombre__icontains=search))
                url_vars += '&search={}'.format(search)

            data["url_vars"] = url_vars

            convenio = ConvenioEmpresa.objects.filter(filtros).order_by('-pk')
            paging = MiPaginador(convenio, 20)
            p = 1
            try:
                paginasesion = 1
                if 'paginador' in request.session:
                    paginasesion = int(request.session['paginador'])
                if 'page' in request.GET:
                    p = int(request.GET['page'])
                else:
                    p = paginasesion
                try:
                    page = paging.page(p)
                except:
                    p = 1
                page = paging.page(p)
            except:
                page = paging.page(p)
            request.session['paginador'] = p
            data['ids'] = id if id else ""
            data['paging'] = paging
            data['rangospaging'] = paging.rangos_paginado(p)
            data['page'] = page
            data['search'] = search if search else ""
            data['convenioempresas'] = page.object_list
            data['total'] = len(convenio)
            return render(request, "adm_convenioempresa/view.html", data)

CORREOS_ADICIONALES_ENVIOS_RESPONSABLES_INTERNOS=['vinculacion@unemi.edu.ec', 'rr.ii@unemi.edu.ec', 'vinculacion@unemi.edu.ec']
#CORREOS_ADICIONALES_ENVIOS_RESPONSABLES_INTERNOS=['jsanchezm2@unemi.edu.ec', 'jplacesc@unemi.edu.ec']

def responsables_convenio():
    pass

def enviar_email_convenio(convenio):
    persona = Persona.objects.get(id=convenio.responsableinterno.id)
    send_html_mail(subject="RESPONSABLES INTERNOS DE LOS CONVENIOS",
                   html_template="emails/texto_responsables_internos_convenios.html",
                   data={'sistema': "SGA",
                       'nombre_convenio': convenio.empresaempleadora.nombre,
                       'nombre_persona': '{} {} {}'.format(persona.nombres, persona.apellido1,
                                                           persona.apellido2)
                   },
                   recipient_list=[persona.emails()]+CORREOS_ADICIONALES_ENVIOS_RESPONSABLES_INTERNOS,
                   recipient_list_cc=[],
                   cuenta=CUENTAS_CORREOS[4][1])

def enviar_email_convenio_departamento(persona, convenio,cargo):
    persona = Persona.objects.get(id=persona.id)
    send_html_mail(subject="RESPONSABLES INTERNOS DE LOS CONVENIOS",
                   html_template="emails/texto_responsables_internos_convenios.html",
                   data={'sistema': "SGA",
                       'nombre_convenio': convenio.empresaempleadora.nombre,
                       'cargo_persona': cargo,
                       'convenio_numocas': convenio.numocas,
                       'tipo_convenio': convenio.tipoconvenio,
                       'fecha_inicio': convenio.fechainicio,
                       'fecha_fin': convenio.fechafinalizacion,
                       'fecha_ocas': convenio.fechaocas,
                       'nombre_persona': '{} {} {}'.format(persona.nombres, persona.apellido1,
                                                           persona.apellido2)
                   },
                   recipient_list=[persona.emails()]+CORREOS_ADICIONALES_ENVIOS_RESPONSABLES_INTERNOS,
                   recipient_list_cc=[],
                   cuenta=CUENTAS_CORREOS[4][1])