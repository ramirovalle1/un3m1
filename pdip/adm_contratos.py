# -*- coding: UTF-8 -*-
import datetime
import json
import random
import sys

import openpyxl
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.db import transaction
from django.db.models import Q
import xlwt
from django.forms import model_to_dict
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse
from django.template import Context
from django.template.loader import get_template
from docx import Document
from xlwt import *
from django.shortcuts import render, redirect
from decorators import secure_module
from postulaciondip.adm_postulacion import weeks_between
from postulaciondip.models import InscripcionInvitacion, HistorialInvitacion, Requisito
from sagest.models import SeccionDepartamento, OpcionSistema
from settings import EMAIL_DOMAIN, SITE_ROOT
from sga.commonviews import adduserdata
from sga.templatetags.sga_extras import encrypt
from .forms import PerfilPuestoDipForm, RequisitoPagoDipForm, CampoContratoDipForm, PlantillaContratoDipForm, \
    ProcesoPagoForm, PasoProcesoPagoForm, ContratoDipForm, ContratoMetodoPago, SecuenciaMemoActPosgradoForm, \
    GestionarContratoDipForm, ArchivoInformesForm, CertificacionPresupuestariaDipForm, ActividadesForm, \
    DocumentoContratosForm, FechaMaximoBitacoraForm, GrupoRevisionPagoContratoForm, HorarioClasesContratoForm, \
    DepartamentoPosgradoForm, GestionPosgradoForm, RecursoPresupuestarioForm, CabeceraRecursoPresupuestarioPosgradoForm, \
    DetalleRecursoPresupuestarioPosgradoForm, ItemRecursoPresupuestarioPosgradoForm, \
    ContratacionConfiguracionRequisitoForm, ContratacionGestionRequisitoForm, RequisitoContratoForm, \
    ConfiguracionGrupoPagoForm, OrdenFirmaInformeActaPagoForm, ConfiguracionActaPagoSolicitadoPorForm, \
    ConfiguracionActaPagoObjetivoForm, ConfiguracionActaPagoMarcoJuridicoReferencialForm, \
    ConfiguracionActaPagoTituloForm, ConfiguracionActaPagoParaForm, ConfiguracionActaPagoForm, \
    RequisitoPagoDipOrdenForm, ConfiguracionActaPagoRecomendacionesForm, ConfiguracionActaPagoConclusionesForm
from sga.funciones import MiPaginador, log, generar_nombre, remover_caracteres_especiales_unicode, convertir_fecha, \
    variable_valor, remover_caracteres_tildes_unicode
from sga.models import Administrativo, Persona, ProfesorMateria, DiasNoLaborable, VariablesGlobales, Carrera, Turno
from sga.funcionesxhtml2pdf import conviert_html_to_pdf_save_informe, conviert_html_to_pdf
from .models import *
from .adm_marcadas_dip import extraervalores
from django.db.models import Sum, Q, F, FloatField
from django.db.models.functions import Coalesce
from datetime import datetime, timedelta
from sagest.commonviews import secuencia_recaudacion, anio_ejercicio
import xlsxwriter
import io
import calendar


def rango_anios():
    if Contratos.objects.exists():
        inicio = datetime.now().year
        fin = Contratos.objects.order_by('anio')[0].anio
        return range(inicio, fin, -1)
    return [datetime.now().date().year]


def fecha_letra(valor):
    mes = ["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre",
           "Noviembre", "Diciembre"]
    d = int(valor[0:2])
    m = int(valor[3:5])
    a = int(valor[6:10])
    if d == 1:
        return u"al %s día del mes de %s del %s" % (numero_a_letras(d), str(mes[m - 1]), numero_a_letras(a))
    else:
        return u"a los %s días del mes de %s del %s" % (numero_a_letras(d), str(mes[m - 1]), numero_a_letras(a))


MONEDA_SINGULAR = 'dolar'
MONEDA_PLURAL = 'dolares'
CENTIMOS_SINGULAR = 'centavo'
CENTIMOS_PLURAL = 'centavos'
MAX_NUMERO = 999999999999

UNIDADES = (
    'cero',
    'uno',
    'dos',
    'tres',
    'cuatro',
    'cinco',
    'seis',
    'siete',
    'ocho',
    'nueve'
)

DECENAS = (
    'diez',
    'once',
    'doce',
    'trece',
    'catorce',
    'quince',
    'dieciseis',
    'diecisiete',
    'dieciocho',
    'diecinueve'
)

DIEZ_DIEZ = (
    'cero',
    'diez',
    'veinte',
    'treinta',
    'cuarenta',
    'cincuenta',
    'sesenta',
    'setenta',
    'ochenta',
    'noventa'
)

CIENTOS = (
    '_',
    'ciento',
    'doscientos',
    'trescientos',
    'cuatroscientos',
    'quinientos',
    'seiscientos',
    'setecientos',
    'ochocientos',
    'novecientos'
)


def numero_a_letras(numero):
    numero_entero = int(numero)
    if numero_entero > MAX_NUMERO:
        raise OverflowError('Número demasiado alto')
    if numero_entero < 0:
        return 'menos %s' % numero_a_letras(abs(numero))
    letras_decimal = ''
    parte_decimal = int(round((abs(numero) - abs(numero_entero)) * 100))
    if parte_decimal > 9:
        letras_decimal = 'punto %s' % numero_a_letras(parte_decimal)
    elif parte_decimal > 0:
        letras_decimal = 'punto cero %s' % numero_a_letras(parte_decimal)
    if numero_entero <= 99:
        resultado = leer_decenas(numero_entero)
    elif numero_entero <= 999:
        resultado = leer_centenas(numero_entero)
    elif numero_entero <= 999999:
        resultado = leer_miles(numero_entero)
    elif numero_entero <= 999999999:
        resultado = leer_millones(numero_entero)
    else:
        resultado = leer_millardos(numero_entero)
    resultado = resultado.replace('uno mil', 'un mil')
    resultado = resultado.strip()
    resultado = resultado.replace(' _ ', ' ')
    resultado = resultado.replace('  ', ' ')
    if parte_decimal > 0:
        resultado = '%s %s' % (resultado, letras_decimal)
    return resultado


def numero_a_moneda(numero):
    if '.' in numero:
        posicion = numero.split('.')
        numero_entero = int(posicion[0])
        parte_decimal = 0
        if posicion[1] != '':
            parte_decimal = int(posicion[1])
    else:
        numero_entero = int(numero)
        parte_decimal = 0
    centimos = ''
    if parte_decimal == 1:
        centimos = CENTIMOS_SINGULAR
    else:
        centimos = CENTIMOS_PLURAL
    moneda = ''
    if numero_entero == 1:
        moneda = MONEDA_SINGULAR
    else:
        moneda = MONEDA_PLURAL
    letras = numero_a_letras(numero_entero)
    letras = letras.replace('uno', 'un')
    letras_decimal = u'%s/100 DOLARES DE LOS ESTADOS UNIDOS DE NORTE AMÉRICA' % (str(parte_decimal))
    letras = u'%s %s' % (letras, letras_decimal)
    return letras


def leer_decenas(numero):
    if numero < 10:
        return UNIDADES[numero]
    decena, unidad = divmod(numero, 10)
    if unidad == 0:
        resultado = DIEZ_DIEZ[decena]
    else:
        if numero <= 19:
            resultado = DECENAS[unidad]
        elif numero <= 29:
            resultado = 'veinti%s' % UNIDADES[unidad]
        else:
            resultado = DIEZ_DIEZ[decena]
            if unidad > 0:
                resultado = '%s y %s' % (resultado, UNIDADES[unidad])
    return resultado


def leer_centenas(numero):
    centena, decena = divmod(numero, 100)
    if numero == 0:
        resultado = 'cien'
    else:
        resultado = CIENTOS[centena]
        if decena > 0:
            resultado = '%s %s' % (resultado, leer_decenas(decena))
    return resultado


def leer_miles(numero):
    millar, centena = divmod(numero, 1000)
    resultado = ''
    if millar == 1:
        resultado = ''
    if 2 <= millar <= 9:
        resultado = UNIDADES[millar]
    elif 10 <= millar <= 99:
        resultado = leer_decenas(millar)
    elif 100 <= millar <= 999:
        resultado = leer_centenas(millar)
    resultado = '%s mil' % resultado
    if centena > 0:
        resultado = '%s %s' % (resultado, leer_centenas(centena))
    return resultado


def leer_millones(numero):
    millon, millar = divmod(numero, 1000000)
    resultado = ''
    if millon == 1:
        resultado = ' un millon '
    if 2 <= millon <= 9:
        resultado = UNIDADES[millon]
    elif 10 <= millon <= 99:
        resultado = leer_decenas(millon)
    elif 100 <= millon <= 999:
        resultado = leer_centenas(millon)
    if millon > 1:
        resultado = '%s millones' % resultado
    if (millar > 0) and (millar <= 999):
        resultado = '%s %s' % (resultado, leer_centenas(millar))
    elif (millar >= 1000) and (millar <= 999999):
        resultado = '%s %s' % (resultado, leer_miles(millar))
    return resultado


def leer_millardos(numero):
    millardo, millon = divmod(numero, 1000000)
    return '%s millones %s' % (leer_miles(millardo), leer_millones(millon))


@login_required(redirect_field_name='ret', login_url='/loginsga')
@secure_module
@transaction.atomic()
def view(request):
    data = {}
    adduserdata(request, data)
    usuario = request.user
    persona = request.session['persona']
    editarCuotaContrato = False
    if request.method == 'POST':
        res_json = []
        action = request.POST['action']

        if action == 'metodoPago':
            try:
                idCuotas = request.POST.getlist('idCuota[]')
                numCuotas = request.POST.getlist('cuota[]')
                valorCuotas = request.POST.getlist('valorCuota[]')
                contrato = ContratoDip.objects.get(pk=request.POST['contratodip'])
                acumulador = 0

                if (valorCuotas):

                    for cuota in valorCuotas:
                        acumulador = acumulador + float(cuota)

                    if acumulador == contrato.rmu:

                        existeCuotasContrato = ContratoDipMetodoPago.objects.filter(
                            contratodip=contrato, status=True).exists()

                        if existeCuotasContrato:
                            # ELIMINO LOS QUE NO VENGAN EN LA LISTA
                            ContratoDipMetodoPago.objects.filter(status=True, contratodip=contrato).exclude(
                                pk__in=idCuotas).update(status=False)
                            for id, numCuota, valorCuota in zip(idCuotas, numCuotas, valorCuotas):

                                if not id == '0':
                                    # por editar
                                    cuota = ContratoDipMetodoPago.objects.get(id=id, status=True)
                                    cuota.contratodip = contrato
                                    cuota.numerocuota = numCuota
                                    cuota.valorcuota = valorCuota
                                    cuota.save()
                                else:
                                    # por editar con cuota añadida
                                    nuevaCuota = ContratoDipMetodoPago(contratodip=contrato,
                                                                       numerocuota=numCuota,
                                                                       valorcuota=valorCuota
                                                                       )
                                    nuevaCuota.save()
                            log(u'Edito las cuotas para contrato dip: %s' % cuota, request, "Edit")
                            return JsonResponse({"result": False}, safe=False)


                        else:
                            # por guardar nuevo
                            for numCuota, valorCuota in zip(numCuotas, valorCuotas):
                                metodoPago = ContratoDipMetodoPago(contratodip=contrato,
                                                                   numerocuota=numCuota,
                                                                   valorcuota=valorCuota
                                                                   )
                                metodoPago.save()
                            log(u'Adicionó cuotas para contrato dip: %s' % metodoPago, request, "add")
                            return JsonResponse({"result": False}, safe=False)

                    else:
                        return JsonResponse(
                            {"result": True,
                             "mensaje": "La suma total de todas las cuotas es mayor o no es igual al valor del contrato."},
                            safe=False)

            except Exception as ex:
                transaction.set_rollback(True)
            return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'addperfil':
            try:
                f = PerfilPuestoDipForm(request.POST)
                if f.is_valid():
                    filtro = PerfilPuestoDip(nombre=f.cleaned_data['nombre'], )
                    filtro.save(request)
                    log(u'Adicionó perfil para contrato dip: %s' % filtro, request, "add")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'addactividades':
            try:
                f = ActividadesForm(request.POST)
                if f.is_valid():
                    filtro = ActividadesPerfil(descripcion=f.cleaned_data['descripcion'], )
                    filtro.save(request)
                    log(u'Adicionó descripcion para actividades de cargos de dip: %s' % filtro, request, "add")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'addsecuencia':
            try:
                f = SecuenciaMemoActPosgradoForm(request.POST)
                if f.is_valid():
                    tipo = None
                    if PlantillaContratoDip.objects.values('id').filter(id=3).exists():
                        tipo = PlantillaContratoDip.objects.get(id=3)
                    filtro = SecuenciaMemoActividadPosgrado(secuencia=f.cleaned_data['secuencia'],
                                                            anioejercicio=secuencia_memo().anioejercicio, tipo=tipo)
                    filtro.save(request)
                    log(u'Adicionó secuencia para contrato dip: %s' % filtro, request, "add")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'editperfil':
            try:
                f = PerfilPuestoDipForm(request.POST)
                if f.is_valid():
                    filtro = PerfilPuestoDip.objects.get(pk=int(request.POST['id']))
                    filtro.nombre = f.cleaned_data['nombre']
                    filtro.save(request)
                    # filtro.actividades.clear()
                    # if 'lista_items1' in request.POST:
                    #     for elemento in json.loads(request.POST['lista_items1']):
                    #         if 'idp' in elemento:
                    #             activiada = ActividadesContratoPerfil.objecst.get(pk = int(elemento['idp']))
                    #         else:
                    #             activiada = ActividadesContratoPerfil(
                    #                 perfil = elemento['idp'],
                    #                 actividad = elemento['id'],
                    #                 obligatorio = elemento['obli']
                    #             )
                    #         activiada.save(request)
                    log(u'Editó perfil para contrato dip: %s' % filtro, request, "edit")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'actividadesext':
            try:
                filtro = ContratoDip.objects.get(pk=int(request.POST['id']))
                # filtro.actividadesextra.clear()
                # for elemento in json.loads(request.POST['lista_items1']):
                #     acti = filtro.actividadesextra.add(elemento['id'])
                # log(u'Editó  contrato dip: %s' % filtro, request, "edit")
                return JsonResponse({"result": False}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'editactvidades':
            try:
                f = ActividadesForm(request.POST)
                if f.is_valid():
                    filtro = ActividadesPerfil.objects.get(pk=int(request.POST['id']))
                    filtro.descripcion = f.cleaned_data['descripcion']
                    filtro.save(request)
                    log(u'Editó actvidad para contrato dip: %s' % filtro, request, "edit")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'editsecuencia':
            try:
                f = SecuenciaMemoActPosgradoForm(request.POST)
                if f.is_valid():
                    tipo = None
                    if PlantillaContratoDip.objects.values('id').filter(id=3).exists():
                        tipo = PlantillaContratoDip.objects.get(id=3)
                    filtro = SecuenciaMemoActividadPosgrado.objects.get(pk=int(request.POST['id']))
                    filtro.anioejercicio = secuencia_memo().anioejercicio
                    filtro.secuencia = f.cleaned_data['secuencia']
                    filtro.tipo = tipo
                    filtro.save(request)
                    log(u'Editó secuencia para contrato dip: %s' % filtro, request, "edit")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'deleteperfil':
            try:
                with transaction.atomic():
                    instancia = PerfilPuestoDip.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino Contrato: %s' % instancia, request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}

        elif action == 'deleterequisitocontratacion':
            try:
                with transaction.atomic():
                    eContratoRequisito = ContratoRequisito.objects.get(pk=int(request.POST['id']))
                    eContratoRequisito.status = False
                    eContratoRequisito.save(request)
                    log(u'Elimino ContratoRequisito: %s' % eContratoRequisito, request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'deleteactperfilcontrato':
            try:
                with transaction.atomic():
                    instancia = ActividadesContratoPerfil.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino la actividad %s del perfil %s' % (instancia.actividad, instancia.perfil), request,
                        'delete')
                    return JsonResponse({'error': False, 'id': instancia.actividad.id})
            except Exception as ex:
                transaction.set_rollback(True)
                return {'error': True, "message": "Error: {}".format(ex)}

        elif action == 'deleteactvidades':
            try:
                with transaction.atomic():
                    instancia = ActividadesPerfil.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino Actividad: %s' % instancia, request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'deletememo':
            try:
                with transaction.atomic():
                    instancia = MemoActividadPosgrado.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino Memo: %s %s' % (instancia, instancia.contrato.persona), request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'deletebitacora':
            try:
                with transaction.atomic():
                    instancia = InformeActividadJornada.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino informe de actividades y jornada laboral: %s %s' % (
                    instancia, instancia.contrato.persona), request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'deleteinftecnico':
            try:
                with transaction.atomic():
                    instancia = InformeTecnico.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino informe tecnico: #%s %s' % (instancia, instancia.contrato.persona), request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'deleteactapago':
            try:
                with transaction.atomic():
                    instancia = ActaPago.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino acta de pago: #%s %s' % (instancia, instancia.contrato.persona), request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'cambiarestado':
            try:
                with transaction.atomic():
                    instancia = HistorialPagoMes.objects.get(pk=int(request.POST['id']))
                    instancia.cancelado = False if instancia.cancelado else True
                    instancia.save(request)
                    log(u'Cambio el estado de pago: #%s %s' % (instancia, instancia.contrato.persona), request, "edit")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'delrecursopresupuestario':
            try:
                with transaction.atomic():
                    pk = request.POST.get('id', '0')
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    instancia = RecursoPresupuestarioPosgrado.objects.get(pk=pk)
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino RecursoPresupuestarioPosgrado: %s' % instancia, request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'addcampo':
            try:
                f = CampoContratoDipForm(request.POST)
                if f.is_valid():

                    if CampoContratoDip.objects.filter(identificador=(f.cleaned_data['identificador']).upper(),
                                                       status=True).exists():
                        return JsonResponse({"result": "bad", "mensaje": "El identificador ya existe."})

                    filtro = CampoContratoDip(descripcion=f.cleaned_data['descripcion'],
                                              tipo=f.cleaned_data['tipo'],
                                              identificador=f.cleaned_data['identificador'],
                                              script=f.cleaned_data['script'])
                    filtro.save(request)
                    log(u'Adicionó campo para contrato dip: %s' % filtro, request, "add")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'editcampo':
            try:
                f = CampoContratoDipForm(request.POST, request.FILES)
                if f.is_valid():
                    filtro = CampoContratoDip.objects.get(pk=int(request.POST['id']), status=True)

                    if CampoContratoDip.objects.filter(identificador=(f.cleaned_data['identificador']).upper(),
                                                       status=True).exclude(pk=int(request.POST['id'])).exists():
                        return JsonResponse({"result": "bad", "mensaje": "El identificador ya existe."})
                    if not filtro.en_uso():
                        filtro.descripcion = f.cleaned_data['descripcion']
                        filtro.tipo = f.cleaned_data['tipo']
                        filtro.script = f.cleaned_data['script']
                    filtro.identificador = f.cleaned_data['identificador']
                    filtro.save(request)
                    log(u'Editó campo para contrato dip: %s' % filtro, request, "edit")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'deletecampo':
            try:
                with transaction.atomic():
                    instancia = CampoContratoDip.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino Campo: %s' % instancia, request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'addplantilla':
            try:
                f = PlantillaContratoDipForm(request.POST, request.FILES)
                if f.is_valid():
                    filtro = PlantillaContratoDip(anio=f.cleaned_data['anio'],
                                                  descripcion=f.cleaned_data['descripcion'],
                                                  vigente=f.cleaned_data['vigente'],
                                                  perfil=f.cleaned_data['perfil'])
                    filtro.save(request)

                    if 'archivo' in request.FILES:
                        newfile = request.FILES['archivo']
                        extension = newfile._name.split('.')
                        tam = len(extension)
                        exte = extension[tam - 1]
                        if newfile.size > 100194304:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {"result": True, "mensaje": u"Error, el tamaño del archivo es mayor a 10 Mb."})
                        if not exte in ['doc', 'docx']:
                            transaction.set_rollback(True)
                            return JsonResponse({"result": True, "mensaje": u"Error, solo archivos .doc, docx"})
                        newfile._name = generar_nombre("plantillaepunemi", newfile._name)
                        filtro.archivo = newfile
                        filtro.save(request)

                    # INSERTA LOS CAMPOS SELECCIONADOS
                    for elemento in json.loads(request.POST['lista_items1']):
                        campoplantillaontratoDip = CampoPlantillaContratoDip(contrato=filtro,
                                                                             campos_id=int(elemento['id']))
                        campoplantillaontratoDip.save(request)

                    log(u'Adicionó plantilla de contrato dip: %s' % filtro, request, "add")
                    return JsonResponse({"result": "ok"}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'editplantilla':
            try:
                f = PlantillaContratoDipForm(request.POST)
                if f.is_valid():
                    datos = json.loads(request.POST['lista_items1'])
                    if len(datos) == 0:
                        return JsonResponse({"result": "bad", "mensaje": u"Seleccione al menos un campo del Contrato."})

                    filtro = PlantillaContratoDip.objects.get(pk=int(request.POST['id']))

                    # if filtro.en_uso():
                    #     return JsonResponse({"result": "bad", "mensaje": u"Seleccione al menos un campo del Contrato."})

                    filtro.anio = f.cleaned_data['anio']
                    filtro.descripcion = f.cleaned_data['descripcion']
                    filtro.vigente = f.cleaned_data['vigente']
                    filtro.perfil = f.cleaned_data['perfil']
                    if 'archivo' in request.FILES:
                        filtro.archivo.delete()
                        newfile = request.FILES['archivo']
                        extension = newfile._name.split('.')
                        tam = len(extension)
                        exte = extension[tam - 1]
                        if newfile.size > 100194304:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {"result": True, "mensaje": u"Error, el tamaño del archivo es mayor a 10 Mb."})
                        if not exte in ['doc', 'docx']:
                            transaction.set_rollback(True)
                            return JsonResponse({"result": True, "mensaje": u"Error, solo archivos .doc, docx"})
                        newfile._name = generar_nombre("plantillaepunemi", newfile._name)
                        filtro.archivo = newfile
                    filtro.save(request)
                    filtro.campoplantillacontratodip_set.all().delete()
                    # INSERTA LOS CAMPOS SELECCIONADOS
                    for elemento in datos:
                        campoplantillacontratoDip = CampoPlantillaContratoDip(contrato=filtro,
                                                                              campos_id=int(elemento['id']))
                        campoplantillacontratoDip.save(request)
                    log(u'Registro modificado contrato plantilla: %s' % filtro, request, "edit")
                    return JsonResponse({"result": "ok"})
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'gestionarcontrato':
            try:
                f = ContratoDipForm(request.POST)
                if f.is_valid():
                    contrato = ContratoDip.objects.get(pk=int(request.POST['id']))
                    contrato.plantilla = f.cleaned_data['plantilla']
                    contrato.fechainicio = f.cleaned_data['fechainicio']
                    contrato.fechafin = f.cleaned_data['fechafin']
                    contrato.codigocontrato = f.cleaned_data['codigocontrato']
                    contrato.rmu = f.cleaned_data['rmu']
                    contrato.iva = f.cleaned_data['ivaAplicado']
                    contrato.valoriva = f.cleaned_data['valorIva']
                    contrato.valortotal = f.cleaned_data['valorTotal']
                    contrato.certificacion = f.cleaned_data['certificacion']
                    contrato.estado = 2

                    contrato.save(request)
                    idContrato = ContratoDip.objects.all().last()
                    numeroDeCuota = f.cleaned_data['numeroCuota']
                    rmu = f.cleaned_data['rmu']
                    valorCuota = round((rmu / numeroDeCuota), 2)
                    sumatoriaCuota = 0
                    for i in range(numeroDeCuota):

                        if i + 1 == numeroDeCuota:
                            valorCuota = round((rmu - sumatoriaCuota), 2)

                        Cuota = ContratoDipMetodoPago(contratodip=idContrato,
                                                      numerocuota=i + 1,
                                                      valorcuota=valorCuota,
                                                      fecha_pago=fecha_pago(cuota=i + 1, contrato=contrato)
                                                      )
                        Cuota.save()
                        sumatoriaCuota += valorCuota
                    codigomaximo = 0
                    for campo in contrato.plantilla.campoplantillacontratodip_set.filter(status=True):
                        if campo.campos.id == 1:

                            contratodipdetalle = ContratoDipDetalle(contratodip=contrato,
                                                                    campo=campo,
                                                                    valor=codigomaximo)
                        else:
                            contratodipdetalle = ContratoDipDetalle(contratodip=contrato,
                                                                    campo=campo)
                        contratodipdetalle.save(request)

                    detalle = generar_detalle_certificacion(request, contrato)
                    if detalle:
                        log(u'Adicionó Detalle de Certificaion: %s' % detalle, request, "add")
                    log(u'Adicionó Contrato : %s' % contrato, request, "add")
                    return JsonResponse({"result": "ok", 'id': contrato.id})
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos. %s" % ex})

        elif action == 'selectrespon':
            try:
                if 'id' in request.POST:
                    res = Gestion.objects.get(status=True, id=int(request.POST['id']))
                    responsable = str(res.responsable if res.responsable else 'No existe')
                    responsablesub = str(res.responsablesubrogante if res.responsablesubrogante else 'No existe')
                    return JsonResponse({"result": "ok", 'responsable': responsable, 'responsablesub': responsablesub})
            except Exception as ex:
                return JsonResponse({"result": "bad", "mensaje": u"Error al consulatar los datos. %s" % (ex.__str__())})

        elif action == 'selectpartida':
            try:
                if 'id' in request.POST:
                    res = CertificacionPresupuestariaDip.objects.get(status=True, id=int(request.POST['id']))
                    codpartida = str(res.codigo if res else 'No existe')
                    return JsonResponse(
                        {"result": "ok", 'codigo': codpartida, 'fechapartida': res.fecha.strftime("%d/%m/%Y")})
            except Exception as ex:
                return JsonResponse({"result": "bad", "mensaje": u"Error al consulatar los datos."})

        elif action == 'validarcodigo':
            try:
                if 'cod' in request.POST:
                    if 'id' in request.POST:
                        res = ContratoDip.objects.values('id').filter(status=True, id=request.POST['id'],
                                                                      codigocontrato=request.POST['cod']).exists()
                        if res:
                            return JsonResponse({"result": "ok"})
                        return JsonResponse({"result": "bad", 'mensaje': 'El codigo esta en uso'})
                    else:
                        res = ContratoDip.objects.values('id').filter(status=True,
                                                                      codigocontrato=request.POST['cod']).exists()
                        if res:
                            return JsonResponse({"result": "bad"})
                        return JsonResponse({"result": "ok", 'mensaje': 'El codigo esta en uso'})
                if 'codpar' in request.POST:
                    if 'id' in request.POST:
                        res = ContratoDip.objects.values('id').filter(status=True, id=request.POST['id'],
                                                                      codigopartida=request.POST['codpar']).exists()
                        if res:
                            return JsonResponse({"result": "ok"})
                        return JsonResponse({"result": "bad", 'mensaje': 'El codigo esta en uso'})
                    else:
                        res = ContratoDip.objects.values('id').filter(status=True,
                                                                      codigopartida=request.POST['codpar']).exists()
                        if res:
                            return JsonResponse({"result": "ok"})
                        return JsonResponse({"result": "bad", 'mensaje': 'El codigo esta en uso'})
            except Exception as ex:
                return JsonResponse({"result": "bad", "mensaje": u"Error al consulatar los datos."})

        elif action == 'editcontrato':
            try:
                f = ContratoDipForm(request.POST)
                f.del_fields('carrera')
                if f.is_valid():
                    filtro = ContratoDip.objects.get(pk=int(request.POST['id']))
                    filtro.codigocontrato = f.cleaned_data['codigocontrato']
                    filtro.plantilla = f.cleaned_data['plantilla']
                    filtro.fechainicio = f.cleaned_data['fechainicio']
                    filtro.fechafin = f.cleaned_data['fechafin']
                    filtro.rmu = f.cleaned_data['rmu']
                    filtro.iva = f.cleaned_data['ivaAplicado']
                    filtro.valoriva = f.cleaned_data['valorIva']
                    filtro.valortotal = f.cleaned_data['valorTotal']
                    filtro.estado = 2
                    filtro.save(request)
                    historial = HistorialContratoDip(
                        contratodip=filtro,
                        observacion=f.cleaned_data['descripcion'],
                        estado=filtro.estado,
                        persona=persona
                    )
                    historial.save(request)
                    log(u'Gestionó Contrato: %s' % filtro, request, "add")
                    numeroDeCuota = f.cleaned_data['numeroCuota']
                    rmu = f.cleaned_data['rmu']
                    valorCuota = round((rmu / numeroDeCuota), 2)
                    sumatoriaCuota = 0
                    for i in range(numeroDeCuota):

                        if i + 1 == numeroDeCuota:
                            valorCuota = round((rmu - sumatoriaCuota), 2)

                        Cuota = ContratoDipMetodoPago(contratodip=filtro,
                                                      numerocuota=i + 1,
                                                      valorcuota=valorCuota,
                                                      fecha_pago=fecha_pago(cuota=i + 1, contrato=filtro)
                                                      )
                        Cuota.save()
                        sumatoriaCuota += valorCuota
                    codigomaximo = 0
                    for campo in filtro.plantilla.campoplantillacontratodip_set.filter(status=True):
                        if campo.campos.id == 1:

                            contratodipdetalle = ContratoDipDetalle(contratodip=filtro,
                                                                    campo=campo,
                                                                    valor=codigomaximo)
                        else:
                            contratodipdetalle = ContratoDipDetalle(contratodip=filtro,
                                                                    campo=campo)
                        contratodipdetalle.save(request)

                    detalle = generar_detalle_certificacion(request, filtro)
                    if detalle:
                        log(u'Adicionó Detalle de Certificaion: %s' % detalle, request, "add")
                    return JsonResponse({"result": "ok"})
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos. %s" % ex})

        elif action == 'editcontratopos':
            try:
                f = ContratoDipForm(request.POST, request.FILES)
                newfile = None
                if f.is_valid():
                    if 'archivo' in request.FILES:
                        newfile = request.FILES['archivo']
                        extension = newfile._name.split('.')
                        tam = len(extension)
                        exte = extension[tam - 1]
                        if newfile.size > 100194304:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {"result": 'bad', "mensaje": u"Error, el tamaño del archivo es mayor a 10 Mb."},
                                safe=False)
                    # if not exte in ['pdf']:
                    #     transaction.set_rollback(True)
                    #     return JsonResponse({"result": 'bad', "mensaje": u"Error, solo archivo .pdf"}, safe=False)
                    filtro = ContratoDip.objects.get(pk=int(request.POST['id']))
                    filtro.codigocontrato = f.cleaned_data['codigocontrato']
                    filtro.validadorgp_id = f.cleaned_data['validadorgp'] if f.cleaned_data['validadorgp'] else None
                    filtro.plantilla = f.cleaned_data['plantilla']
                    filtro.fechainicio = f.cleaned_data['fechainicio']
                    filtro.fechafin = f.cleaned_data['fechafin']
                    filtro.rmu = f.cleaned_data['rmu']
                    filtro.iva = f.cleaned_data['ivaAplicado']
                    filtro.valoriva = f.cleaned_data['valorIva']
                    filtro.valortotal = f.cleaned_data['valorTotal']
                    if filtro.estado != 5: filtro.estado = 2
                    filtro.gestion = f.cleaned_data['seccion']
                    filtro.certificacion = f.cleaned_data['certificacion']
                    filtro.cargo = f.cleaned_data['cargo']
                    filtro.tipogrupo = f.cleaned_data['tipogrupo']
                    filtro.tipopago = f.cleaned_data['tipopago']

                    lis_excluid = []
                    lis_excluid_area = []
                    for dep in f.cleaned_data['nombreareaprograma']:
                        if not ContratoAreaPrograma.objects.values('id').filter(status=True, contrato=filtro,
                                                                                gestion=dep).exists():
                            areapr = ContratoAreaPrograma(
                                contrato=filtro,
                                gestion=dep
                            )
                            areapr.save(request)
                            lis_excluid_area.append(areapr.id)
                        else:
                            areapr = ContratoAreaPrograma.objects.filter(status=True, contrato=filtro,
                                                                         gestion=dep).order_by('-id').first()
                            lis_excluid_area.append(areapr.id)
                    del_arpro = ContratoAreaPrograma.objects.filter(status=True, contrato=filtro).exclude(
                        id__in=lis_excluid_area)
                    for del_ar in del_arpro:
                        del_ar.status = False
                        del_ar.save(request)
                    for carrera in f.cleaned_data['carrera']:
                        if not ContratoCarrera.objects.values('id').filter(status=True, contrato=filtro,
                                                                           carrera=carrera).exists():
                            cc = ContratoCarrera(contrato=filtro, carrera=carrera)
                            cc.save(request)
                            his_ = HistorialContratoDipCarreras(contratocarrera=cc, carrera=carrera)
                            his_.save(request)
                            lis_excluid.append(cc.id)
                        else:
                            cc = ContratoCarrera.objects.filter(status=True, contrato=filtro, carrera=carrera).order_by(
                                '-id').first()
                            lis_excluid.append(cc.id)
                    con_carrera = ContratoCarrera.objects.filter(status=True, contrato=filtro).exclude(
                        id__in=lis_excluid)
                    for reg in con_carrera:
                        reg.status = False
                        reg.save(request)

                    if newfile: filtro.archivo = newfile
                    # directory_principal = os.path.join(SITE_STORAGE, 'media', 'contratosepunemi','contrato')
                    # try:
                    #     os.stat(directory_principal)
                    # except:
                    #     os.mkdir(directory_principal)
                    filtro.save(request)
                    historial = HistorialContratoDip(contratodip=filtro, observacion=f.cleaned_data['descripcion'],
                                                     estado=filtro.estado, persona=persona)
                    historial.save(request)
                    log(u'Gestionó Contrato: %s' % filtro, request, "edit")
                    codigomaximo = 0
                    for campo in filtro.plantilla.campoplantillacontratodip_set.filter(status=True):
                        if campo.campos.id == 1:

                            contratodipdetalle = ContratoDipDetalle(contratodip=filtro,
                                                                    campo=campo,
                                                                    valor=codigomaximo)
                        else:
                            contratodipdetalle = ContratoDipDetalle(contratodip=filtro,
                                                                    campo=campo)
                        contratodipdetalle.save(request)

                    # detalle = generar_detalle_certificacion(request, filtro)
                    # if detalle:
                    #     log(u'Adicionó Detalle de Certificaion: %s' % detalle, request, "add")
                    # else:
                    #     transaction.set_rollback(True)
                    #     mensaje = 'El valor excede a lo establecido en la certificaicon presupuestaria'
                    #     return  JsonResponse({'result':'bad','mensaje':mensaje})
                    return JsonResponse({"result": "ok"})
                else:
                    raise NameError(f'{[{k: v[0]} for k, v in f.errors.items()]}')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al guardar los datos. %s" % ex})

        elif action == 'deletecontrato':
            try:
                with transaction.atomic():
                    instancia = ContratoDip.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    certificacion = instancia.certificacion
                    certificacion.saldo += instancia.valortotal
                    certificacion.save(request)

                    detalleContrato = ContratoDipDetalle.objects.filter(contratodip=int(request.POST['id']))
                    if DetalleCertificacionPresupuestariaDip.objects.values('id').filter(status=True,
                                                                                         certificado=certificacion,
                                                                                         contratodip=instancia).exists():
                        detalle = DetalleCertificacionPresupuestariaDip.objects.filter(status=True,
                                                                                       certificado=certificacion,
                                                                                       contratodip=instancia).order_by(
                            '-id').first()
                        detalle.status = False
                        detalle.save(request)
                    for detalle in detalleContrato:
                        detalle.status = False
                        detalle.save(request)

                    log(u'Elimino Contrato: %s' % instancia, request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                transaction.set_rollback(True)
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'deleteplantilla':
            try:
                with transaction.atomic():
                    instancia = PlantillaContratoDip.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    camposPlantilla = CampoPlantillaContratoDip.objects.filter(contrato=int(request.POST['id']))

                    for campo in camposPlantilla:
                        campo.status = False
                        campo.save(request)

                    log(u'Elimino Plantilla: %s' % instancia, request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'detalle_plantilla':
            try:
                if 'id' in request.POST:
                    data['detallePlantilla'] = PlantillaContratoDip.objects.get(pk=request.POST['id'],
                                                                                status=True)
                    data['camposPlantilla'] = CampoPlantillaContratoDip.objects.filter(contrato=request.POST['id'],
                                                                                       status=True)
                    template = get_template("adm_contratodip/detalle_plantilla.html")
                    return JsonResponse({"result": 'ok', 'data': template.render(data)})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al obtener los datos."})

        elif action == 'vigente':
            try:
                contrato = PlantillaContratoDip.objects.get(pk=request.POST['id'])
                contrato.vigente = True if request.POST['val'] == 'y' else False
                contrato.save(request)
                log(u'Cambió vigencia de contrato dip: %s (%s)' % (contrato, contrato.vigente),
                    request, "edit")
                return JsonResponse({"result": "ok"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad"})

        elif action == 'editcontratodetalle':
            try:
                datos = request.POST['lista1']
                ids = request.POST['id']
                contratopersona = ContratoDip.objects.filter(pk=int(ids))[0]
                nombre_plantilla = contratopersona.plantilla.archivo.file.name
                nombre_contrato = str(contratopersona.id) + ".docx"
                try:
                    os.stat(os.path.join(SITE_ROOT, 'media', 'contratoepunemi', 'contrato'))
                except:
                    os.mkdir(os.path.join(SITE_ROOT, 'media', 'contratoepunemi', 'contrato'))
                direccion_contrato = os.path.join(SITE_ROOT, 'media', 'contratoepunemi', 'contrato')
                filename_contrato = os.path.join(direccion_contrato, nombre_contrato)
                # guarda la direccion
                cantidad_parrafo = 0
                if datos:
                    document = Document(nombre_plantilla)
                    for elementos in datos.split('######'):
                        elemento = elementos.split(';;;;')
                        if int(elemento[0]) == 1:
                            for elementos_aux in datos.split('######'):
                                elemento_aux = elementos_aux.split(';;;;')
                                if int(elemento_aux[0]) == 44:
                                    fecha_aux = elemento_aux[1]
                            anio = fecha_aux.split('-')[2]
                            mes = fecha_aux.split('-')[1]
                            valor1 = u"%s-%s-%s" % (anio, mes, contratopersona.persona.id)
                        else:
                            valor1 = elemento[1]

                        campoId = CampoPlantillaContratoDip.objects.get(campos=elemento[0],
                                                                        contrato=contratopersona.plantilla)

                        if ContratoDipDetalle.objects.filter(contratodip=contratopersona, campo=campoId).exists():
                            contratodipdetalle = \
                                ContratoDipDetalle.objects.filter(contratodip=contratopersona, campo=campoId)[0]
                            contratodipdetalle.valor = valor1
                            contratodipdetalle.save(request)
                        else:
                            contratodipdetalle = ContratoDipDetalle(contratodip=contratopersona, campo=campoId,
                                                                    valor=valor1)
                            contratodipdetalle.save(request)
                        parrafo = document.paragraphs
                        cantidad_parrafo = parrafo.__len__()
                        n = 0
                        campo = CampoContratoDip.objects.get(pk=int(elemento[0]))
                        if campo.script[:11] == 'JAVASCRIPT:':
                            campo1 = valor1
                            valor = ''
                            try:
                                valor = eval(campo.script[11:])
                            except:
                                pass
                            campo_buscar = '${CAMPO' + str(campo.identificador) + '}'

                            for n in range(cantidad_parrafo):
                                for run in parrafo[n].runs:
                                    if campo_buscar in run.text:
                                        run.text = run.text.replace(campo_buscar, str(valor))
                        else:
                            campo_buscar = '${CAMPO' + str(campo.identificador) + '}'
                            for n in range(cantidad_parrafo):
                                for run in parrafo[n].runs:
                                    if campo_buscar in run.text:
                                        run.text = run.text.replace(campo_buscar, str(valor1))
                    n = 0
                    persona_nombre = contratopersona.persona.nombre_titulo()
                    persona_cedula = contratopersona.persona.cedula
                    persona_ruc = u'%s001' % contratopersona.persona.cedula
                    cadena1 = '${EMPLEADO}'
                    cadena2 = '${CONTRATO}'
                    cadena3 = '${FECHAINICIO}'
                    cadena4 = '${CEDULA}'
                    cadena5 = '${RUC}'

                    for n in range(cantidad_parrafo):
                        for run in parrafo[n].runs:
                            run.text = run.text.replace(cadena1, persona_nombre)
                            run.text = run.text.replace(cadena2, contratopersona.codigocontrato)
                            run.text = run.text.replace(cadena3, str(contratopersona.fechainicio))
                            run.text = run.text.replace(cadena4, persona_cedula)
                            run.text = run.text.replace(cadena5, str(persona_ruc))
                    document.save(filename_contrato)
                contratopersona.archivo.name = "contratoepunemi/contrato/%s" % nombre_contrato
                contratopersona.save(request)
                log(u'Contrato Persona: %s' % contratopersona, request, "edit")
                return JsonResponse({"result": "ok"})

            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": "Error al procesar los datos."})

        elif action == 'addcontrato':
            try:
                newfile = None
                form = ContratoDipForm(request.POST, request.FILES)
                if form.is_valid():
                    tipogrupo = int(form.cleaned_data['tipogrupo'])
                    tipopago = int(form.cleaned_data['tipopago'])
                    if tipogrupo == 0:
                        return JsonResponse({"result": 'bad', "mensaje": f"Complete el campo tipo grupo"},safe=False)

                    if tipopago == 0:
                        return JsonResponse({"result": 'bad', "mensaje": f"Complete el campo tipo pago"},safe=False)

                    # if not 'archivo' in request.FILES:
                    #     return JsonResponse({"result": 'bad', "mensaje": f"Suba el contrato legalizado"},safe=False)

                    if 'archivo' in request.FILES:
                        newfile = request.FILES['archivo']
                        extension = newfile._name.split('.')
                        tam = len(extension)
                        exte = extension[tam - 1]
                        if newfile.size > 100194304:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {"result": 'bad', "mensaje": u"Error, el tamaño del archivo es mayor a 10 Mb."},
                                safe=False)
                        # if not exte in ['pdf']:
                        #     transaction.set_rollback(True)
                        #     return JsonResponse({"result": 'bad', "mensaje": u"Error, solo archivo .pdf"}, safe=False)

                        _name = remover_caracteres_tildes_unicode(remover_caracteres_especiales_unicode(extension[0])).lower().replace(' ', '_').replace('-','_')
                        newfile._name = generar_nombre(u"%s_" % _name, f"{_name}.pdf")
                    contrato = ContratoDip(codigocontrato=form.cleaned_data['codigocontrato'],
                                           gestion=form.cleaned_data['seccion'],
                                           persona_id=form.cleaned_data['persona'],
                                           validadorgp_id=form.cleaned_data['validadorgp'] if form.cleaned_data[
                                               'validadorgp'] else None,
                                           plantilla=form.cleaned_data['plantilla'],
                                           fechainicio=form.cleaned_data['fechainicio'],
                                           fechafin=form.cleaned_data['fechafin'],
                                           valortotal=form.cleaned_data['valorTotal'],
                                           valoriva=form.cleaned_data['valorIva'],
                                           descripcion=form.cleaned_data['descripcion'],
                                           rmu=form.cleaned_data['rmu'],
                                           iva=form.cleaned_data['ivaAplicado'],
                                           cargo=form.cleaned_data['cargo'],
                                           tipogrupo=form.cleaned_data['tipogrupo'],
                                           tipopago=form.cleaned_data['tipopago'],
                                           manual=True,
                                           certificacion=form.cleaned_data['certificacion'],
                                           estado=2,
                                           archivo=newfile)

                    contrato.save(request)
                    contrato.notificar_contrato_subido_para_registrar_analista_validador(request)
                    for dep in form.cleaned_data['nombreareaprograma']:
                        if not ContratoAreaPrograma.objects.values('id').filter(status=True, contrato=contrato,
                                                                                gestion=dep).exists():
                            areapr = ContratoAreaPrograma(
                                contrato=contrato,
                                gestion=dep
                            )
                            areapr.save(request)

                    for carrera in form.cleaned_data['carrera']:
                        if not ContratoCarrera.objects.values('id').filter(contrato=contrato, carrera=carrera,
                                                                           status=True).exists():
                            cc = ContratoCarrera(contrato=contrato, carrera=carrera)
                            cc.save()
                            his_ = HistorialContratoDipCarreras(contratocarrera=cc, carrera=carrera)
                            his_.save(request)
                    log(u'Agregar contrato dip: %s (%s)' % (contrato, contrato.codigocontrato), request, 'add')
                else:
                    raise NameError(f'{[{k: v[0]} for k, v in form.errors.items()]}')
                return JsonResponse({"result": "ok"},
                                    safe=False)  # return JsonResponse({"result": "ok","id":contrato.id}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": f"Error al procesar los datos. {ex}"})

        elif action == 'addarchivoevidencias':
            try:
                form = ArchivoInformesForm(request.POST, request.FILES)
                if form.is_valid():
                    newfile = request.FILES['archivo']
                    extension = newfile._name.split('.')
                    tam = len(extension)
                    exte = extension[tam - 1]
                    if newfile.size > 100194304:
                        transaction.set_rollback(True)
                        return JsonResponse(
                            {"result": 'bad', "mensaje": u"Error, el tamaño del archivo es mayor a 10 Mb."}, safe=False)
                    if not exte in ['pdf']:
                        transaction.set_rollback(True)
                        return JsonResponse({"result": 'bad', "mensaje": u"Error, solo archivo .pdf"}, safe=False)
                    if int(request.POST['tipo']) == 1:
                        arch = MemoActividadPosgrado.objects.get(pk=request.POST['id'])
                    elif int(request.POST['tipo']) == 2:
                        arch = InformeActividadJornada.objects.get(pk=request.POST['id'])
                    elif int(request.POST['tipo']) == 3:
                        arch = InformeTecnico.objects.get(pk=request.POST['id'])
                    elif int(request.POST['tipo']) == 4:
                        arch = ActaPago.objects.get(pk=request.POST['id'])
                    arch.archivofirmado = newfile
                    arch.save(request)
                    return JsonResponse({'result': 'ok', 'mensaje': 'Archivo guardado con éxito'}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": 'bad', "mensaje": u"Error al procesar los datos."}, safe=False)

        elif action == 'finalizarcontrato':
            try:
                idcon = int(encrypt(request.POST['idcon']))
                if ContratoDip.objects.values('id').filter(status=True, pk=idcon).exists():
                    contrato = ContratoDip.objects.get(status=True, pk=idcon)
                    contrato.estado = 5
                    contrato.fechafinalizacion = request.POST['fecha']
                    contrato.save(request)
                    return JsonResponse({'result': 'ok'}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": 'bad', "mensaje": u"Error al procesar los datos."}, safe=False)

        elif action == 'addcertificacion':
            try:
                f = CertificacionPresupuestariaDipForm(request.POST)
                if variable_valor('VALIDAR_CERTIFICACION_PRESUPUESTARIA_POS'):
                    if 'fecha' in request.POST:
                        if convertir_fecha(request.POST['fecha']) < datetime.now().date():
                            return JsonResponse({"result": True,
                                                 "mensaje": "*La fecha ingresada debe ser mayor o igual a la fecha actual. %s" % datetime.now().date()},
                                                safe=False)
                if f.is_valid():
                    filtro = CertificacionPresupuestariaDip(partida=f.cleaned_data['partida'],
                                                            descripcion=f.cleaned_data['descripcion'],
                                                            codigo=f.cleaned_data['codigo'],
                                                            valor=f.cleaned_data['valor'],
                                                            fecha=f.cleaned_data['fecha'],
                                                            saldo=f.cleaned_data['valor'])
                    filtro.save(request)

                    if 'archivo' in request.FILES:
                        newfile = request.FILES['archivo']
                        extension = newfile._name.split('.')
                        tam = len(extension)
                        exte = extension[tam - 1]
                        if newfile.size > 100194304:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {"result": True, "mensaje": u"Error, el tamaño del archivo es mayor a 10 Mb."})
                        if not exte.lower() in ['png', 'pdf', 'jpg', 'jpeg']:
                            transaction.set_rollback(True)
                            return JsonResponse({"result": True, "mensaje": u"Error, solo archivos .png,pdf,jpg,jpeg"})
                        newfile._name = generar_nombre("certificacionepunemi", newfile._name)
                        filtro.archivo = newfile
                        filtro.save(request)
                    log(u'Adicionó certificacion dip: %s' % filtro, request, "add")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError('Error')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'addrecursopresupuesatario':
            try:
                f = RecursoPresupuestarioForm(request.POST)
                if f.is_valid():
                    eRecursoPresupuestarioPosgrado = RecursoPresupuestarioPosgrado(
                        descripcion=f.cleaned_data['descripcion'])
                    eRecursoPresupuestarioPosgrado.save(request)
                    log(u"Agregó recurso", request, 'add')
                    return JsonResponse({"result": True})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'editrecursopresupuesatario':
            try:
                f = RecursoPresupuestarioForm(request.POST)
                pk = request.POST.get('id', '0')
                eRecursoPresupuestarioPosgrado = RecursoPresupuestarioPosgrado.objects.get(pk=pk)
                if f.is_valid():
                    eRecursoPresupuestarioPosgrado.descripcion = f.cleaned_data['descripcion']
                    eRecursoPresupuestarioPosgrado.save(request)
                    log(u"Edito recurso", request, 'add')
                    return JsonResponse({"result": True})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'add_programa_maestria':
            try:
                f = CabeceraRecursoPresupuestarioPosgradoForm(request.POST)
                if f.is_valid():

                    eCabeceraRecursoPresupuestarioPosgrado = CabeceraRecursoPresupuestarioPosgrado(
                        recursopresupuestarioPosgrado_id=int(request.POST.get('id_secundario', '0')),
                        malla=f.cleaned_data['malla'],
                        periodo=f.cleaned_data['periodo']
                    )
                    eCabeceraRecursoPresupuestarioPosgrado.save(request)
                    log(u"Agregó recurso presupuestario programa", request, 'add')
                    return JsonResponse({"result": True})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'edit_programa_maestria':
            try:
                f = CabeceraRecursoPresupuestarioPosgradoForm(request.POST)
                pk = request.POST.get('id', '0')
                eCabeceraRecursoPresupuestarioPosgrado = CabeceraRecursoPresupuestarioPosgrado.objects.get(pk=pk)
                if f.is_valid():
                    eCabeceraRecursoPresupuestarioPosgrado.malla = f.cleaned_data['malla']
                    eCabeceraRecursoPresupuestarioPosgrado.periodo = f.cleaned_data['periodo']
                    eCabeceraRecursoPresupuestarioPosgrado.save(request)
                    log(u"Editar recurso presupuestario programa", request, 'add')
                    return JsonResponse({"result": True})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'add_subitems':
            try:
                f = DetalleRecursoPresupuestarioPosgradoForm(request.POST)
                if f.is_valid():

                    eDetalleRecursoPresupuestarioPosgrado = DetalleRecursoPresupuestarioPosgrado(
                        itemrecursopresupuestarioposgrado_id=int(request.POST.get('id_secundario', '0')),
                        desglosemoduloadictar=f.cleaned_data['desglosemoduloadictar'],
                        horaspormodulo=f.cleaned_data['horaspormodulo'],
                        valor_x_hora=f.cleaned_data['valor_x_hora'],
                        categoriadocente=f.cleaned_data['categoriadocente']
                    )
                    eDetalleRecursoPresupuestarioPosgrado.save(request)
                    log(u"Agregó sub item recurso presupuestario programa", request, 'add')
                    return JsonResponse({"result": True})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'add_items':
            try:
                f = ItemRecursoPresupuestarioPosgradoForm(request.POST)
                if f.is_valid():

                    eItemRecursoPresupuestarioPosgrado = ItemRecursoPresupuestarioPosgrado(
                        cabecerarecursopresupuestarioposgrado_id=int(request.POST.get('id_secundario', '0')),
                        total_paralelos=f.cleaned_data['total_paralelos'],
                        modulos_a_dictar=f.cleaned_data['modulos_a_dictar'],
                    )
                    eItemRecursoPresupuestarioPosgrado.save(request)
                    log(u"Agregó item recurso presupuestario programa", request, 'add')
                    return JsonResponse({"result": True})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'edit_items':
            try:
                f = ItemRecursoPresupuestarioPosgradoForm(request.POST)
                pk = request.POST.get('id', '0')
                if pk == 0:
                    raise NameError("Parametro no encontrado")

                eItemRecursoPresupuestarioPosgrado = ItemRecursoPresupuestarioPosgrado.objects.get(pk=pk)
                if f.is_valid():
                    eItemRecursoPresupuestarioPosgrado.total_paralelos = f.cleaned_data['total_paralelos']
                    eItemRecursoPresupuestarioPosgrado.modulos_a_dictar = f.cleaned_data['modulos_a_dictar']
                    eItemRecursoPresupuestarioPosgrado.save(request)
                    log(u"editar item recurso presupuestario programa", request, 'add')
                    return JsonResponse({"result": True})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'edit_subitems':
            try:
                f = DetalleRecursoPresupuestarioPosgradoForm(request.POST)
                pk = request.POST.get('id', '0')
                if pk == 0:
                    raise NameError("Parametro no encontrado")

                eDetalleRecursoPresupuestarioPosgrado = DetalleRecursoPresupuestarioPosgrado.objects.get(pk=pk)
                if f.is_valid():
                    eDetalleRecursoPresupuestarioPosgrado.desglosemoduloadictar = f.cleaned_data[
                        'desglosemoduloadictar']
                    eDetalleRecursoPresupuestarioPosgrado.horaspormodulo = f.cleaned_data['horaspormodulo']
                    eDetalleRecursoPresupuestarioPosgrado.valor_x_hora = f.cleaned_data['valor_x_hora']
                    eDetalleRecursoPresupuestarioPosgrado.categoriadocente = f.cleaned_data['categoriadocente']
                    eDetalleRecursoPresupuestarioPosgrado.save(request)
                    log(u"editar item recurso presupuestario programa", request, 'add')
                    return JsonResponse({"result": True})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'delete_subitems':
            try:
                eDetalleRecursoPresupuestarioPosgrado = DetalleRecursoPresupuestarioPosgrado.objects.get(
                    pk=int(request.POST['id']))
                log(u'Elimino  eDetalleRecursoPresupuestarioPosgrado: %s' % eDetalleRecursoPresupuestarioPosgrado,
                    request, "del")
                eDetalleRecursoPresupuestarioPosgrado.status = False
                eDetalleRecursoPresupuestarioPosgrado.save(request)
                return JsonResponse({"result": True, "error": False})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})

        elif action == 'delete_items':
            try:
                eItemRecursoPresupuestarioPosgrado = ItemRecursoPresupuestarioPosgrado.objects.get(
                    pk=int(request.POST['id']))
                log(u'Elimino  eItemRecursoPresupuestarioPosgrado: %s' % eItemRecursoPresupuestarioPosgrado,
                    request, "del")
                eItemRecursoPresupuestarioPosgrado.status = False
                eItemRecursoPresupuestarioPosgrado.save(request)
                return JsonResponse({"result": True, "error": False})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})

        elif action == 'delete_programa_maestria':
            try:
                eCabeceraRecursoPresupuestarioPosgrado = CabeceraRecursoPresupuestarioPosgrado.objects.get(
                    pk=int(request.POST['id']))
                eCabeceraRecursoPresupuestarioPosgrado.status = False
                eCabeceraRecursoPresupuestarioPosgrado.save(request)
                log(u'Elimino  programa maestria recurso presupuestario: %s' % eCabeceraRecursoPresupuestarioPosgrado,
                    request, "del")

                return JsonResponse({"result": True, "error": False})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})

        elif action == 'editcertificacion':
            try:
                f = CertificacionPresupuestariaDipForm(request.POST)
                if f.is_valid():
                    filtro = CertificacionPresupuestariaDip.objects.get(pk=int(request.POST['id']))
                    filtro.valor = f.cleaned_data['valor']
                    filtro.descripcion = f.cleaned_data['descripcion']
                    filtro.fecha = f.cleaned_data['fecha']
                    filtro.partida = f.cleaned_data['partida']
                    filtro.codigo = f.cleaned_data['codigo']
                    if 'archivo' in request.FILES:
                        filtro.archivo.delete()
                        newfile = request.FILES['archivo']
                        extension = newfile._name.split('.')
                        tam = len(extension)
                        exte = extension[tam - 1]
                        if newfile.size > 100194304:
                            transaction.set_rollback(True)
                            return JsonResponse(
                                {"result": True, "mensaje": u"Error, el tamaño del archivo es mayor a 10 Mb."})
                        if not exte.lower() in ['png', 'pdf', 'jpg', 'jpeg']:
                            raise NameError("Error, solo archivos .png,pdf,jpg,jpeg")
                        newfile._name = generar_nombre("certificacionepunemi", newfile._name)
                        filtro.archivo = newfile
                    filtro.save(request)
                    log(u'Editó certificacion de contrato dip: %s' % filtro, request, "edit")
                    return JsonResponse({"result": False}, safe=False)
                else:
                    raise NameError(f'{[{k: v[0]} for k, v in f.errors.items()]}')
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": f"{ex.__str__()}({sys.exc_info()[-1].tb_lineno})"},
                                    safe=False)

        elif action == 'deletecertificacion':
            try:
                id = request.POST['id']
                filtro = CertificacionPresupuestariaDip.objects.get(pk=int(id))
                if not filtro.esta_en_contrato():
                    filtro.status = False
                    filtro.save(request)
                    log(u'Eliminó certificacion de contrato dip: %s' % filtro, request, "del")
                    return JsonResponse(
                        {"result": False, "mensaje": u"Certificacion eliminada correctamnete!", 'error': False})
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": u"Error, esta certificaion esta en un contrato"})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'adddocument':
            try:
                f = DocumentoContratosForm(request.POST)
                if not f.is_valid():
                    raise NameError(str(f.errors))
                documentoconrato = DocumentoContrato(
                    nombre=f.cleaned_data['nombre'],
                    codigo=f.cleaned_data['codigo'],
                    campo=f.cleaned_data['campo']
                )
                documentoconrato.save(request)
                log(f"Agrego documento contrato: {documentoconrato}", request, 'add')
                return JsonResponse({"result": False}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'editdocument':
            try:
                f = DocumentoContratosForm(request.POST)
                filtro = DocumentoContrato.objects.get(id=request.POST['id'], status=True)
                if not f.is_valid():
                    raise NameError(str(f.errors))
                filtro.nombre = f.cleaned_data['nombre']
                filtro.codigo = f.cleaned_data['codigo']
                filtro.campo = f.cleaned_data['campo']
                filtro.save(request)
                log(f"Edito documento contrato: {filtro}", request, 'add')
                return JsonResponse({"result": False}, safe=False)
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

        elif action == 'deletedocument':
            try:
                with transaction.atomic():
                    instancia = DocumentoContrato.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino Campo: %s' % instancia, request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)

        elif action == 'addhorariocontrato':
            try:
                f = HorarioClasesContratoForm(request.POST)
                eConvocatoria = ContratoDip.objects.get(pk=request.POST['id'])
                if not f.is_valid():
                    raise NameError(f'{[{k: v[0]} for k, v in f.errors.items()]}')
                dia, inicio, fin = f.cleaned_data['dia'], f.cleaned_data['inicio'], f.cleaned_data['fin']

                if inicio < eConvocatoria.fechainicio or fin < eConvocatoria.fechainicio or fin > eConvocatoria.fechafin or inicio > eConvocatoria.fechafin:
                    raise NameError(f"La fecha ingresada no se encuenta en el rango de fechas inicio/fin del módulo.")
                if HorarioPlanificacionContrato.objects.filter(dia=dia, contrato=eConvocatoria, inicio=inicio, fin=fin,
                                                               status=True).exists():
                    raise NameError(f"Ya existe un horario registrado en este dia y fecha.")
                horario = HorarioPlanificacionContrato(dia=dia, contrato=eConvocatoria, inicio=inicio, fin=fin)
                horario.save(request)
                horario.turno.clear()
                horario.turno.set(f.cleaned_data.get('turno'))
                log(f"Agregó horario a contrato {eConvocatoria}", request, 'add')
                return JsonResponse({"result": True})
            except Exception as ex:
                err_ = f'{ex}({sys.exc_info()[-1].tb_lineno})'
                return JsonResponse({"result": False, 'mensaje': err_})

        elif action == 'edithorariocontrato':
            try:
                f = HorarioClasesContratoForm(request.POST)
                horario = HorarioPlanificacionContrato.objects.get(status=True, id=request.POST['id'])
                eConvocatoria = horario.contrato
                if not f.is_valid():
                    raise NameError(f'{[{k: v[0]} for k, v in f.errors.items()]}')
                dia, inicio, fin = f.cleaned_data['dia'], f.cleaned_data['inicio'], f.cleaned_data['fin']

                if inicio < eConvocatoria.fechainicio or fin < eConvocatoria.fechainicio or fin > eConvocatoria.fechafin or inicio > eConvocatoria.fechafin:
                    raise NameError(f"La fecha ingresada no se encuenta en el rango de fechas inicio/fin del contrato.")
                if HorarioPlanificacionContrato.objects.filter(dia=dia, contrato=eConvocatoria, inicio=inicio, fin=fin,
                                                               status=True).exclude(id=horario.id).exists():
                    raise NameError(f"Ya existe un horario registrado en este dia y fecha.")
                horario.dia = dia
                horario.inicio = inicio
                horario.fin = fin
                horario.save(request)
                horario.turno.clear()
                horario.turno.set(f.cleaned_data.get('turno'))
                log(f"Agregó horario a contrato {eConvocatoria}", request, 'change')
                return JsonResponse({"result": True})
            except Exception as ex:
                err_ = f'{ex}({sys.exc_info()[-1].tb_lineno})'
                return JsonResponse({"result": False, 'mensaje': err_})

        elif action == 'updatebitacora':
            try:
                form = FechaMaximoBitacoraForm(request.POST)
                if not form.is_valid():
                    raise NameError(f"{[{k: v[0]} for k, v in form.errors.items()]}")
                id = encrypt(request.POST['id'])
                registro = ContratoDip.objects.get(status=True, id=int(id))
                # registro.bitacora = not registro.bitacora
                registro.fechaaplazo = form.cleaned_data['fechaaplazo']
                registro.save(request)
                if not persona.usuario.is_superuser: log(
                    "Actualizo permiso para registrar actividades en bitácora: %s" % (registro.__str__()), request,
                    'change')
                res_json = {"result": False}
            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": True, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'addgrouprevision':
            try:
                form = GrupoRevisionPagoContratoForm(request.POST)
                form.edit(request.POST['persona'])
                ids_items = request.POST.get('lista_items1', None)
                if ids_items: form.edit_personacontrato(json.loads(ids_items))
                if not form.is_valid():
                    raise NameError(f'{[{k: v[0]} for k, v in form.errors.items()]}')
                if GrupoRevisionPago.objects.filter(status=True, persona=form.cleaned_data['persona']):
                    raise NameError("Ya existe un grupo con esa persona")
                grupo = GrupoRevisionPago(
                    nombre=form.cleaned_data['nombre'],
                    persona=form.cleaned_data['persona']
                )
                grupo.save(request)
                for cont in form.cleaned_data['personacontrato']:
                    if GrupoRevisionPagoContrato.objects.filter(status=True, personacontrato=cont).exclude(
                            gruporevision=grupo).exists():
                        raise NameError(f"La persona {cont}, se encuentra ya en grupo de revisión")
                    rev = GrupoRevisionPagoContrato(
                        gruporevision=grupo,
                        personacontrato=cont
                    )
                    rev.save(request)
                log(f"Agrego grupo de revisión: {grupo}", request, 'add')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'addconfiguracionrequisitocontratacion':
            try:
                form = ContratacionConfiguracionRequisitoForm(request.POST)
                if form.is_valid():
                    eContratacionConfiguracionRequisito = ContratacionConfiguracionRequisito(
                        nombre=form.cleaned_data['nombre'],
                        activo=form.cleaned_data['activo'],
                    )
                    eContratacionConfiguracionRequisito.save(request)
                    log(f"Agrego configuracion requisito contratacion: {eContratacionConfiguracionRequisito}", request,
                        'add')
                res_json = {"result": True}

            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'add_orden_firma_acta_pago':
            try:
                f = OrdenFirmaInformeActaPagoForm(request.POST)
                if f.is_valid():
                    eOrdenFirmaActaPago = OrdenFirmaActaPago(
                        responsabilidadfirma=f.cleaned_data['responsabilidadfirma'], orden=f.cleaned_data['orden'])
                    eOrdenFirmaActaPago.save(request)
                    log(u"Add orden firma acta pago", request, 'add')
                    return JsonResponse({"result": "ok"})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'add_formato_acta_pago':
            try:
                f = ConfiguracionActaPagoForm(request.POST)
                if f.is_valid():
                    eConfiguracionActaPago = ConfiguracionActaPago(titulo=f.cleaned_data['titulo'])
                    eConfiguracionActaPago.save(request)
                    log(u"Add configuracion firma acta pago", request, 'add')
                    return JsonResponse({"result": "ok"})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'addconfiguracionrequisitopago':
            try:
                form = ConfiguracionGrupoPagoForm(request.POST)
                if form.is_valid():
                    eGrupoRequisitoPago = GrupoRequisitoPago(
                        descripcion=form.cleaned_data['descripcion'],
                        tipogrupo=form.cleaned_data['tipogrupo'],
                        activo=form.cleaned_data['activo']
                    )
                    eGrupoRequisitoPago.save(request)
                    log(f"Agrego configuracion requisito pago: {eGrupoRequisitoPago}", request, 'add')
                res_json = {"result": True}

            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'addgestorrequisitocontratacion':
            try:
                form = ContratacionGestionRequisitoForm(request.POST)
                if form.is_valid():
                    eRequisitoContratacionConfiguracionRequisito = RequisitoContratacionConfiguracionRequisito(
                        contratacionconfiguracionrequisito_id=int(request.POST['id']),
                        requisito=form.cleaned_data['requisito'],
                        tipo=form.cleaned_data['tipo'],
                        opcional=form.cleaned_data['opcional'],
                    )
                    eRequisitoContratacionConfiguracionRequisito.save(request)
                    log(f"Agrego  requisito contratacion: {eRequisitoContratacionConfiguracionRequisito}", request,
                        'add')
                res_json = {"result": True}

            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'addgestorrequisitopago':
            try:
                form = RequisitoPagoDipForm(request.POST)
                if form.is_valid():
                    eRequisitoPagoGrupoRequisito = RequisitoPagoGrupoRequisito(
                        gruporequisitopago_id=int(request.POST['id']),
                        requisitopagodip=form.cleaned_data['requisito'],
                        orden=form.cleaned_data['orden'],
                        opcional=form.cleaned_data['opcional'],
                    )
                    eRequisitoPagoGrupoRequisito.save(request)
                    log(f"Agrego configuracion requisito pago: {eRequisitoPagoGrupoRequisito}", request, 'add')
                res_json = {"result": True}

            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'deletegrouprevision':
            try:
                id = int(encrypt(request.POST['id']))
                reg = GrupoRevisionPago.objects.get(status=True, id=id)
                reg.status = False
                reg.save(request)
                for grcon in GrupoRevisionPagoContrato.objects.filter(status=True, gruporevision=reg):
                    grcon.status = False
                    grcon.save(request)
                log(f"Eliminó al departamento una gestión de posgrado: {reg}", request, 'del')
                res_json = {"error": False}
            except Exception as ex:
                err_ = "{}({})".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {'error': True, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'deleteconfiguracionrequisitocontratacion':
            try:
                id = int(encrypt(request.POST['id']))
                eContratacionConfiguracionRequisito = ContratacionConfiguracionRequisito.objects.get(status=True, id=id)
                eContratacionConfiguracionRequisito.status = False
                eContratacionConfiguracionRequisito.save(request)

                log(f"Eliminó configuracion de requisitos contratacion: {eContratacionConfiguracionRequisito}", request,
                    'del')
                res_json = {"error": False}
            except Exception as ex:
                err_ = "{}({})".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {'error': True, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'deleteconfiguracionrequisitopago':
            try:
                id = int(encrypt(request.POST['id']))
                eGrupoRequisitoPago = GrupoRequisitoPago.objects.get(status=True, id=id)
                eGrupoRequisitoPago.status = False
                eGrupoRequisitoPago.save(request)

                log(f"Eliminó configuracion de requisitos pago: {eGrupoRequisitoPago}", request, 'del')
                res_json = {"error": False}
            except Exception as ex:
                err_ = "{}({})".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {'error': True, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'deletegestorrequisitocontratacion':
            try:
                id = int(encrypt(request.POST['id']))
                eRequisitoContratacionConfiguracionRequisito = RequisitoContratacionConfiguracionRequisito.objects.get(
                    status=True, id=id)
                eRequisitoContratacionConfiguracionRequisito.status = False
                eRequisitoContratacionConfiguracionRequisito.save(request)

                log(f"Eliminó requisitos contratacion: {eRequisitoContratacionConfiguracionRequisito}", request, 'del')
                res_json = {"error": False}
            except Exception as ex:
                err_ = "{}({})".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {'error': True, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'deletegestorrequisitopago':
            try:
                id = int(encrypt(request.POST['id']))
                eRequisitoPagoGrupoRequisito = RequisitoPagoGrupoRequisito.objects.get(status=True, id=id)
                eRequisitoPagoGrupoRequisito.status = False
                eRequisitoPagoGrupoRequisito.save(request)

                log(f"Eliminó requisitos contratacion: {eRequisitoPagoGrupoRequisito}", request, 'del')
                res_json = {"error": False}
            except Exception as ex:
                err_ = "{}({})".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {'error': True, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'editgrouprevision':
            try:
                id = request.POST.get('id', None)
                grupo = GrupoRevisionPago.objects.get(status=True, id=int(encrypt(id)))
                form = GrupoRevisionPagoContratoForm(request.POST)
                form.edit(request.POST['persona'])
                ids_items = request.POST.get('lista_items1', None)
                if ids_items: form.edit_personacontrato(json.loads(ids_items))
                if not form.is_valid():
                    raise NameError(f'{[{k: v[0]} for k, v in form.errors.items()]}')
                if GrupoRevisionPago.objects.filter(status=True, persona=form.cleaned_data['persona']).exclude(
                        id=grupo.id):
                    raise NameError("Ya existe un grupo con esa persona")
                grupo.nombre = form.cleaned_data['nombre']
                grupo.persona = form.cleaned_data['persona']
                grupo.save(request)
                lis_excl = []
                for cont in form.cleaned_data['personacontrato']:
                    if GrupoRevisionPagoContrato.objects.filter(status=True, personacontrato=cont).exclude(
                            gruporevision=grupo).exists():
                        raise NameError(f"La persona {cont}, se encuentra ya en grupo de revisión")
                    if GrupoRevisionPagoContrato.objects.filter(status=True, personacontrato=cont,
                                                                gruporevision=grupo).exists():
                        rev = GrupoRevisionPagoContrato.objects.filter(status=True, personacontrato=cont,
                                                                       gruporevision=grupo).order_by('-id').first()
                    else:
                        rev = GrupoRevisionPagoContrato(
                            gruporevision=grupo,
                            personacontrato=cont
                        )
                        rev.save(request)
                    lis_excl.append(rev.id)
                grupocontr_borrar = GrupoRevisionPagoContrato.objects.filter(status=True, gruporevision=grupo).exclude(
                    id__in=lis_excl)
                for gbr in grupocontr_borrar:
                    gbr.status = False
                    gbr.save(request)
                log(f"Editó grupo de revisión: {grupo}", request, 'change')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'editconfiguracionrequisitocontratacion':
            try:
                id = request.POST.get('id', None)
                eContratacionConfiguracionRequisito = ContratacionConfiguracionRequisito.objects.get(status=True,
                                                                                                     id=int(
                                                                                                         encrypt(id)))
                form = ContratacionConfiguracionRequisitoForm(request.POST)
                if form.is_valid():
                    eContratacionConfiguracionRequisito.nombre = form.cleaned_data['nombre']
                    eContratacionConfiguracionRequisito.activo = form.cleaned_data['activo']
                    eContratacionConfiguracionRequisito.save(request)
                    log(f"Editó eContratacionConfiguracionRequisito de revisión: {eContratacionConfiguracionRequisito}",
                        request, 'change')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'editconfiguracionrequisitopago':
            try:
                id = request.POST.get('id', None)
                eGrupoRequisitoPago = GrupoRequisitoPago.objects.get(status=True, id=int(encrypt(id)))
                form = ConfiguracionGrupoPagoForm(request.POST)
                if form.is_valid():
                    eGrupoRequisitoPago.descripcion = form.cleaned_data['descripcion']
                    eGrupoRequisitoPago.tipogrupo = form.cleaned_data['tipogrupo']
                    eGrupoRequisitoPago.activo = form.cleaned_data['activo']
                    eGrupoRequisitoPago.save(request)
                    log(f"Editó eGrupoRequisitoPago de revisión: {eGrupoRequisitoPago}", request, 'change')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'editgestorrequisitocontratacion':
            try:
                id = request.POST.get('id', None)
                eRequisitoContratacionConfiguracionRequisito = RequisitoContratacionConfiguracionRequisito.objects.get(
                    status=True, id=int(encrypt(id)))
                form = ContratacionGestionRequisitoForm(request.POST)
                if form.is_valid():
                    eRequisitoContratacionConfiguracionRequisito.requisito = form.cleaned_data['requisito']
                    eRequisitoContratacionConfiguracionRequisito.tipo = form.cleaned_data['tipo']
                    eRequisitoContratacionConfiguracionRequisito.opcional = form.cleaned_data['opcional']
                    eRequisitoContratacionConfiguracionRequisito.save(request)
                    log(f"Editó eRequisitoContratacionConfiguracionRequisito de revisión: {eRequisitoContratacionConfiguracionRequisito}",
                        request, 'change')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'editgestorrequisitopago':
            try:
                id = request.POST.get('id', None)
                eRequisitoPagoGrupoRequisito = RequisitoPagoGrupoRequisito.objects.get(status=True, id=int(encrypt(id)))
                form = RequisitoPagoDipForm(request.POST)
                if form.is_valid():
                    eRequisitoPagoGrupoRequisito.requisitopagodip = form.cleaned_data['requisito']
                    eRequisitoPagoGrupoRequisito.orden = form.cleaned_data['orden']
                    eRequisitoPagoGrupoRequisito.opcional = form.cleaned_data['opcional']
                    eRequisitoPagoGrupoRequisito.save(request)
                    log(f"Editó eRequisitoPagoGrupoRequisito de revisión: {eRequisitoPagoGrupoRequisito}", request,
                        'change')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'editgestorrequisitopagoorden':
            try:
                id = request.POST.get('id', None)
                eRequisitoPagoGrupoRequisito = RequisitoPagoGrupoRequisito.objects.get(status=True, id=int(encrypt(id)))
                form = RequisitoPagoDipOrdenForm(request.POST)
                if form.is_valid():
                    eRequisitoPagoGrupoRequisito.orden = form.cleaned_data['orden']
                    eRequisitoPagoGrupoRequisito.opcional = form.cleaned_data['opcional']
                    eRequisitoPagoGrupoRequisito.save(request)
                    log(f"Editó eRequisitoPagoGrupoRequisito de revisión: {eRequisitoPagoGrupoRequisito}", request,
                        'change')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "Error: {}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'adddepposgrado':
            try:
                form = DepartamentoPosgradoForm(request.POST)
                id_resp = request.POST['responsable']
                if not id_resp:
                    raise NameError("El campo responsable es obligatorio!")
                form.edit(id_resp)
                ids_items = request.POST.get('lista_items1', None)
                if ids_items: form.res_subrogante(json.loads(ids_items))
                if not form.is_valid():
                    raise NameError(f"{[{k: v[0]} for k, v in form.errors.items()]}")
                dep = Departamento(
                    nombre=form.cleaned_data['nombre'],
                    responsable=form.cleaned_data['responsable'],
                )
                dep.save(request)
                dep.responsable_subrogante.set(form.cleaned_data['responsable_subrogante'])
                log(f"Agregó departamento de posgrado: {dep.__str__()}", request, 'add')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "{}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'editdepposgrado':
            try:
                id = encrypt(request.POST.get('id', None))
                filtro = Departamento.objects.get(status=True, id=int(id))
                form = DepartamentoPosgradoForm(request.POST)
                id_resp = request.POST['responsable']
                if not id_resp:
                    raise NameError("El campo responsable es obligatorio!")
                form.edit(id_resp)
                ids_items = request.POST.get('lista_items1', None)
                if ids_items: form.res_subrogante(json.loads(ids_items))
                if not form.is_valid():
                    raise NameError(f"{[{k: v[0]} for k, v in form.errors.items()]}")
                filtro.nombre = form.cleaned_data['nombre']
                filtro.responsable = form.cleaned_data['responsable']
                filtro.save(request)
                filtro.responsable_subrogante.clear()
                filtro.responsable_subrogante.set(form.cleaned_data['responsable_subrogante'])
                log(f"Edito departamento de posgrado: {filtro.__str__()}", request, 'change')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "{}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'deletedepposgrado':
            try:
                id = int(encrypt(request.POST['id']))
                reg = Departamento.objects.get(status=True, id=id)
                reg.status = False
                reg.save(request)
                log(f"Eliminó departamento de posgrado: {reg}", request, 'del')
                res_json = {"error": False}
            except Exception as ex:
                err_ = "{}({})".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {'error': True, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'addgestposgrado':
            try:
                id_dep = request.POST.get('ids', None)
                departamento = Departamento.objects.get(status=True, id=int(encrypt(id_dep)))
                form = GestionPosgradoForm(request.POST)
                id_resp = request.POST['responsable']
                id_respsub = request.POST['responsablesubrogante']
                if not id_resp:
                    raise NameError("El campo responsable es obligatorio!")
                form.edit(id_resp)
                if id_respsub: form.res_subrogante(id_respsub)
                if not form.is_valid():
                    raise NameError(f"{[{k: v[0]} for k, v in form.errors.items()]}")
                gest = Gestion(
                    departamento=departamento,
                    gestion=form.cleaned_data['gestion'],
                    cargo=form.cleaned_data['cargo'],
                    responsable=form.cleaned_data['responsable'],
                    responsablesubrogante=form.cleaned_data['responsablesubrogante']
                )
                gest.save(request)
                log(f"Agregó al departamento una gestión de posgrado: {gest.__str__()}", request, 'add')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "{}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'subir_requisito_contrato':
            try:
                id_contrato = int(request.POST.get('idc', '0'))
                if id_contrato == 0:
                    raise NameError("Parametro no encontrado  contrato")

                id_requisito = int(request.POST.get('id', '0'))
                if id_requisito == 0:
                    raise NameError("Parametro no encontrado  requisito")

                eContratoDip = ContratoDip.objects.get(pk=id_contrato)
                eRequisito = Requisito.objects.get(pk=id_requisito)
                form = RequisitoContratoForm(request.POST, request.FILES)
                if form.is_valid() and request.FILES.get('archivo', None):
                    newfile = request.FILES.get('archivo')
                    if newfile:
                        if newfile.size > 6291456:
                            return JsonResponse({"result": "bad", "mensaje": u"Error, archivo mayor a 6 Mb."})
                        else:
                            newfilesd = newfile._name
                            ext = newfilesd[newfilesd.rfind("."):].lower()
                            if ext == '.pdf':
                                _name = generar_nombre(f"{eRequisito.__str__()}", '')
                                _name = remover_caracteres_tildes_unicode(
                                    remover_caracteres_especiales_unicode(_name)).lower().replace(' ', '_').replace('-',
                                                                                                                    '_')
                                newfile._name = generar_nombre(u"%s_" % _name, f"{_name}.pdf")

                                eContratoRequisito = ContratoRequisito(
                                    contratodip=eContratoDip,
                                    requisito=eRequisito,
                                    archivo=newfile,
                                    fecha_caducidad=form.cleaned_data['fecha_caducidad']
                                )
                                eContratoRequisito.save(request)
                            else:
                                return JsonResponse({"result": "bad", "mensaje": u"Error, Solo archivos PDF"})
                else:
                    raise NameError(f"{[{k: v[0]} for k, v in form.errors.items()]}")

                res_json = {"result": True}

            except Exception as ex:
                err_ = "{}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'actualizar_fecha_caducidad_requisito':
            try:
                pk = int(request.POST.get('id', '0'))
                if pk == 0:
                    raise NameError("Parametro no encontrado  contrato")

                form = RequisitoContratoForm(request.POST)
                form.del_field('archivo')
                if form.is_valid():
                    eContratoRequisito = ContratoRequisito.objects.get(pk=pk)
                    eContratoRequisito.fecha_caducidad = form.cleaned_data['fecha_caducidad']
                    eContratoRequisito.save(request)
                else:
                    raise NameError(f"{[{k: v[0]} for k, v in form.errors.items()]}")

                res_json = {"result": True}

            except Exception as ex:
                err_ = "{}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'actualizar_requisito_contratacion':
            try:
                pk = int(request.POST.get('id', '0'))
                if pk == 0:
                    raise NameError("Parametro no encontrado  contrato")

                eContratoRequisito = ContratoRequisito.objects.get(pk=pk)
                form = RequisitoContratoForm(request.POST, request.FILES)
                if form.is_valid() and request.FILES.get('archivo', None):
                    newfile = request.FILES.get('archivo')
                    if newfile:
                        if newfile.size > 6291456:
                            return JsonResponse({"result": "bad", "mensaje": u"Error, archivo mayor a 6 Mb."})
                        else:
                            newfilesd = newfile._name
                            ext = newfilesd[newfilesd.rfind("."):].lower()
                            if ext == '.pdf':
                                _name = generar_nombre(f"{eContratoRequisito.requisito.__str__()}", '')
                                _name = remover_caracteres_tildes_unicode(
                                    remover_caracteres_especiales_unicode(_name)).lower().replace(' ', '_').replace('-',
                                                                                                                    '_')
                                newfile._name = generar_nombre(u"%s_" % _name, f"{_name}.pdf")
                                eContratoRequisito.archivo = newfile
                                eContratoRequisito.fecha_caducidad = form.cleaned_data['fecha_caducidad']
                                eContratoRequisito.save(request)
                            else:
                                return JsonResponse({"result": "bad", "mensaje": u"Error, Solo archivos PDF"})
                else:
                    raise NameError(f"{[{k: v[0]} for k, v in form.errors.items()]}")

                res_json = {"result": True}

            except Exception as ex:
                err_ = "{}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'editgestposgrado':
            try:
                id = encrypt(request.POST.get('id', None))
                filtro = Gestion.objects.get(status=True, id=int(id))
                form = GestionPosgradoForm(request.POST)
                id_resp = request.POST['responsable']
                id_respsub = request.POST['responsablesubrogante']
                if not id_resp:
                    raise NameError("El campo responsable es obligatorio!")
                form.edit(id_resp)
                if id_respsub: form.res_subrogante(id_respsub)
                if not form.is_valid():
                    raise NameError(f"{[{k: v[0]} for k, v in form.errors.items()]}")
                filtro.gestion = form.cleaned_data['gestion']
                filtro.cargo = form.cleaned_data['cargo']
                filtro.responsable = form.cleaned_data['responsable']
                filtro.responsablesubrogante = form.cleaned_data['responsablesubrogante']
                filtro.save(request)
                log(f"Edito al departamento una gestión de posgrado: {filtro.__str__()}", request, 'change')
                res_json = {"result": True}
            except Exception as ex:
                err_ = "{}. En la linea {}".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {"result": False, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'edit_orden_firma_acta_pago':
            try:
                f = OrdenFirmaInformeActaPagoForm(request.POST)
                eOrdenFirmaActaPago = OrdenFirmaActaPago.objects.get(pk=request.POST.get('id'))
                if f.is_valid():
                    eOrdenFirmaActaPago.responsabilidadfirma = f.cleaned_data.get('responsabilidadfirma')
                    eOrdenFirmaActaPago.orden = f.cleaned_data.get('orden')
                    eOrdenFirmaActaPago.save(request)
                    log(u"Editó orden firma acta de pago", request, 'edit')
                    return JsonResponse({"result": True})
                else:
                    return JsonResponse({"result": False, "mensaje": f"Error al validar los datos.",
                                         "form": [{k: v[0]} for k, v in f.errors.items()]})
            except Exception as ex:
                transaction.rollback()
                return JsonResponse({"result": False, "mensaje": f"Solicitud incorrecta. {ex.__str__()}"})

        elif action == 'deletegestposgrado':
            try:
                id = int(encrypt(request.POST['id']))
                reg = Gestion.objects.get(status=True, id=id)
                reg.status = False
                reg.save(request)
                log(f"Eliminó al departamento una gestión de posgrado: {reg}", request, 'del')
                res_json = {"error": False}
            except Exception as ex:
                err_ = "{}({})".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {'error': True, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'delhorariocontrato':
            try:
                id = int(request.POST['id'])
                reg = HorarioPlanificacionContrato.objects.get(status=True, id=id)
                eConvocatoria = reg.contrato
                reg.status = False
                reg.save(request)
                log(f"Agregó horario a contrato {eConvocatoria}", request, 'del')
                res_json = {"error": False}
            except Exception as ex:
                err_ = "{}({})".format(ex, sys.exc_info()[-1].tb_lineno)
                transaction.set_rollback(True)
                res_json = {'error': True, "mensaje": err_}
            return JsonResponse(res_json)

        elif action == 'delete_orden_firma_acta_pago':
            try:
                eOrdenFirmaActaPago = OrdenFirmaActaPago.objects.get(pk=int(request.POST['id']))
                log(u'Elimino orden firma: %s' % eOrdenFirmaActaPago, request, "del")
                eOrdenFirmaActaPago.status = False
                eOrdenFirmaActaPago.save(request)
                return JsonResponse({"result": True, "error": False})
            except Exception as ex:
                transaction.set_rollback(True)
                return JsonResponse({"result": "bad", "mensaje": u"Error al eliminar los datos."})

        elif action == 'save-titulo-acta-pago':
            try:
                f = ConfiguracionActaPagoTituloForm(request.POST)
                eConfiguracionActaPagos = ConfiguracionActaPago.objects.filter(status = True)
                if f.is_valid():
                    if eConfiguracionActaPagos.exists():
                        eConfiguracionActaPago = eConfiguracionActaPagos.first()
                        eConfiguracionActaPago.titulo = f.cleaned_data['titulo']
                        eConfiguracionActaPago.save(request)

                    else:
                        eConfiguracionActaPago = ConfiguracionActaPago (
                            titulo=f.cleaned_data['titulo']
                        )
                        eConfiguracionActaPago.save(request)

                return JsonResponse({"result": True, "error": False})

            except Exception as ex:
                return HttpResponseRedirect("/adm_contratodip?action=configuracionactapago&info=%s" % ex.__str__())

        elif action == 'save-objetivo-acta-pago':
            try:
                f = ConfiguracionActaPagoObjetivoForm(request.POST)
                eConfiguracionActaPagos = ConfiguracionActaPago.objects.filter(status = True)
                if f.is_valid():
                    if eConfiguracionActaPagos.exists():
                        eConfiguracionActaPago = eConfiguracionActaPagos.first()
                        eConfiguracionActaPago.objetivo = f.cleaned_data['objetivo']
                        eConfiguracionActaPago.save(request)

                    else:
                        eConfiguracionActaPago = ConfiguracionActaPago (
                            objetivo=f.cleaned_data['objetivo']
                        )
                        eConfiguracionActaPago.save(request)


                return JsonResponse({"result": True, "error": False})
            except Exception as ex:
                return HttpResponseRedirect("/adm_contratodip?action=configuracionactapago&info=%s" % ex.__str__())

        elif action == 'save-marco-juridico-acta-pago':
            try:
                f = ConfiguracionActaPagoMarcoJuridicoReferencialForm(request.POST)
                eConfiguracionActaPagos = ConfiguracionActaPago.objects.filter(status = True)
                if f.is_valid():
                    if eConfiguracionActaPagos.exists():
                        eConfiguracionActaPago = eConfiguracionActaPagos.first()
                        eConfiguracionActaPago.marcojuridicoreferencial = f.cleaned_data['marcojuridicoreferencial']
                        eConfiguracionActaPago.save(request)


                    else:
                        eConfiguracionActaPago = ConfiguracionActaPago (
                            marcojuridicoreferencial=f.cleaned_data['marcojuridicoreferencial']
                        )
                        eConfiguracionActaPago.save(request)

                return JsonResponse({"result": True, "error": False})
            except Exception as ex:
                return HttpResponseRedirect("/adm_contratodip?action=configuracionactapago&info=%s" % ex.__str__())

        elif action == 'save-conclusion-acta-pago':
            try:
                f = ConfiguracionActaPagoConclusionesForm(request.POST)
                eConfiguracionActaPagos = ConfiguracionActaPago.objects.filter(status = True)
                if f.is_valid():
                    if eConfiguracionActaPagos.exists():
                        eConfiguracionActaPago = eConfiguracionActaPagos.first()
                        eConfiguracionActaPago.conclusiones = f.cleaned_data['conclusiones']
                        eConfiguracionActaPago.save(request)
                    else:
                        eConfiguracionActaPago = ConfiguracionActaPago (
                            conclusiones=f.cleaned_data['conclusiones']
                        )
                        eConfiguracionActaPago.save(request)

                return JsonResponse({"result": True, "error": False})
            except Exception as ex:
                return HttpResponseRedirect("/adm_contratodip?action=configuracionactapago&info=%s" % ex.__str__())

        elif action == 'save-recomendaciones-acta-pago':
            try:
                f = ConfiguracionActaPagoRecomendacionesForm(request.POST)
                eConfiguracionActaPagos = ConfiguracionActaPago.objects.filter(status = True)
                if f.is_valid():
                    if eConfiguracionActaPagos.exists():
                        eConfiguracionActaPago = eConfiguracionActaPagos.first()
                        eConfiguracionActaPago.recomendaciones = f.cleaned_data['recomendaciones']
                        eConfiguracionActaPago.save(request)
                    else:
                        eConfiguracionActaPago = ConfiguracionActaPago (
                            recomendaciones=f.cleaned_data['recomendaciones']
                        )
                        eConfiguracionActaPago.save(request)

                return JsonResponse({"result": True, "error": False})
            except Exception as ex:
                return HttpResponseRedirect("/adm_contratodip?action=configuracionactapago&info=%s" % ex.__str__())

        elif action == 'save-solicitadopor-acta-pago':
            try:
                f = ConfiguracionActaPagoSolicitadoPorForm(request.POST)
                id_solicitado_por = request.POST['solicitadopor']
                if not id_solicitado_por:
                    raise NameError("El campo solicitado por es obligatorio!")
                f.edit(id_solicitado_por)
                eConfiguracionActaPagos = ConfiguracionActaPago.objects.filter(status = True)
                if f.is_valid():
                    if eConfiguracionActaPagos.exists():
                        eConfiguracionActaPago = eConfiguracionActaPagos.first()
                        eConfiguracionActaPago.solicitadopor = f.cleaned_data['solicitadopor']
                        eConfiguracionActaPago.save(request)

                    else:
                        eConfiguracionActaPago = ConfiguracionActaPago (
                            solicitadopor=f.cleaned_data['solicitadopor']
                        )
                        eConfiguracionActaPago.save(request)


                return JsonResponse({"result": True, "error": False})
            except Exception as ex:
                return HttpResponseRedirect("/adm_contratodip?action=configuracionactapago&info=%s" % ex.__str__())

        elif action == 'save-para-acta-pago':
            try:
                f = ConfiguracionActaPagoParaForm(request.POST)
                id_para = request.POST['para']
                if not id_para:
                    raise NameError("El campo para  es obligatorio!")
                f.edit(id_para)
                eConfiguracionActaPagos = ConfiguracionActaPago.objects.filter(status = True)
                if f.is_valid():
                    if eConfiguracionActaPagos.exists():
                        eConfiguracionActaPago = eConfiguracionActaPagos.first()
                        eConfiguracionActaPago.para = f.cleaned_data['para']
                        eConfiguracionActaPago.save(request)

                    else:
                        eConfiguracionActaPago = ConfiguracionActaPago (
                            para=f.cleaned_data['para']
                        )
                        eConfiguracionActaPago.save(request)


                return JsonResponse({"result": True, "error": False})
            except Exception as ex:
                return HttpResponseRedirect("/adm_contratodip?action=configuracionactapago&info=%s" % ex.__str__())


        elif action == 'delete_seccion_acta_pago':
            try:
                with transaction.atomic():
                    instancia = ConfiguracionActaPagoItem.objects.get(pk=int(request.POST['id']))
                    instancia.status = False
                    instancia.save(request)
                    log(u'Elimino seccion acta pago posgrado: %s' % instancia, request, "delete")
                    res_json = {"error": False}
            except Exception as ex:
                res_json = {'error': True, "message": "Error: {}".format(ex)}
            return JsonResponse(res_json, safe=False)


        return JsonResponse({"result": "bad", "mensaje": u"Solicitud incorrecta. %s" % ex.__str__()})
    else:
        if 'action' in request.GET:
            data['action'] = action = request.GET['action']

            if action == 'addperfil':
                try:
                    data['form2'] = PerfilPuestoDipForm()
                    data['add'] = True
                    template = get_template("adm_contratodip/modal/formperfil.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'addactividades':
                try:
                    data['form2'] = ActividadesForm()
                    data['add'] = True
                    template = get_template("adm_contratodip/modal/formperfil.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'addsecuencia':
                try:
                    data['form2'] = SecuenciaMemoActPosgradoForm()
                    template = get_template("adm_contratodip/modal/formsecuencia.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'editperfil':
                try:
                    data['filtro'] = filtro = PerfilPuestoDip.objects.get(pk=int(request.GET['id']))
                    data['form2'] = PerfilPuestoDipForm(initial=model_to_dict(filtro))
                    data['actividades'] = ActividadesPerfil.objects.filter(status=True).order_by('-id')
                    if filtro.actividadesperfil().exists():
                        data['acti'] = filtro.actividadesperfil()
                    template = get_template("adm_contratodip/modal/formperfil.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'addactividadperfil':
                try:
                    actividades2 = []
                    if 'actividades' in request.GET:
                        actividades = json.loads(request.GET['actividades'])
                    if 'idperfil' in request.GET:
                        idp = int(request.GET['idperfil'])
                        for act in actividades:
                            if not ActividadesContratoPerfil.objects.filter(status=True, perfil_id=idp,
                                                                            actividad_id=act).exists():
                                acti = ActividadesContratoPerfil(
                                    perfil_id=idp,
                                    actividad_id=act
                                )
                                acti.save(request)
                                actividades2.append(
                                    {'actividad': acti.actividad.id, 'perfil': acti.perfil.id, 'id': acti.pk,
                                     'descripcion': acti.actividad.descripcion})
                                log('Adiciono la actividad %s al cargo %s' % (acti.actividad.descripcion, acti.perfil),
                                    request, 'add')
                    if 'idcontrato' in request.GET:
                        idcon = int(request.GET['idcontrato'])
                        if not ContratoDip.objects.filter(status=True, pk=idcon).exists():
                            return JsonResponse({'result': 'bad'})
                        contrato = ContratoDip.objects.get(status=True, pk=idcon)
                        for act in actividades:
                            if not ActividadesContratoPerfil.objects.filter(status=True, contrato=contrato,
                                                                            perfil=contrato.cargo,
                                                                            actividad_id=act).exists():
                                acti = ActividadesContratoPerfil(
                                    contrato_id=idcon,
                                    actividad_id=act
                                )
                                acti.save(request)
                                actividades2.append(
                                    {'actividad': acti.actividad.id, 'contrato': acti.contrato.id, 'id': acti.pk,
                                     'descripcion': acti.actividad.descripcion})
                                log('Adiciono la actividad %s al contrato %s' % (
                                acti.actividad.descripcion, acti.contrato), request, 'add')
                    return JsonResponse({'result': True, 'actividades': actividades2}, safe=True)
                except Exception as ex:
                    pass

            elif action == 'cambiarestadoact':
                try:
                    if ActividadesContratoPerfil.objects.filter(pk=int(request.GET['id'])).exists():
                        acti = ActividadesContratoPerfil.objects.filter(pk=int(request.GET['id']))[0]
                        acti.obligatoria = json.loads(request.GET['valor'])
                        acti.save(request)
                        log('Se cambio el estado obligatorio de la actividad: %s' % (acti.actividad), request, 'edit')
                        return JsonResponse({'result': True}, safe=True)
                    return JsonResponse({'result': False}, safe=True)
                except Exception as ex:
                    transaction.set_rollback(True)
                    return JsonResponse({"result": True, "mensaje": "Datos erróneos, intente nuevamente."}, safe=False)

            elif action == 'actividadesext':
                try:
                    data['contratodip'] = contrato = ContratoDip.objects.get(status=True, pk=int(request.GET['id']))
                    data['actividades'] = filtro = ActividadesPerfil.objects.filter(status=True)
                    if contrato.actividadescontratoperfil_set.filter(status=True).all():
                        data['acti'] = contrato.actividadescontratoperfil_set.filter(status=True).all()
                    template = get_template("adm_contratodip/modal/formactividades.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'editactvidades':
                try:
                    data['filtro'] = filtro = ActividadesPerfil.objects.get(pk=int(request.GET['id']))
                    data['form2'] = ActividadesForm(initial=model_to_dict(filtro))
                    data['add'] = True
                    template = get_template("adm_contratodip/modal/formperfil.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'editsecuencia':
                try:
                    data['filtro'] = filtro = SecuenciaMemoActividadPosgrado.objects.get(pk=int(request.GET['id']))
                    data['form2'] = SecuenciaMemoActPosgradoForm(initial=model_to_dict(filtro))
                    template = get_template("adm_contratodip/modal/formperfil.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'addcampo':
                try:
                    form2 = CampoContratoDipForm()
                    form2.add()
                    data['form2'] = form2
                    template = get_template("adm_contratodip/modal/formcampo.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'editcampo':
                try:

                    data['filtro'] = filtro = CampoContratoDip.objects.get(pk=int(request.GET['id']), status=True)
                    form2 = CampoContratoDipForm(initial=model_to_dict(filtro))
                    if filtro.en_uso():
                        form2.editar()
                    data['form2'] = form2
                    template = get_template("adm_contratodip/modal/formcampo.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'addplantilla':

                try:
                    data['title'] = u'Nueva Plantilla'
                    data['form'] = PlantillaContratoDipForm()
                    data['campos'] = CampoContratoDip.objects.filter(status=True)

                    return render(request, "adm_contratodip/addplantilla.html", data)
                except Exception as ex:
                    pass

            elif action == 'editplantilla':
                try:
                    data['title'] = u'Modificación de Plantilla'
                    data['plantilla'] = plantilla = PlantillaContratoDip.objects.get(pk=int(request.GET['id']),
                                                                                     status=True)
                    data['camposseleccion'] = CampoPlantillaContratoDip.objects.filter(contrato=plantilla)
                    data['campos'] = CampoContratoDip.objects.filter(status=True)
                    data['form'] = form = PlantillaContratoDipForm(initial=model_to_dict(plantilla))
                    return render(request, "adm_contratodip/editplantilla.html", data)
                except Exception as ex:
                    pass

            elif action == 'memo':
                try:
                    if not (
                    HistorialPagoMes.objects.values('id').filter(status=True, contrato_id=int(request.GET['id']),
                                                                 fecha_pago__month=int(request.GET['desde']),
                                                                 fecha_pago__year=int(request.GET['anio'])).exists()):
                        historial = HistorialPagoMes(
                            contrato_id=int(request.GET['id']),
                            fecha_pago=date(year=int(request.GET['anio']), month=int(request.GET['desde']),
                                            day=date.today().day),
                            cancelado=False
                        )
                        historial.save(request)
                        log('Agrega historial de pago %s' % (historial), request, 'add')
                    else:
                        historial = HistorialPagoMes.objects.get(status=True, contrato_id=int(request.GET['id']),
                                                                 fecha_pago__month=int(request.GET['desde']),
                                                                 fecha_pago__year=int(request.GET['anio']))
                    ids = request.GET['id']
                    desde = request.GET['desde']
                    contratopersona = ContratoDip.objects.filter(pk=int(ids))[0]
                    cumplehoras = TrabajadorDiaJornada.objects.filter(persona=contratopersona.persona, mes=int(desde),
                                                                      status=True)

                    if desde == '':
                        return JsonResponse({'result': False, 'mensaje': 'Debe ingresar una fecha válida'})
                    if MemoActividadPosgrado.objects.filter(status=True, contrato=contratopersona,
                                                            mes=int(desde)).exists():
                        return JsonResponse({'result': False, 'mensaje': 'Ya existe el registro con el mes de %s' % (
                        MESES_CHOICES[int(desde) - 1][1])})
                    if PlantillaInformes.objects.values('id').filter(tipo=0, vigente=True, anio=2022,
                                                                     status=True).exists():
                        informe = PlantillaInformes.objects.filter(tipo=0, vigente=True, anio=2022, status=True)[0]
                    else:
                        return JsonResponse({'result': False, 'mensaje': 'No Existe la Plantilla'})
                    # desde = date(desde)
                    secpos = \
                    SecuenciaMemoActividadPosgrado.objects.filter(tipo=PlantillaContratoDip.objects.last(), status=True,
                                                                  anioejercicio__anioejercicio=datetime.now().year)[0]
                    memo = contratopersona.secuencia_codigo()
                    nombre_plantilla = os.path.join(SITE_ROOT, 'media', informe.archivo.name)
                    nombre_contrato = 'Memo_' + str(contratopersona.id) + random.randint(
                        1, 10000).__str__() + ".docx"
                    direccion_contrato = os.path.join(SITE_ROOT, 'media', 'contratoepunemi', 'memo')
                    filename_contrato = os.path.join(direccion_contrato, nombre_contrato)
                    memos = MemoActividadPosgrado(secuencia=memo, secuenciamemo=secpos, contrato=contratopersona,
                                                  mes=int(desde), historialpago=historial)
                    # guarda la direccion
                    cantidad_parrafo = 0
                    document = Document(nombre_plantilla)
                    parrafo = document.paragraphs
                    table = document.tables
                    cantidad_parrafo = parrafo.__len__()
                    cantidad_tablas = table.__len__()
                    n = 0
                    persona_nombre = contratopersona.persona
                    persona_cedula = contratopersona.persona.cedula
                    rmu = contratopersona.rmu
                    iva = contratopersona.valoriva
                    total = contratopersona.valortotal
                    cargo = contratopersona.cargo
                    codcontrato = contratopersona.codigocontrato
                    partida = contratopersona.certificacion.codigo
                    fecha = str(date.today().day) + " de " + str(
                        MESES_CHOICES[int(date.today().month) - 1][1]) + " del " + str(date.today().year)
                    fechaletra = MESES_CHOICES[int(desde) - 1][1] + ' del ' + str(date.today().year)
                    secuencia = '0' + str(memos) if memos.secuencia > 9 else '00' + str(memos)
                    cadena1 = '${secuencia}'
                    cadena2 = '${fecha}'
                    cadena3 = '${letra}'
                    cadena4 = '${cargo}'
                    cadena5 = '${trabajador}'
                    cadena6 = '${contrato}'
                    cadena7 = '${partida}'
                    cadena8 = '${rmu}'
                    cadena9 = '${iva}'
                    cadena10 = '${total_mes}'

                    for n in range(cantidad_parrafo):
                        for run in parrafo[n].runs:
                            run.text = run.text.replace(cadena1, str(secuencia.upper()))
                            run.text = run.text.replace(cadena2, str(fecha))
                            run.text = run.text.replace(cadena3, str(fechaletra))
                            run.text = run.text.replace(cadena4, str(cargo))
                            run.text = run.text.replace(cadena5, str(persona_nombre))
                            run.text = run.text.replace(cadena6, str(codcontrato))
                            run.text = run.text.replace(cadena7, str(partida))
                    for k in range(cantidad_tablas):
                        for run in table[k].rows[1].cells:
                            run.text = run.text.replace(cadena2, str(MESES_CHOICES[int(desde) - 1][1]))
                            run.text = run.text.replace(cadena4, str(cargo))
                            run.text = run.text.replace(cadena5, str(persona_nombre))
                            run.text = run.text.replace(cadena8, str(rmu))
                            run.text = run.text.replace(cadena9, str(iva))
                            run.text = run.text.replace(cadena10, str(total))
                    try:
                        os.stat(direccion_contrato)
                    except:
                        os.mkdir(direccion_contrato)
                    document.save(filename_contrato)

                    memos.archivo.name = "contratoepunemi/memo/%s" % nombre_contrato
                    memos.save(request)
                    log(u'Contrato Persona: %s' % contratopersona, request, "edit")
                    return JsonResponse({"result": True, "mensaje": 'Memo generado correctamente'})

                except Exception as ex:
                    transaction.set_rollback(True)
                    return JsonResponse({"result": "bad", "mensaje": "Error al procesar los datos."})

            elif action == 'bitacora':
                try:
                    ids = request.GET['id']
                    desde = request.GET['desde']
                    data['distributivo'] = contratopersona = ContratoDip.objects.filter(pk=int(ids))[0]
                    data['persona'] = personad = Persona.objects.get(pk=contratopersona.persona.id)
                    cumplehoras = TrabajadorDiaJornada.objects.filter(persona=contratopersona.persona, mes=int(desde),
                                                                      status=True)
                    if InformeActividadJornada.objects.values('id').filter(mes=int(desde), status=True,
                                                                           contrato=contratopersona).exists():
                        return JsonResponse({'result': False, 'mensaje': 'Ya existe el registro con el mes de %s' % (
                        MESES_CHOICES[int(desde) - 1][1])})
                    secinforme = contratopersona.secuencia_informe()
                    if not (
                    HistorialPagoMes.objects.values('id').filter(status=True, contrato_id=int(request.GET['id']),
                                                                 fecha_pago__month=int(request.GET['desde']),
                                                                 fecha_pago__year=int(request.GET['anio'])).exists()):
                        historial = HistorialPagoMes(
                            contrato_id=int(request.GET['id']),
                            fecha_pago=date(year=int(request.GET['anio']), month=int(request.GET['desde']),
                                            day=date.today().day),
                            cancelado=False
                        )
                        historial.save(request)
                        log('Agrega historial de pago %s' % (historial), request, 'add')
                    else:
                        historial = HistorialPagoMes.objects.get(status=True, contrato_id=int(request.GET['id']),
                                                                 fecha_pago__month=int(request.GET['desde']),
                                                                 fecha_pago__year=int(request.GET['anio']))
                    name_file = 'Informe_actividad_' + str(secinforme) + random.randint(1, 10000).__str__() + '.pdf'
                    folder = 'contratoepunemi/informe/'
                    try:
                        os.stat(os.path.join(SITE_STORAGE, 'media', folder))
                    except:
                        os.mkdir(os.path.join(SITE_STORAGE, 'media', folder))
                    secpos = \
                    SecuenciaMemoActividadPosgrado.objects.filter(tipo=PlantillaContratoDip.objects.last(), status=True,
                                                                  anioejercicio__anioejercicio=datetime.now().year)[0]
                    informe = InformeActividadJornada(mes=int(desde), contrato=contratopersona, secuenciageneral=secpos,
                                                      secuencia=secinforme, historialpago=historial)
                    # data = extraervalores(contratopersona.persona.id,request.GET['desde'],request.GET['anio'],True if 'h' in request.GET else False)
                    data['bitacora'] = BitacoraActividadDiaria.objects.filter(status=True,
                                                                              persona=contratopersona.persona,
                                                                              fecha__month=int(desde))
                    data['secuencia'] = '0' + str(informe) if informe.secuencia > 9 else '00' + str(informe)
                    # data['dia'] = range(1,calendar.monthrange(int(request.GET['anio']),int(request.GET['desde']))[1]+1)
                    data['mes'] = request.GET['desde']
                    data['anio'] = request.GET['anio']
                    data['mes_nombre'] = MESES_CHOICES[int(desde) - 1][1]
                    pdf = conviert_html_to_pdf_save_informe('adm_contratodip/detalle_jornada_pdf.html',
                                                            {'pagesize': 'A4', 'data': data}, folder,
                                                            name_file)
                    if pdf:
                        informe.archivo.name = os.path.join(folder, name_file)
                        informe.save(request)
                        mensaje = {"result": True, "mensaje": 'Informe generado correctamente'}
                    else:
                        mensaje = {"result": False, "mensaje": 'Error al guardar pdf'}
                    return JsonResponse(mensaje)
                except Exception as ex:
                    print(ex)
                    transaction.set_rollback(True)
                    return JsonResponse({"result": False, "mensaje": "Error al procesar los datos."})

            elif action == 'inftecnico':
                try:
                    ids = request.GET['id']
                    desde = request.GET['desde']
                    data['persona'] = contratopersona = ContratoDip.objects.filter(pk=int(ids))[0]
                    fechas = TrabajadorDiaJornada.objects.values_list('fecha').filter(persona=contratopersona.persona,
                                                                                      anio=2022, mes=int(desde),
                                                                                      status=True).order_by('fecha')
                    dias_no_laborable = DiasNoLaborable.objects.values_list('fecha').filter(fecha__in=fechas).exclude(
                        periodo__isnull=False)
                    data['dias'] = cumplehoras = TrabajadorDiaJornada.objects.filter(persona=contratopersona.persona,
                                                                                     anio=2022, mes=int(desde),
                                                                                     status=True).exclude(
                        fecha__in=dias_no_laborable).order_by('fecha')
                    if InformeTecnico.objects.values('id').filter(mes=int(desde), status=True,
                                                                  contrato=contratopersona).exists():
                        return JsonResponse({'result': False, 'mensaje': 'Ya existe el registro con el mes de %s' % (
                        MESES_CHOICES[int(desde) - 1][1])})
                    if not InformeActividadJornada.objects.values('id').filter(mes=int(desde), status=True,
                                                                               contrato=contratopersona).exists():
                        return JsonResponse(
                            {'result': False, 'mensaje': 'No ha generado el reporte de asistencia y actividades'})
                    if not (
                    HistorialPagoMes.objects.values('id').filter(status=True, contrato_id=int(request.GET['id']),
                                                                 fecha_pago__month=int(request.GET['desde']),
                                                                 fecha_pago__year=int(request.GET['anio'])).exists()):
                        historial = HistorialPagoMes(
                            contrato_id=int(request.GET['id']),
                            fecha_pago=date(year=int(request.GET['anio']), month=int(request.GET['desde']),
                                            day=date.today().day),
                            cancelado=False
                        )
                        historial.save(request)
                        log('Agrega historial de pago %s' % (historial), request, 'add')
                    else:
                        historial = HistorialPagoMes.objects.get(status=True, contrato_id=int(request.GET['id']),
                                                                 fecha_pago__month=int(request.GET['desde']),
                                                                 fecha_pago__year=int(request.GET['anio']))
                    data['info'] = InformeActividadJornada.objects.filter(mes=int(desde), status=True,
                                                                          contrato=contratopersona).last()
                    d = data['info'].archivo.name.split('\\')
                    data['nombrearchivo'] = d[-1].split('.')[0]
                    secinforme = contratopersona.secuencia_inftecnico()

                    name_file = 'Informe_Tecnico_' + str(secinforme) + random.randint(1, 10000).__str__() + '.pdf'
                    folder = 'contratoepunemi/inftec/'
                    try:
                        os.stat(os.path.join(SITE_STORAGE, 'media', folder))
                    except:
                        os.mkdir(os.path.join(SITE_STORAGE, 'media', folder))
                    secpos = \
                    SecuenciaMemoActividadPosgrado.objects.filter(tipo=PlantillaContratoDip.objects.last(), status=True,
                                                                  anioejercicio__anioejercicio=datetime.now().year)[0]
                    informe = InformeTecnico(mes=int(desde), contrato=contratopersona, secuenciageneral=secpos,
                                             secuencia=secinforme, historialpago=historial)
                    # data = extraervalores(contratopersona.persona.id,request.GET['desde'],request.GET['anio'],True if 'h' in request.GET else False)
                    data['fehca'] = date.today()
                    data['mes'] = MESES_CHOICES[datetime.today().month - 1]
                    data['secuencia'] = '0' + str(informe) if informe.secuencia > 9 else '00' + str(informe)
                    pdf = conviert_html_to_pdf_save_informe('adm_contratodip/informe_tecnico_pdf.html',
                                                            {'pagesize': 'A4', 'data': data}, folder,
                                                            name_file)
                    if pdf:
                        informe.archivo.name = os.path.join(folder, name_file)
                        informe.save(request)
                        mensaje = {"result": True, "mensaje": 'Informe generado correctamente'}
                    else:
                        mensaje = {"result": False, "mensaje": 'Error al guardar pdf'}
                    return JsonResponse(mensaje)
                except Exception as ex:
                    print(ex)
                    transaction.set_rollback(True)
                    return JsonResponse({"result": False, "mensaje": "Error al procesar los datos."})

            elif action == 'actapago':
                try:
                    ids = request.GET['id']
                    desde = request.GET['desde']
                    data['elaborado'] = persona
                    data['persona'] = contratopersona = ContratoDip.objects.filter(pk=int(ids))[0]
                    cumplehoras = TrabajadorDiaJornada.objects.filter(persona=contratopersona.persona, mes=int(desde),
                                                                      status=True)
                    if not InformeTecnico.objects.values('id').filter(mes=int(desde), status=True,
                                                                      contrato=contratopersona).exists():
                        return JsonResponse({'result': False, 'mensaje': 'No ha generado el Informe'})
                    if not InformeActividadJornada.objects.values('id').filter(mes=int(desde), status=True,
                                                                               contrato=contratopersona).exists():
                        return JsonResponse(
                            {'result': False, 'mensaje': 'No ha generado el reporte de asistencia y actividades'})
                    if ActaPago.objects.values('id').filter(mes=int(desde), status=True,
                                                            contrato=contratopersona).exists():
                        return JsonResponse({'result': False, 'mensaje': 'Ya existe el registro con el mes de %s' % (
                        MESES_CHOICES[int(desde) - 1][1])})
                    secinforme = contratopersona.secuencia_inftecnico()
                    if not (
                    HistorialPagoMes.objects.values('id').filter(status=True, contrato_id=int(request.GET['id']),
                                                                 fecha_pago__month=int(request.GET['desde']),
                                                                 fecha_pago__year=int(request.GET['anio'])).exists()):
                        historial = HistorialPagoMes(
                            contrato_id=int(request.GET['id']),
                            fecha_pago=date(year=int(request.GET['anio']), month=int(request.GET['desde']),
                                            day=date.today().day),
                            cancelado=False
                        )
                        historial.save(request)
                        log('Agrega historial de pago %s' % (historial), request, 'add')
                    else:
                        historial = HistorialPagoMes.objects.get(status=True, contrato_id=int(request.GET['id']),
                                                                 fecha_pago__month=int(request.GET['desde']),
                                                                 fecha_pago__year=int(request.GET['anio']))
                    name_file = 'Acta_Pago_' + str(secinforme) + random.randint(1, 10000).__str__() + '.pdf'
                    folder = 'contratoepunemi/inftec/'
                    try:
                        os.stat(os.path.join(SITE_STORAGE, 'media', folder))
                    except:
                        os.mkdir(os.path.join(SITE_STORAGE, 'media', folder))
                    secpos = \
                    SecuenciaMemoActividadPosgrado.objects.filter(tipo=PlantillaContratoDip.objects.last(), status=True,
                                                                  anioejercicio__anioejercicio=datetime.now().year)[0]
                    informe = ActaPago(mes=int(desde), contrato=contratopersona, secuenciageneral=secpos,
                                       secuencia=secinforme, historialpago=historial)
                    # data = extraervalores(contratopersona.persona.id,request.GET['desde'],request.GET['anio'],True if 'h' in request.GET else False)
                    data['inftecnicod'] = contratopersona.informetecnico_set.filter(status=True, mes=int(desde))[0]
                    data['actividadjornada'] = \
                    contratopersona.informeactividadjornada_set.filter(status=True, mes=int(desde))[0]
                    data['fehca'] = date.today()
                    data['mes'] = MESES_CHOICES[datetime.today().month - 1]
                    data['secuencia'] = '0' + str(informe) if informe.secuencia > 9 else '00' + str(informe)
                    pdf = conviert_html_to_pdf_save_informe('adm_contratodip/acta_pdf.html',
                                                            {'pagesize': 'A4', 'data': data}, folder,
                                                            name_file)
                    if pdf:
                        informe.archivo.name = os.path.join(folder, name_file)
                        informe.save(request)
                        mensaje = {"result": True, "mensaje": 'Acta generado correctamente'}
                    else:
                        mensaje = {"result": False, "mensaje": 'Error al guardar pdf'}
                    return JsonResponse(mensaje)
                except Exception as ex:
                    print(ex)
                    transaction.set_rollback(True)
                    return JsonResponse({"result": False, "mensaje": "Error al procesar los datos."})

            elif action == 'gestionarcontrato':
                try:
                    data['title'] = u'Nuevo Contrato Persona'
                    form = GestionarContratoDipForm()
                    # form.fields['persona'].queryset = Persona.objects.none()
                    data['form'] = form
                    return render(request, "adm_contratodip/gestionarcontrato.html", data)
                except Exception as ex:
                    pass

            elif action == 'addcontrato':
                try:
                    data['title'] = u'Nuevo Contrato Persona'
                    form = ContratoDipForm()
                    form.del_fields('validadorgp')
                    data['form'] = form
                    return render(request, "adm_contratodip/addcontrato.html", data)
                except Exception as ex:
                    pass

            elif action == 'editcontratopos':
                try:
                    data['title'] = 'Editar Contrato persona'
                    data['filtro'] = filtro = ContratoDip.objects.get(status=True, pk=int(request.GET['id']))
                    data['responsable'] = str(filtro.gestion.responsable if filtro.gestion else '')
                    data['responsablesub'] = str(filtro.gestion.responsablesubrogante if filtro.gestion else '')
                    data['codigo'] = str(filtro.certificacion.codigo if filtro.certificacion else '')
                    # form = ContratoDipForm(initial=model_to_dict(filtro))
                    form = ContratoDipForm(initial={
                        'codigocontrato': filtro.codigocontrato,
                        'seccion': filtro.gestion,
                        'persona': filtro.persona,
                        'cargo': filtro.cargo,
                        'tipogrupo': filtro.tipogrupo,
                        'tipopago': filtro.tipopago,
                        'plantilla': filtro.plantilla,
                        'certificacion': filtro.certificacion,
                        'fechainicio': filtro.fechainicio,
                        'fechafin': filtro.fechafin,
                        'rmu': filtro.rmu,
                        'ivaAplicado': filtro.iva,
                        'valoriva': filtro.valoriva,
                        'valorTotal': filtro.valortotal,
                        'descripcion': filtro.descripcion,
                        'archivo': filtro.archivo,
                        'validadorgp': filtro.validadorgp.id if filtro.validadorgp else '',
                        'nombreareaprograma': filtro.contratoareaprograma_set.values_list('gestion_id',
                                                                                          flat=True).filter(
                            status=True),
                        'carrera': filtro.contratocarrera_set.values_list('carrera_id', flat=True).filter(status=True),
                    })
                    data['form'] = form
                    return render(request, "adm_contratodip/editcontratopos.html", data)
                except Exception as ex:
                    print(ex)
                    print(sys.exc_info()[-1].tb_lineno)

            elif action == 'editcontrato':
                try:
                    data['title'] = u'Editar Contrato'
                    data['filtro'] = filtro = ContratoDip.objects.get(pk=int(request.GET['id']))
                    data['totalCuotas'] = ContratoDipMetodoPago.objects.filter(status=True, contratodip=int(
                        request.GET['id'])).count()
                    form = ContratoDipForm(initial=model_to_dict(filtro))
                    form.del_fields('carrera')
                    data['form'] = form
                    return render(request, "adm_contratodip/editcontrato.html", data)
                except Exception as ex:
                    pass

            elif action == 'buscarpersonas':
                try:
                    id = request.GET['id']
                    q = request.GET['q'].upper().strip()
                    s = q.split(" ")
                    # querybase = Persona.objects.filter(status=True)
                    querybase = Persona.objects.filter(status=True).order_by(
                        'apellido1')
                    if len(s) == 1:
                        per = querybase.filter((Q(nombres__icontains=q) | Q(apellido1__icontains=q) | Q(
                            cedula__icontains=q) | Q(apellido2__icontains=q) | Q(cedula__contains=q)),
                                               Q(status=True)).distinct()[:15]
                    elif len(s) == 2:
                        per = querybase.filter((Q(apellido1__contains=s[0]) & Q(apellido2__contains=s[1])) |
                                               (Q(nombres__icontains=s[0]) & Q(nombres__icontains=s[1])) |
                                               (Q(nombres__icontains=s[0]) & Q(apellido1__contains=s[1]))).filter(
                            status=True).distinct()[:15]
                    else:
                        per = querybase.filter(
                            (Q(nombres__contains=s[0]) & Q(apellido1__contains=s[1]) & Q(apellido2__contains=s[2])) |
                            (Q(nombres__contains=s[0]) & Q(nombres__contains=s[1]) & Q(
                                apellido1__contains=s[2]))).filter(status=True).distinct()[:15]
                    data = {"result": "ok",
                            "results": [{"id": x.id, "name": "{} - {}".format(x.cedula, x.nombre_completo())} for x in
                                        per]}
                    return JsonResponse(data)
                except Exception as ex:
                    pass

            elif action == 'recurso':
                try:
                    data['title'] = u'Configuraciones'
                    eRecursoPresupuestarioPosgrado = RecursoPresupuestarioPosgrado.objects.filter(status=True).order_by(
                        '-id')
                    paging = MiPaginador(eRecursoPresupuestarioPosgrado, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['listado'] = page.object_list
                    return render(request, "adm_contratodip/configuraciones/recursopresupuestario.html", data)
                except Exception as ex:
                    pass

            elif action == 'conf_recurso_presupuestario':
                try:
                    data['title'] = u'Recurso Presupuestario'
                    eRecursoPresupuestarioPosgrado = RecursoPresupuestarioPosgrado.objects.filter(status=True).order_by(
                        '-id')
                    paging = MiPaginador(eRecursoPresupuestarioPosgrado, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data['listado'] = page.object_list
                    return render(request, "adm_contratodip/configuraciones/conf_recurso_presupuestario.html", data)
                except Exception as ex:
                    pass

            elif action == 'gestionarrecursos':
                try:
                    data['title'] = u'Gestionar Recurso Presupuestario'
                    pk = int(request.GET.get('id', '0'))
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    eRecursoPresupuestarioPosgrado = RecursoPresupuestarioPosgrado.objects.get(pk=pk)
                    data['eRecursoPresupuestarioPosgrado'] = eRecursoPresupuestarioPosgrado
                    return render(request, "adm_contratodip/configuraciones/gestionarrecursos.html", data)
                except Exception as ex:
                    pass

            elif action == 'configuraciones':
                try:
                    secuencia_memo()
                    data['title'] = u'Configuraciones'
                    data['campos'] = CampoContratoDip.objects.filter(status=True).order_by('-id')
                    data['recursopresupuestario'] = RecursoPresupuestarioPosgrado.objects.filter(status=True).order_by(
                        '-id')
                    data['documentos'] = DocumentoContrato.objects.filter(status=True).order_by('-id')
                    data['perfiles'] = PerfilPuestoDip.objects.filter(status=True).order_by('-id')
                    data['actividades'] = ActividadesPerfil.objects.filter(status=True).order_by('-id')
                    data['modeloscontratos'] = PlantillaContratoDip.objects.filter(status=True).order_by('-anio', '-pk')
                    data['secuencia'] = SecuenciaMemoActividadPosgrado.objects.filter(status=True).order_by('-id')
                    search, filtro, url_vars = request.GET.get('s', ''), Q(status=True), ''
                    if search:
                        filtro = filtro & Q(descripcion__icontains=search) | Q(codigo__icontains=search) | Q(
                            partida__icontains=search)
                        url_vars += '&s=' + search
                        data['search'] = search
                    listado = CertificacionPresupuestariaDip.objects.filter(status=True).filter(filtro).order_by('-id')
                    paging = MiPaginador(listado, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['page'] = page
                    data["url_vars"] = url_vars
                    data['listado'] = page.object_list
                    data['totcount'] = listado.count()
                    data['email_domain'] = EMAIL_DOMAIN
                    return render(request, "adm_contratodip/configuraciones/configuraciones.html", data)
                except Exception as ex:
                    pass

            elif action == 'addrecursopresupuesatario':
                try:
                    f = RecursoPresupuestarioForm()
                    data['form2'] = f
                    template = get_template('adm_contratodip/configuraciones/modal/formModal.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})

                except Exception as ex:
                    pass

            elif action == 'editrecursopresupuesatario':
                try:
                    pk = request.GET.get('id', '0')
                    eRecursoPresupuestarioPosgrado = RecursoPresupuestarioPosgrado.objects.get(pk=pk)
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    f = RecursoPresupuestarioForm(initial={
                        'descripcion': eRecursoPresupuestarioPosgrado.descripcion
                    })
                    data['form2'] = f
                    data['id'] = pk
                    template = get_template('adm_contratodip/configuraciones/modal/formModal.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})

                except Exception as ex:
                    pass

            elif action == 'add_programa_maestria':
                try:
                    f = CabeceraRecursoPresupuestarioPosgradoForm()
                    data['form2'] = f
                    pk = int(request.GET.get('id', '0'))
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    data['id_secundario'] = pk
                    template = get_template('adm_contratodip/configuraciones/modal/formModal.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})

                except Exception as ex:
                    pass

            elif action == 'edit_programa_maestria':
                try:
                    pk = int(request.GET.get('id', '0'))
                    eCabeceraRecursoPresupuestarioPosgrado = CabeceraRecursoPresupuestarioPosgrado.objects.get(pk=pk)
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    f = CabeceraRecursoPresupuestarioPosgradoForm(initial={
                        'malla': eCabeceraRecursoPresupuestarioPosgrado.malla,
                        'periodo': eCabeceraRecursoPresupuestarioPosgrado.periodo
                    })
                    data['form2'] = f
                    data['id'] = eCabeceraRecursoPresupuestarioPosgrado.pk
                    template = get_template('adm_contratodip/configuraciones/modal/formModal.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})

                except Exception as ex:
                    pass

            elif action == 'add_subitems':
                try:
                    f = DetalleRecursoPresupuestarioPosgradoForm()
                    data['form2'] = f
                    pk = int(request.GET.get('id', '0'))
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    data['id_secundario'] = pk
                    template = get_template('adm_contratodip/configuraciones/modal/formModal.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})

                except Exception as ex:
                    pass

            elif action == 'add_items':
                try:
                    f = ItemRecursoPresupuestarioPosgradoForm()
                    data['form2'] = f
                    pk = int(request.GET.get('id', '0'))
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    data['id_secundario'] = pk
                    template = get_template('adm_contratodip/configuraciones/modal/formModal.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})

                except Exception as ex:
                    pass

            elif action == 'edit_items':
                try:
                    pk = int(request.GET.get('id', '0'))
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    eItemRecursoPresupuestarioPosgrado = ItemRecursoPresupuestarioPosgrado.objects.get(pk=pk)
                    f = ItemRecursoPresupuestarioPosgradoForm(initial={
                        'total_paralelos': eItemRecursoPresupuestarioPosgrado.total_paralelos,
                        'modulos_a_dictar': eItemRecursoPresupuestarioPosgrado.modulos_a_dictar,
                    })
                    data['form2'] = f

                    data['id'] = eItemRecursoPresupuestarioPosgrado.pk
                    template = get_template('adm_contratodip/configuraciones/modal/formModal.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})

                except Exception as ex:
                    pass

            elif action == 'edit_subitems':
                try:
                    pk = int(request.GET.get('id', '0'))
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    eDetalleRecursoPresupuestarioPosgrado = DetalleRecursoPresupuestarioPosgrado.objects.get(pk=pk)
                    f = DetalleRecursoPresupuestarioPosgradoForm(initial={
                        'desglosemoduloadictar': eDetalleRecursoPresupuestarioPosgrado.desglosemoduloadictar,
                        'horaspormodulo': eDetalleRecursoPresupuestarioPosgrado.horaspormodulo,
                        'valor_x_hora': eDetalleRecursoPresupuestarioPosgrado.valor_x_hora,
                        'categoriadocente': eDetalleRecursoPresupuestarioPosgrado.categoriadocente
                    })
                    data['form2'] = f

                    data['id'] = eDetalleRecursoPresupuestarioPosgrado.pk
                    template = get_template('adm_contratodip/configuraciones/modal/formModal.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})

                except Exception as ex:
                    pass

            elif action == 'addcampocontrato':
                try:
                    data['form2'] = CampoContratoDipForm()
                    template = get_template("adm_contratodip/modal/formcampo.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'addcertificacion':
                try:
                    data['form2'] = CertificacionPresupuestariaDipForm()
                    data['VALIDAR_CERTIFICACION_PRESUPUESTARIA_POS'] = variable_valor(
                        'VALIDAR_CERTIFICACION_PRESUPUESTARIA_POS')
                    template = get_template("adm_certificacion/modal/formcertificacion.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'editcertificacion':
                try:
                    data['filtro'] = filtro = CertificacionPresupuestariaDip.objects.get(pk=int(request.GET['id']))
                    data['form2'] = CertificacionPresupuestariaDipForm(initial=model_to_dict(filtro))
                    data['VALIDAR_CERTIFICACION_PRESUPUESTARIA_POS'] = variable_valor(
                        'VALIDAR_CERTIFICACION_PRESUPUESTARIA_POS')
                    template = get_template("adm_certificacion/modal/formcertificacion.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'detallecertificacion':
                data['detalle'] = detalle = CertificacionPresupuestariaDip.objects.get(
                    pk=int(request.GET['id'])).detalles_certificacion()
                template = get_template("adm_certificacion/modal/detallecertificacion.html")
                return JsonResponse({"result": True, 'data': template.render(data)})

            elif action == 'buscarpersonas':
                try:
                    id = request.GET['id']
                    q = request.GET['q'].upper().strip()
                    s = q.split(" ")
                    querybase = Persona.objects.filter(status=True)
                    if len(s) == 1:
                        per = querybase.filter((Q(nombres__icontains=q) | Q(apellido1__icontains=q) | Q(
                            cedula__icontains=q) | Q(apellido2__icontains=q) | Q(cedula__contains=q)),
                                               Q(status=True)).distinct()[:15]
                    elif len(s) == 2:
                        per = querybase.filter((Q(apellido1__contains=s[0]) & Q(apellido2__contains=s[1])) |
                                               (Q(nombres__icontains=s[0]) & Q(nombres__icontains=s[1])) |
                                               (Q(nombres__icontains=s[0]) & Q(apellido1__contains=s[1]))).filter(
                            status=True).distinct()[:15]
                    else:
                        per = querybase.filter(
                            (Q(nombres__contains=s[0]) & Q(apellido1__contains=s[1]) & Q(apellido2__contains=s[2])) |
                            (Q(nombres__contains=s[0]) & Q(nombres__contains=s[1]) & Q(
                                apellido1__contains=s[2]))).filter(status=True).distinct()[:15]
                    data = {"result": "ok",
                            "results": [{"id": x.id, "name": "{} - {}".format(x.cedula, x.nombre_completo())} for x in
                                        per]}
                    return JsonResponse(data)
                except Exception as ex:
                    pass

            elif action == 'descargaplantilla':
                try:
                    plantilla = PlantillaContratoDip.objects.get(pk=request.GET['id'])
                    campos = plantilla.campoplantillacontratodip_set.filter(status=True)
                    __author__ = 'Unemi'

                    output = io.BytesIO()
                    workbook = xlsxwriter.Workbook(output)
                    ws = workbook.add_worksheet('plantilla_')
                    ws.set_column(0, 100, 60)

                    formatoceldagris = workbook.add_format(
                        {'align': 'center', 'border': 1, 'text_wrap': True, 'fg_color': '#B6BFC0'})

                    ws.write(0, 0, 'CAMPO', formatoceldagris)
                    ws.write(0, 1, 'VALOR', formatoceldagris)
                    ws.write(0, 2, 'UNIDAD ORGÁNICA', formatoceldagris)
                    ws.write(0, 3, 'DEPARTAMENTO', formatoceldagris)
                    ws.write(0, 4, 'DENOMINACION PUESTO', formatoceldagris)
                    ws.write(0, 5, 'PUESTO', formatoceldagris)
                    ws.write(0, 6, 'REMUNERACIÓN', formatoceldagris)
                    ws.write(0, 7, 'RMU', formatoceldagris)
                    cont = 1
                    for campo in campos:
                        ws.write(cont, 0, str(campo.campos), formatoceldagris)
                        ws.write(cont, 1,
                                 str(campo.campos.identificador if campo.campos.identificador else 'INDEFINIDO'),
                                 formatoceldagris)
                        cont += 1
                    workbook.close()
                    output.seek(0)
                    filename = 'plantilla_%s.xlsx' % (plantilla.descripcion)
                    response = HttpResponse(output,
                                            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    response['Content-Disposition'] = 'attachment; filename=%s' % filename
                    return response
                except Exception as ex:
                    pass

            elif action == 'metodoPago':
                try:
                    existeCuotasContrato = ContratoDipMetodoPago.objects.filter(contratodip=request.GET['id']).exists()

                    if existeCuotasContrato:
                        data['cuotasContrato'] = ContratoDipMetodoPago.objects.filter(contratodip=request.GET['id'],
                                                                                      status=True)

                    data['personaContrato'] = ContratoDip.objects.get(id=request.GET['id'])
                    template = get_template("adm_contratodip/modal/formmetodopago.html")
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'detalle_cuota':
                try:
                    if 'id' in request.GET:
                        data['cuotasContrato'] = ContratoDipMetodoPago.objects.filter(contratodip=request.GET['id'],
                                                                                      status=True)
                        template = get_template("adm_contratodip/detallecuotapago.html")
                        return JsonResponse({"result": 'ok', 'data': template.render(data)})
                except Exception as ex:
                    transaction.set_rollback(True)
                    return JsonResponse({"result": "bad", "mensaje": u"Error al obtener los datos."})

            elif action == 'detalle_contrato':
                try:
                    if 'id' in request.GET:
                        data['contrato'] = ContratoDip.objects.get(id=request.GET['id'],
                                                                   status=True)
                        template = get_template("adm_contratodip/detallecontrato.html")
                        return JsonResponse({"result": 'ok', 'data': template.render(data)})
                except Exception as ex:
                    transaction.set_rollback(True)
                    return JsonResponse({"result": "bad", "mensaje": u"Error al obtener los datos."})

            elif action == 'editcontratodetalle':
                try:
                    data['title'] = u'Modificar detalle del contrato'
                    data['contratopersona'] = contratopersona = \
                        ContratoDip.objects.filter(pk=request.GET['id'], status=True)[0]
                    data['contratoscampos'] = CampoPlantillaContratoDip.objects.filter(
                        contrato=contratopersona.plantilla)
                    data['campos'] = CampoPlantillaContratoDip.objects.filter(status=True,
                                                                              contrato=contratopersona.plantilla).order_by(
                        'campos__id')
                    return render(request, "adm_contratodip/editcontratodetalle.html", data)
                except Exception as ex:
                    pass

            elif action == 'viewmemos':
                try:
                    data['memos'] = memo = MemoActividadPosgrado.objects.filter(status=True,
                                                                                contrato=int(request.GET['idcon']))
                    data['contrato'] = ContratoDip.objects.get(id=int(request.GET['idcon']), status=True)
                    data['title'] = 'Registro de Memorandums'
                    return render(request, "adm_contratodip/viewmemos.html", data)
                except Exception as ex:
                    transaction.set_rollback(True)

            elif action == 'viewbitacora':
                try:
                    data['bitacora'] = bitacora = InformeActividadJornada.objects.filter(status=True, contrato=int(
                        encrypt(request.GET['idcon'])))
                    data['contrato'] = ContratoDip.objects.get(id=encrypt(request.GET['idcon']))
                    data['title'] = 'Registro de Informe de Actividades'
                    return render(request, "adm_contratodip/viewbitacora.html", data)
                except Exception as ex:
                    transaction.set_rollback(True)

            elif action == 'viewinftecnico':
                try:
                    data['contrato'] = contratopersona = ContratoDip.objects.get(id=request.GET['idcon'])
                    data['tecnico'] = InformeTecnico.objects.filter(status=True, contrato=contratopersona)
                    # data['sec'] = secinforme = contratopersona.secuencia_inftecnico()
                    # data['fehca'] = date.today()
                    # data['mes'] =MESES_CHOICES[datetime.today().month-1]
                    data['title'] = 'Registro de Informe Técnico'
                    return render(request, 'adm_contratodip/viewinftecnico.html', data)
                    # return conviert_html_to_pdf('adm_contratodip/informe_tecnico_pdf.html',{'pagesize': 'A4', 'data': data})
                except Exception as ex:
                    transaction.set_rollback(True)

            elif action == 'viewactapago':
                try:
                    data['contrato'] = contratopersona = ContratoDip.objects.get(id=request.GET['idcon'])
                    data['actapago'] = ActaPago.objects.filter(status=True, contrato=contratopersona)
                    # data['sec'] = secinforme = contratopersona.secuencia_actapago()
                    # data['fehca'] = date.today()
                    # data['mes'] =MESES_CHOICES[datetime.today().month-1]
                    data['title'] = 'Registro de Acta de Pago'
                    return render(request, 'adm_contratodip/viewactapago.html', data)
                    # return conviert_html_to_pdf('adm_contratodip/acta_pdf.html',{'pagesize': 'A4', 'data': data})
                except Exception as ex:
                    transaction.set_rollback(True)

            elif action == 'viewreportes':
                try:
                    data['contrato'] = contratopersona = ContratoDip.objects.get(status=True,
                                                                                 id=encrypt(request.GET['idcon']))
                    data['historial'] = HistorialPagoMes.objects.filter(status=True, contrato=contratopersona).order_by(
                        'fecha_pago')
                    data['actapago'] = ActaPago.objects.filter(status=True, contrato=contratopersona)
                    data['tecnico'] = InformeTecnico.objects.filter(status=True, contrato=contratopersona)
                    data['bitacora'] = InformeActividadJornada.objects.filter(status=True, contrato=contratopersona)
                    data['memos'] = MemoActividadPosgrado.objects.filter(status=True, contrato=contratopersona)
                    data['title'] = 'Registro de Reportes'
                    return render(request, 'adm_contratodip/viewreportes.html', data)
                except Exception as ex:
                    pass

            elif action == 'consultaarchivos':
                try:
                    data['contrato'] = contratopersona = ContratoDip.objects.get(status=True,
                                                                                 id=encrypt(request.GET['idcon']))
                    data['historial'] = historial = HistorialPagoMes.objects.get(status=True, contrato=contratopersona,
                                                                                 id=encrypt(request.GET['idhist']))
                    data['actapago'] = ActaPago.objects.filter(status=True, contrato=contratopersona,
                                                               fecha_creacion__year=historial.fecha_pago.year,
                                                               mes=historial.fecha_pago.month).first()
                    data['tecnico'] = InformeTecnico.objects.filter(status=True, contrato=contratopersona,
                                                                    fecha_creacion__year=historial.fecha_pago.year,
                                                                    mes=historial.fecha_pago.month).first()
                    data['bitacora'] = InformeActividadJornada.objects.filter(status=True, contrato=contratopersona,
                                                                              fecha_creacion__year=historial.fecha_pago.year,
                                                                              mes=historial.fecha_pago.month).first()
                    data['memos'] = MemoActividadPosgrado.objects.filter(status=True, contrato=contratopersona,
                                                                         fecha_creacion__year=historial.fecha_pago.year,
                                                                         mes=historial.fecha_pago.month).first()
                    if ContratoDip.objects.filter(status=True, id=encrypt(request.GET['idcon']),
                                                  persona=persona).exists():
                        data['trabaja'] = True
                    template = get_template("adm_contratodip/archivospagos.html")
                    return JsonResponse({"result": 'ok', 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'addarchivoevidencias':
                try:
                    if int(request.GET['tipo']) == 1:
                        arch = MemoActividadPosgrado.objects.get(status=True, pk=int(encrypt(request.GET['id'])))
                    elif int(request.GET['tipo']) == 2:
                        arch = InformeActividadJornada.objects.get(status=True, pk=int(encrypt(request.GET['id'])))
                    elif int(request.GET['tipo']) == 3:
                        arch = InformeTecnico.objects.get(status=True, pk=int(encrypt(request.GET['id'])))
                    else:
                        arch = ActaPago.objects.get(status=True, pk=int(encrypt(request.GET['id'])))
                    if arch.archivofirmado:
                        data['filtro'] = arch
                        data['id'] = encrypt(request.GET['id'])
                        data['tipo'] = request.GET['tipo']
                        data['form2'] = ArchivoInformesForm(initial={'archivo': arch.archivofirmado})
                    else:
                        form2 = ArchivoInformesForm()
                        data['tipo'] = request.GET['tipo']
                        data['id'] = encrypt(request.GET['id'])
                        data['form2'] = form2
                    template = get_template('adm_contratodip/modal/addarchivosinfo.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'reportescontratos':
                try:
                    contratos = ContratoDip.objects.filter(status=True).all()
                    __autor__ = 'unemi'
                    output = io.BytesIO()
                    workbook = xlsxwriter.Workbook(output)
                    ws = workbook.add_worksheet('contratos')
                    titulo = workbook.add_format(
                        {'font_name': 'Calibri', 'border': 1, 'align': 'center', 'font_size': 12, 'valign': 'vcenter',
                         'bg_color': '#E4E5DF'})
                    style2 = workbook.add_format(
                        {'text_wrap': True, 'font_name': 'Calibri', 'border': 1, 'align': 'center', 'font_size': 12,
                         'valign': 'vcenter'})
                    money = workbook.add_format(
                        {'text_wrap': True, 'font_name': 'Calibri', 'border': 1, 'align': 'center', 'font_size': 12,
                         'valign': 'vcenter', 'num_format': '$#,##0.00'})
                    row_num = 1
                    columns = [
                        ('No. Contrato', 15),
                        ('No. Partida', 15),
                        ('Nombres', 30),
                        ('Apellidos', 30),
                        ('Cédula', 20),
                        ('Cargo', 70),
                        ('Departamento', 70),
                        ('RMU', 20),
                        ('Valor IVA', 20),
                        ('Valor Total', 20),
                    ]
                    for col_num in range(len(columns)):
                        ws.write(row_num, col_num, columns[col_num][0], titulo)
                        ws.set_column(col_num, col_num, columns[col_num][1])
                    row_num += 1
                    for data in contratos:
                        ws.write(row_num, 0, data.codigocontrato, style2)
                        ws.write(row_num, 1, data.certificacion.codigo, style2)
                        ws.write(row_num, 2, data.persona.nombres, style2)
                        ws.write(row_num, 3, data.persona.apellido1 + ' ' + data.persona.apellido2, style2)
                        ws.write(row_num, 4, str(data.persona.cedula), style2)
                        ws.write(row_num, 5, data.cargo.__str__(), style2)
                        ws.write(row_num, 6, data.seccion.departamento.__str__(), style2)
                        ws.write(row_num, 7, data.rmu, money)
                        ws.write(row_num, 8, data.valoriva, money)
                        ws.write(row_num, 9, data.valortotal, money)
                        row_num += 1
                    workbook.close()
                    output.seek(0)
                    filename = 'contratados' + random.randint(1, 10000).__str__() + '.xlsx'
                    response = HttpResponse(output,
                                            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    response['Content-Disposition'] = 'attachment; filename=%s' % filename
                    return response
                except Exception as ex:
                    pass

            elif action == 'adddocument':
                try:
                    data['title'] = u'Agregar documento'
                    data['form2'] = DocumentoContratosForm()
                    data['add'] = True
                    template = get_template("adm_contratodip/modal/formperfil.html")
                    return JsonResponse({'result': True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({'result': False, 'mensaje': ''})

            elif action == 'editdocument':
                try:
                    data['title'] = u'Agregar documento'
                    data['filtro'] = filtro = DocumentoContrato.objects.get(id=request.GET['id'], status=True)
                    data['form2'] = DocumentoContratosForm(initial=model_to_dict(filtro))
                    data['add'] = True
                    template = get_template("adm_contratodip/modal/formperfil.html")
                    return JsonResponse({'result': True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({'result': False, 'mensaje': ''})

            elif action == 'horariocontrato':
                try:
                    data['title'] = u"Gestión de horarios"
                    id = request.GET.get('id', None)
                    eContratoDip = ContratoDip.objects.get(status=True, pk=int(encrypt(id)))
                    data['idp'] = eContratoDip.id
                    data['eContrato'] = eContratoDip
                    data['hoy_str'] = datetime.now().date()
                    return render(request, "adm_contratodip/horariocontrato.html", data)
                except Exception as ex:
                    pass

            elif action == 'cargareventos':
                try:
                    idp = int(encrypt(request.GET['id']))
                    fecha = datetime.now().date()
                    filtro = Q(status=True, contrato_id=idp)
                    value = request.GET.get('value', '')
                    if value:
                        q = request.GET['value'].upper().strip()
                        s = q.split(" ")
                        if len(s) == 1:
                            filtro = filtro & (Q(persona__nombres__icontains=q) |
                                               Q(persona__apellido1__icontains=q) |
                                               Q(persona__cedula__icontains=q) |
                                               Q(persona__apellido2__icontains=q) |
                                               Q(persona__cedula__contains=q))
                        elif len(s) == 2:
                            filtro = filtro & (
                                        (Q(persona__apellido1__contains=s[0]) & Q(persona__apellido2__contains=s[1])) |
                                        (Q(persona__nombres__icontains=s[0]) & Q(persona__nombres__icontains=s[1])) |
                                        (Q(persona__nombres__icontains=s[0]) & Q(persona__apellido1__contains=s[1])))
                        else:
                            filtro = filtro & ((Q(persona__nombres__contains=s[0]) & Q(
                                persona__apellido1__contains=s[1]) & Q(persona__apellido2__contains=s[2])) |
                                               (Q(persona__nombres__contains=s[0]) & Q(
                                                   persona__nombres__contains=s[1]) & Q(
                                                   persona__apellido1__contains=s[2])))
                    else:
                        finicio = datetime.strptime(request.GET['finicio'], '%d/%m/%Y')
                        ffin = datetime.strptime(request.GET['ffin'], '%d/%m/%Y')
                        filtro = filtro & Q(Q(inicio__range=(finicio, ffin)) | Q(fin__range=(finicio, ffin)))
                    cronogramas = HorarioPlanificacionContrato.objects.filter(filtro)
                    if cronogramas and value:
                        fecha = cronogramas.order_by('inicio').first().fecha
                    event_list = []
                    for cronograma in cronogramas:
                        start_date = cronograma.inicio.strftime('%Y-%m-%dT%H:%M:%S')
                        end_date = cronograma.fin.strftime('%Y-%m-%dT%H:%M:%S')
                        event_list.append({
                            'id': cronograma.id,
                            'title': cronograma.contrato.persona.nombre_normal_minus(),
                            'extendedProps': {
                                'description': [i.__str__() for i in cronograma.turno.all()],
                                'dia': cronograma.get_dia_display(),
                                'id_estado': '',
                                'puede_eliminar': True,
                                'color_estado': '',
                                'fecha_c': str(cronograma.inicio) if cronograma.inicio else '',
                                'fecha_d': str(cronograma.fin) if cronograma.fin else '',
                                'eventColor': cronograma.get_color_dia()
                            },
                            'start': start_date,
                            'end': end_date,
                            'backgroundColor': cronograma.get_color_dia(),
                        })
                    return JsonResponse({"fecha_e": str(fecha), 'eventos': event_list}, safe=False)
                except Exception as ex:
                    return JsonResponse({'result': False, 'mensaje': '{}'.format(ex)})

            elif action == 'addhorariocontrato':
                try:
                    data['id'] = id = request.GET.get('id')
                    fecha = datetime.strptime(request.GET['dia'], '%Y-%m-%d').date() if request.GET.get('dia') else None
                    model = HorarioPlanificacionContrato.objects.filter(contrato_id=id, status=True).first()
                    f = HorarioClasesContratoForm(initial={'dia': fecha.weekday() + 1, 'inicio': fecha, 'fin': fecha})
                    data['form2'] = f
                    template = get_template('adm_contratodip/modal/formhorario.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    err_ = f'{ex}({sys.exc_info()[-1].tb_lineno})'
                    return JsonResponse({"result": False, 'mensaje': err_})

            elif action == 'edithorariocontrato':
                try:
                    data['id'] = id = request.GET.get('id')
                    model = HorarioPlanificacionContrato.objects.get(id=id, status=True)
                    f = HorarioClasesContratoForm(initial=model_to_dict(model))
                    data['form2'] = f
                    template = get_template('adm_contratodip/modal/formhorario.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    err_ = f'{ex}({sys.exc_info()[-1].tb_lineno})'
                    return JsonResponse({"result": False, 'mensaje': err_})

            elif action == 'updatebitacora':
                try:
                    data['title'] = u'Agregar documento'
                    data['id'] = encrypt(request.GET['id'])
                    form = FechaMaximoBitacoraForm()
                    template = get_template("adm_contratodip/modal/formmodal.html")
                    data['form'] = form
                    res_js = {'result': True, 'data': template.render(data)}
                except Exception as ex:
                    err_ = f"Ocurrio un error: {ex.__str__()}. En la linea {sys.exc_info()[-1].tb_lineno}"
                    res_js = {'result': False, 'mensaje': err_}
                return JsonResponse(res_js)

            elif action == 'horarios':
                try:
                    data['title'] = 'Horarios de actividades'
                    data['turnos'] = turnos = Turno.objects.filter(status=True, mostrar=True, sesion_id=20)
                    data['semana'] = [[1, 'Lunes'], [2, 'Martes'], [3, 'Miercoles'], [4, 'Jueves'], [5, 'Viernes'],
                                      [6, 'Sabado'], [6, 'Domingo']]
                    hoy = datetime.now().date()
                    data['sumaactividad'] = 0
                    data['suma'] = 0
                    return render(request, 'adm_contratodip/horarios.html', data)
                except Exception as ex:
                    err_ = f'Lo sentimos, tuvimos fallas al ingresara la configuración de horario: {ex.__str__()}({sys.exc_info()[-1].tb_lineno})'
                    return HttpResponseRedirect(f'{request.path}?info={err_}')

            elif action == 'gruporevision':
                try:
                    data['title'] = 'Grupo de revisión'
                    filtro, url_vars, search = Q(status=True), '&action=gruporevision', request.GET.get('s', None)
                    ePersona_id = ContratoDip.objects.filter(status=True, validadorgp__isnull=False).values_list(
                        'validadorgp_id', flat=True).distinct()
                    if search:
                        data['search'] = search
                        ss = search.split(' ')
                        if len(ss) == 1:
                            filtro = filtro & Q(Q(nombre__icontains=ss) |
                                                Q(apellido1__icontains=search) |
                                                Q(apellido1__icontains=search) |
                                                Q(cedula__icontains=search) |
                                                Q(email__icontains=search) |
                                                Q(emailinst__icontains=search))

                    query = ContratoDip.objects.filter(filtro).filter(persona_id__in=ePersona_id)
                    paging = MiPaginador(query, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['page'] = page
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['url_vars'] = url_vars
                    data['listado'] = paging.object_list
                    return render(request, 'adm_contratodip/viewgruporevision.html', data)
                except Exception as ex:
                    err_ = f"{ex}({sys.exc_info()[-1].tb_lineno})"
                    return HttpResponseRedirect(f'{request.path}?info={err_}')

            elif action == 'configuracionrequisitocontratacion':
                try:
                    data['title'] = 'Configuración requisito contratación'
                    filtro, url_vars, search = Q(
                        status=True), '&action=configuracionrequisitocontratacion', request.GET.get('s', None)
                    query = ContratacionConfiguracionRequisito.objects.filter(filtro).distinct()
                    paging = MiPaginador(query, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['page'] = page
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['url_vars'] = url_vars
                    data['ContratacionConfiguracionRequisito'] = paging.object_list
                    return render(request, 'adm_contratodip/viewconfiguracionrequisitocontratacion.html', data)
                except Exception as ex:
                    err_ = f"{ex}({sys.exc_info()[-1].tb_lineno})"
                    return HttpResponseRedirect(f'{request.path}?info={err_}')

            elif action == 'configuracionrequisitopago':
                try:
                    data['title'] = 'Configuración requisito pago'
                    filtro, url_vars, search = Q(status=True), '&action=configuracionrequisitopago', request.GET.get(
                        's', None)
                    query = GrupoRequisitoPago.objects.filter(filtro).distinct()
                    paging = MiPaginador(query, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['page'] = page
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['url_vars'] = url_vars
                    data['GrupoRequisitoPago'] = paging.object_list
                    return render(request, 'adm_contratodip/pago/viewconfiguracionrequisitopago.html', data)
                except Exception as ex:
                    err_ = f"{ex}({sys.exc_info()[-1].tb_lineno})"
                    return HttpResponseRedirect(f'{request.path}?info={err_}')

            elif action == 'configuracionactapago':
                try:
                    data['title'] = 'Configuración acta pago'
                    request.session['confgenerales'] = 1
                    pk = int(request.GET.get('id','0'))
                    if pk== 0:
                        raise NameError("Parametro no encontrado")
                    eConfiguracionActaPago = ConfiguracionActaPago.objects.get(pk=pk)

                    eConfiguracionActaPagoSolicitadoPorForm = ConfiguracionActaPagoSolicitadoPorForm()
                    eConfiguracionActaPagoObjetivoForm = ConfiguracionActaPagoObjetivoForm()
                    eConfiguracionActaPagoMarcoJuridicoReferencialForm = ConfiguracionActaPagoMarcoJuridicoReferencialForm()
                    eConfiguracionActaPagoTituloForm = ConfiguracionActaPagoTituloForm()
                    eConfiguracionActaPagoParaForm = ConfiguracionActaPagoParaForm()
                    eConfiguracionActaPagoConclusionesForm = ConfiguracionActaPagoConclusionesForm()
                    eConfiguracionActaPagoRecomendacionesForm = ConfiguracionActaPagoRecomendacionesForm()

                    eConfiguracionActaPagoSolicitadoPorForm.edit(eConfiguracionActaPago.solicitadopor_id)
                    eConfiguracionActaPagoParaForm.edit(eConfiguracionActaPago.para_id)
                    eConfiguracionActaPagoTituloForm.fields['titulo'].initial = eConfiguracionActaPago.titulo
                    eConfiguracionActaPagoObjetivoForm.fields['objetivo'].initial = eConfiguracionActaPago.objetivo
                    eConfiguracionActaPagoMarcoJuridicoReferencialForm.fields['marcojuridicoreferencial'].initial = eConfiguracionActaPago.marcojuridicoreferencial
                    eConfiguracionActaPagoConclusionesForm.fields['conclusiones'].initial = eConfiguracionActaPago.conclusiones
                    eConfiguracionActaPagoRecomendacionesForm.fields['recomendaciones'].initial = eConfiguracionActaPago.recomendaciones


                    data['form_solicitado_por'] = eConfiguracionActaPagoSolicitadoPorForm
                    data['form_objetivo'] = eConfiguracionActaPagoObjetivoForm
                    data['form_marco_juridico_referencial'] = eConfiguracionActaPagoMarcoJuridicoReferencialForm
                    data['form_titulo'] = eConfiguracionActaPagoTituloForm
                    data['form_para'] = eConfiguracionActaPagoParaForm
                    data['configuracion'] = eConfiguracionActaPago
                    data['form_conclusion'] = eConfiguracionActaPagoConclusionesForm
                    data['form_recomendacion'] = eConfiguracionActaPagoRecomendacionesForm
                    return render(request, 'adm_contratodip/confgenerales/conf_acta_pago.html', data)
                except Exception as ex:
                    err_ = f"{ex}({sys.exc_info()[-1].tb_lineno})"
                    return HttpResponseRedirect(f'{request.path}?info={err_}')


            elif action == 'listadoformatoactapago':
                try:
                    data['title'] = 'Listado de formato de acta pago'
                    request.session['confgenerales'] = 1
                    eConfiguracionActaPago = ConfiguracionActaPago.objects.filter(status=True)

                    data['eConfiguracionActaPago'] = eConfiguracionActaPago
                    return render(request, 'adm_contratodip/confgenerales/formatosactapago.html', data)
                except Exception as ex:
                    err_ = f"{ex}({sys.exc_info()[-1].tb_lineno})"
                    return HttpResponseRedirect(f'{request.path}?info={err_}')

            elif action == 'orden_firma_acta_pago':
                try:
                    request.session['confgenerales'] = 2
                    eOrdenFirmaActaPago = OrdenFirmaActaPago.objects.filter(status=True).order_by('orden')
                    data['eOrdenFirmaActaPago'] = eOrdenFirmaActaPago
                    return render(request, "adm_contratodip/confgenerales/orden_firma.html", data)
                except Exception as ex:
                    pass

            elif action == 'add_orden_firma_acta_pago':
                try:
                    f = OrdenFirmaInformeActaPagoForm()
                    data['form2'] = f
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'add_formato_acta_pago':
                try:
                    f = ConfiguracionActaPagoForm()
                    data['form2'] = f
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'addgrouprevision':
                try:
                    data['form2'] = GrupoRevisionPagoContratoForm()
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "mensaje": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'addconfiguracionrequisitocontratacion':
                try:
                    data['form2'] = ContratacionConfiguracionRequisitoForm()
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "mensaje": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'addconfiguracionrequisitopago':
                try:
                    data['form2'] = ConfiguracionGrupoPagoForm()
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "mensaje": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'editgrouprevision':
                try:
                    data['id'] = id = request.GET.get('id', None)
                    id = encrypt(id)
                    filtro = GrupoRevisionPago.objects.get(status=True, id=id)
                    form = GrupoRevisionPagoContratoForm(
                        initial=model_to_dict(filtro)
                    )
                    ids_sub = filtro.gruporevisionpagocontrato_set.values_list('personacontrato_id', flat=True).filter(
                        status=True)

                    form.fields['personacontrato'].intial = ids_sub
                    form.edit(filtro.persona.id)
                    data['ids_sub'] = ids_sub
                    form.edit_personacontrato(ids_sub)
                    data['form2'] = form
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "message": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'editconfiguracionrequisitocontratacion':
                try:
                    data['id'] = id = request.GET.get('id', None)
                    id = encrypt(id)
                    eContratacionConfiguracionRequisito = ContratacionConfiguracionRequisito.objects.get(status=True,
                                                                                                         id=id)
                    form = ContratacionConfiguracionRequisitoForm(
                        initial=model_to_dict(eContratacionConfiguracionRequisito)
                    )
                    data['form2'] = form
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "message": f"Lo sentimos, {ex.__str__()}"})


            elif action == 'editconfiguracionrequisitopago':
                try:
                    data['id'] = id = request.GET.get('id', None)
                    id = encrypt(id)
                    eGrupoRequisitoPago = GrupoRequisitoPago.objects.get(status=True, id=id)
                    form = ConfiguracionGrupoPagoForm(
                        initial=model_to_dict(eGrupoRequisitoPago)
                    )
                    data['form2'] = form
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "message": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'editgestorrequisitocontratacion':
                try:
                    data['id'] = id = request.GET.get('id', None)
                    id = encrypt(id)
                    eRequisitoContratacionConfiguracionRequisito = RequisitoContratacionConfiguracionRequisito.objects.get(
                        status=True, id=id)
                    form = ContratacionGestionRequisitoForm(
                        initial=model_to_dict(eRequisitoContratacionConfiguracionRequisito)
                    )
                    data['form2'] = form
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "message": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'editgestorrequisitopago':
                try:
                    data['id'] = id = request.GET.get('id', None)
                    id = encrypt(id)
                    eRequisitoPagoGrupoRequisito = RequisitoPagoGrupoRequisito.objects.get(status=True, id=id)
                    form = RequisitoPagoDipForm(
                        initial=model_to_dict(eRequisitoPagoGrupoRequisito)
                    )
                    form.fields['requisito'].initial = eRequisitoPagoGrupoRequisito.requisitopagodip
                    data['form2'] = form
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "message": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'editgestorrequisitopagoorden':
                try:
                    data['id'] = id = request.GET.get('id', None)
                    id = encrypt(id)
                    eRequisitoPagoGrupoRequisito = RequisitoPagoGrupoRequisito.objects.get(status=True, id=id)
                    form = RequisitoPagoDipOrdenForm(
                        initial={
                            'orden': eRequisitoPagoGrupoRequisito.orden,
                            'opcional': eRequisitoPagoGrupoRequisito.opcional
                        }
                    )
                    data['form2'] = form
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "message": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'buscarpersona':
                try:
                    q = request.GET['q'].upper().strip()
                    s = q.split(" ")
                    filtro = Q(usuario__isnull=False, status=True)
                    if len(s) == 1:
                        filtro &= ((Q(nombres__icontains=q) | Q(apellido1__icontains=q) | Q(cedula__icontains=q) | Q(
                            apellido2__icontains=q) | Q(cedula__contains=q)))
                    elif len(s) == 2:
                        filtro &= ((Q(apellido1__contains=s[0]) & Q(apellido2__contains=s[1])) |
                                   (Q(nombres__icontains=s[0]) & Q(nombres__icontains=s[1])) |
                                   (Q(nombres__icontains=s[0]) & Q(apellido1__contains=s[1])))
                    else:
                        filtro &= ((Q(nombres__contains=s[0]) & Q(apellido1__contains=s[1]) & Q(
                            apellido2__contains=s[2])) |
                                   (Q(nombres__contains=s[0]) & Q(nombres__contains=s[1]) & Q(
                                       apellido1__contains=s[2])))

                    per = Persona.objects.filter(filtro).exclude(cedula='').order_by('apellido1', 'apellido2',
                                                                                     'nombres').distinct()[:15]
                    return JsonResponse({"result": "ok", "results": [{"id": x.id, "name": "%s %s" % (
                    f"<img src='{x.get_foto()}' width='25' height='25' style='border-radius: 20%;' alt='...'>",
                    x.nombre_completo_inverso())} for x in per]})
                except Exception as ex:
                    pass

            elif action == 'depposgrado':
                try:
                    data['title'] = 'Departamentos posgrado'
                    filtro, url_vars, search = Q(status=True), '&action=depposgrado', request.GET.get('s', None)

                    if search:
                        data['search'] = search
                        ss = search.split(' ')
                        if len(ss) == 1:
                            filtro = filtro & Q(Q(nombre__icontains=ss) |
                                                Q(responsable__persona__apellido1__icontains=search) |
                                                Q(responsable__persona__apellido1__icontains=search) |
                                                Q(responsable__persona__cedula__icontains=search) |
                                                Q(responsable__persona__email__icontains=search) |
                                                Q(responsable__persona__emailinst__icontains=search))

                    query = Departamento.objects.filter(filtro)
                    paging = MiPaginador(query, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['page'] = page
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['url_vars'] = url_vars
                    data['listado'] = page.object_list
                    return render(request, 'adm_contratodip/viewdepartamentos.html', data)
                except Exception as ex:
                    pass

            elif action == 'adddepposgrado':
                try:
                    data['form2'] = DepartamentoPosgradoForm()
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "mensaje": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'editdepposgrado':
                try:
                    data['id'] = id = request.GET.get('id', None)
                    id = encrypt(id)
                    filtro = Departamento.objects.get(status=True, id=id)
                    form = DepartamentoPosgradoForm(
                        initial=model_to_dict(filtro)
                    )
                    form.edit(filtro.responsable.id)
                    data['ids_sub'] = ids_sub = filtro.responsable_subrogante.values_list('id', flat=True)
                    form.res_subrogante(ids_sub)
                    data['form2'] = form
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "message": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'gestposgrado':
                try:
                    data['title'] = 'Gestiones'
                    data['id'] = id = request.GET.get('id')
                    filtro, url_vars, search = Q(status=True, departamento_id=int(
                        encrypt(id))), f'&action=gestposgrado&id={id}', request.GET.get('s', None)

                    if search:
                        data['search'] = search
                        ss = search.split(' ')
                        if len(ss) == 1:
                            filtro = filtro & Q(Q(gestion__icontains=ss) | Q(cargo__icontains=ss) |
                                                Q(responsable__persona__apellido1__icontains=search) |
                                                Q(responsable__persona__apellido1__icontains=search) |
                                                Q(responsable__persona__cedula__icontains=search) |
                                                Q(responsable__persona__email__icontains=search) |
                                                Q(responsable__persona__emailinst__icontains=search))

                    query = Gestion.objects.filter(filtro)
                    paging = MiPaginador(query, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['url_vars'] = url_vars
                    data['listado'] = page.object_list
                    data['page'] = page
                    return render(request, 'adm_contratodip/viewgestiones.html', data)
                except Exception as ex:
                    pass

            elif action == 'gestionarconfiguracionrequisitocontratacion':
                try:
                    data['title'] = 'gestionar configuracion requisito contratación'
                    id = int(encrypt(request.GET.get('id')))
                    if id == 0:
                        raise NameError("Parametro no encontrado")
                    filtro, url_vars, search = Q(status=True,
                                                 contratacionconfiguracionrequisito_id=id), f'&action=gestionarconfiguracionrequisitocontratacion&id={id}', request.GET.get(
                        's', None)
                    eContratacionConfiguracionRequisito = ContratacionConfiguracionRequisito.objects.get(pk=id)
                    eRequisitoContratacionConfiguracionRequisito = RequisitoContratacionConfiguracionRequisito.objects.filter(
                        filtro)

                    paging = MiPaginador(eRequisitoContratacionConfiguracionRequisito, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['url_vars'] = url_vars
                    data['eRequisitoContratacionConfiguracionRequisito'] = page.object_list
                    data['page'] = page
                    data['eContratacionConfiguracionRequisito'] = eContratacionConfiguracionRequisito
                    return render(request, 'adm_contratodip/viewgestionarrequisitocontratacion.html', data)
                except Exception as ex:
                    line_err = f"Error en la linea {sys.exc_info()[-1].tb_lineno}"
                    err_ = f"Ocurrio un error, {ex.__str__()}. {line_err}"
                    return HttpResponseRedirect(f"{request.path}?info={err_}")

            elif action == 'gestionarconfiguracionrequisitopago':
                try:
                    data['title'] = 'gestionar configuración requisito pago'
                    id = int(encrypt(request.GET.get('id')))
                    if id == 0:
                        raise NameError("Parametro no encontrado")
                    filtro, url_vars, search = Q(status=True,
                                                 gruporequisitopago_id=id), f'&action=gestionarconfiguracionrequisitopago&id={id}', request.GET.get(
                        's', None)
                    eGrupoRequisitoPago = GrupoRequisitoPago.objects.get(pk=id)
                    eRequisitoPagoGrupoRequisito = RequisitoPagoGrupoRequisito.objects.filter(filtro).order_by('orden')

                    paging = MiPaginador(eRequisitoPagoGrupoRequisito, 20)
                    p = 1
                    try:
                        paginasesion = 1
                        if 'paginador' in request.session:
                            paginasesion = int(request.session['paginador'])
                        if 'page' in request.GET:
                            p = int(request.GET['page'])
                        else:
                            p = paginasesion
                        try:
                            page = paging.page(p)
                        except:
                            p = 1
                        page = paging.page(p)
                    except:
                        page = paging.page(p)
                    request.session['paginador'] = p
                    data['paging'] = paging
                    data['rangospaging'] = paging.rangos_paginado(p)
                    data['url_vars'] = url_vars
                    data['eRequisitoPagoGrupoRequisito'] = page.object_list
                    data['page'] = page
                    data['eGrupoRequisitoPago'] = eGrupoRequisitoPago
                    return render(request, 'adm_contratodip/pago/viewgestionarrequisitopago.html', data)
                except Exception as ex:
                    line_err = f"Error en la linea {sys.exc_info()[-1].tb_lineno}"
                    err_ = f"Ocurrio un error, {ex.__str__()}. {line_err}"
                    return HttpResponseRedirect(f"{request.path}?info={err_}")


            elif action == 'requisitocontratacion':
                try:
                    data['title'] = 'Gestionar requisito contratación'
                    id = int(encrypt(request.GET.get('id')))
                    if id == 0:
                        raise NameError("Parametro no encontrado")
                    data['eContratoDip'] = eContratoDip = ContratoDip.objects.get(pk=id)

                    return render(request, 'adm_contratodip/viewrequisitoscontrato.html', data)
                except Exception as ex:
                    line_err = f"Error en la linea {sys.exc_info()[-1].tb_lineno}"
                    err_ = f"Ocurrio un error, {ex.__str__()}. {line_err}"
                    return HttpResponseRedirect(f"{request.path}?info={err_}")

            elif action == 'subir_requisito_contrato':
                try:
                    pk = int(request.GET.get('id', '0'))
                    id_contrato = int(request.GET.get('id_contrato', '0'))
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    data['form2'] = RequisitoContratoForm()
                    data['id'] = pk
                    data['idc'] = id_contrato
                    eContratoDip = ContratoDip.objects.get(pk=id_contrato)
                    eRequisito = Requisito.objects.get(pk=pk)
                    data['action'] = action
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'actualizar_requisito_contratacion':
                try:
                    pk = int(request.GET.get('id', '0'))
                    id_contrato = int(request.GET.get('id_contrato', '0'))
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    eContratoRequisito = ContratoRequisito.objects.get(pk=pk)
                    data['form2'] = RequisitoContratoForm(initial=model_to_dict(eContratoRequisito))
                    data['id'] = pk
                    data['idc'] = id_contrato

                    data['action'] = action
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'actualizar_fecha_caducidad_requisito':
                try:
                    pk = int(request.GET.get('id', '0'))
                    if pk == 0:
                        raise NameError("Parametro no encontrado")
                    eContratoRequisito = ContratoRequisito.objects.get(pk=pk)
                    form = RequisitoContratoForm(initial={
                        'fecha_caducidad': eContratoRequisito.fecha_caducidad
                    })
                    form.del_field('archivo')
                    data['form2'] = form
                    data['id'] = eContratoRequisito.pk
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    pass

            elif action == 'addgestposgrado':
                try:
                    data['ids'] = request.GET.get('id', None)
                    data['form2'] = GestionPosgradoForm()
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "mensaje": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'addgestorrequisitocontratacion':
                try:
                    data['id'] = request.GET.get('id', None)
                    data['form2'] = ContratacionGestionRequisitoForm()
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "mensaje": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'addgestorrequisitopago':
                try:
                    data['id'] = request.GET.get('id', None)
                    data['form2'] = RequisitoPagoDipForm()
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "mensaje": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'editgestposgrado':
                try:
                    data['id'] = id = request.GET.get('id', None)
                    id = encrypt(id)
                    filtro = Gestion.objects.get(status=True, id=id)
                    form = GestionPosgradoForm(
                        initial=model_to_dict(filtro)
                    )
                    form.edit(filtro.responsable.id)
                    if filtro.responsablesubrogante:
                        form.res_subrogante(filtro.responsablesubrogante.id)
                    data['form2'] = form
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "message": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'edit_orden_firma_acta_pago':
                try:
                    eOrdenFirmaActaPago = OrdenFirmaActaPago.objects.get(pk=request.GET.get('id'))
                    f = OrdenFirmaInformeActaPagoForm(initial=model_to_dict(eOrdenFirmaActaPago))
                    data['form2'] = f
                    data['id'] = eOrdenFirmaActaPago.pk
                    template = get_template('adm_contratodip/modal/formmodalgen.html')
                    return JsonResponse({"result": True, 'data': template.render(data)})
                except Exception as ex:
                    return JsonResponse({"result": False, "message": f"Lo sentimos, {ex.__str__()}"})

            elif action == 'historial_de_pago_contrato_posgrado':
                id = request.GET.get('id',0)
                eContratoDip = ContratoDip.objects.get(pk=id)
                data['estructura'] = eContratoDip.estructura_de_pagos_contrato()

                template = get_template('adm_contratodip/modal/modal_pagos_contrato.html')
                return JsonResponse({"result": True, 'data': template.render(data)})
        else:
            data['title'] = u'Gestión de posgrado / pago'
            search = None
            ids = None
            listado = ContratoDip.objects.filter(status=True).order_by('-id')
            url_vars = ''
            if 's' in request.GET:
                search = request.GET['s'].strip().upper()
                ss = search.split(' ')
                url_vars += '&s=' + search
                if len(ss) == 1:
                    listado = listado.filter(Q(codigocontrato__icontains=search) |
                                             Q(plantilla__descripcion__icontains=search) |
                                             Q(persona__apellido1__icontains=search) |
                                             Q(persona__apellido1__icontains=search) |
                                             Q(persona__cedula__icontains=search) |
                                             Q(persona__email__icontains=search) |
                                             Q(persona__emailinst__icontains=search) |
                                             Q(plantilla__perfil__nombre__icontains=search) |
                                             Q(descripcion__icontains=search))

                else:
                    listado = listado.filter(Q(persona__apellido1__icontains=ss[0]) &
                                             Q(persona__apellido2__icontains=ss[1])).distinct()
            elif 'id' in request.GET:
                ids = int(request.GET['id'])
                listado = listado.filter(id=ids)

            paging = MiPaginador(listado, 25)
            p = 1
            try:
                paginasesion = 1
                if 'paginador' in request.session:
                    paginasesion = int(request.session['paginador'])
                if 'page' in request.GET:
                    p = int(request.GET['page'])
                else:
                    p = paginasesion
                try:
                    page = paging.page(p)
                except:
                    p = 1
                page = paging.page(p)
            except:
                page = paging.page(p)
            request.session['paginador'] = p
            data['paging'] = paging
            data['rangospaging'] = paging.rangos_paginado(p)
            data['usuario'] = usuario
            data['page'] = page
            data['search'] = search if search else ""
            data['ids'] = ids if ids else ""
            data['listado'] = page.object_list
            data['totcount'] = listado.count()
            data['email_domain'] = EMAIL_DOMAIN
            data['fecha'] = datetime.now().date()
            data['url_vars'] = url_vars
            return render(request, 'adm_contratodip/view.html', data)


def generar_detalle_certificacion(request, contrato):
    with transaction.atomic():
        try:
            certi = contrato.certificacion
            if certi.saldo < Decimal(contrato.valortotal):
                return None
            certi.saldo -= Decimal(contrato.valortotal)
            certi.save(request)
            log(u'Se actualizo el saldo de la certificacion: %s' % certi, request, "edit")
            detalle = DetalleCertificacionPresupuestariaDip(certificado=certi, contratodip=contrato,
                                                            valor=contrato.valortotal,
                                                            fecha=datetime.now(), tipo=1, saldo=contrato.valortotal)
            detalle.save(request)
            return detalle
        except Exception as e:
            print(e)
            transaction.set_rollback(True)
        return None


def fecha_pago(cuota, contrato):
    fecha = contrato.fechainicio
    return fecha.replace(day=1) + relativedelta(months=cuota) - timedelta(days=1)
    # nextmonth = fecha.month+cuota
    # if nextmonth == 13:
    #     return fecha.replace(month=nextmonth-1, day=31)
    # if ((nextmonth-2)/12) in [1, 2, 3, 4, 5]:
    #     return fecha.replace(year=fecha.year+int(((nextmonth-2)/12)), month=1, day=1) - timedelta(days=1)
    # return fecha.replace(month=nextmonth, day=1) - timedelta(days=1)


def secuencia_memo():
    anio = anio_ejercicio().anioejercicio
    secuencia_choices = ('sgapagpos',)
    plantilla = None
    if PlantillaContratoDip.objects.values('id').filter(status=True).last():
        contrato = PlantillaContratoDip.objects.filter(status=True).last()
        if SecuenciaMemoActividadPosgrado.objects.values('id').filter(status=True, anioejercicio__anioejercicio=anio,
                                                                      tipo=contrato).exists():
            secuencia = SecuenciaMemoActividadPosgrado.objects.filter(status=True, anioejercicio__anioejercicio=anio,
                                                                      tipo=contrato)[0]
        else:
            secuencia = SecuenciaMemoActividadPosgrado(anioejercicio=anio_ejercicio(), tipo=contrato,
                                                       secuencia=secuencia_choices[0])
            secuencia.save()
    else:
        if SecuenciaMemoActividadPosgrado.objects.values('id').filter(status=True,
                                                                      anioejercicio__anioejercicio=anio).exists():
            secuencia = SecuenciaMemoActividadPosgrado.objects.filter(status=True, anioejercicio__anioejercicio=anio)[0]
        else:
            secuencia = SecuenciaMemoActividadPosgrado(anioejercicio=anio_ejercicio(), secuencia=secuencia_choices[0])
            secuencia.save()

    return secuencia
